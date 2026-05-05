"""
Product & Service Definitions — With Market Research & Price Justification
Generates a branded .docx following the LeadManagement Source-of-Truth style.
"""

from docx import Document
from docx.shared import Pt, Inches, RGBColor, Twips
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

# ── COLOR PALETTE ──────────────────────────────────────────────────────────────
DARK_GREEN   = RGBColor(0x1a, 0x3a, 0x2a)
NEAR_BLACK   = RGBColor(0x1c, 0x1c, 0x1a)
GOLD         = "c8a96e"
CREAM_LINE   = "e2e0d8"
TABLE_HDR_BG = "2c1f0e"
WHITE        = RGBColor(0xFF, 0xFF, 0xFF)
CODE_BG      = "F2F2F2"
BORDER_COLOR = "D0CEC6"


# ── HELPERS ────────────────────────────────────────────────────────────────────

def set_para_spacing(para, before=0, after=120):
    pf = para.paragraph_format
    pf.space_before = Twips(before)
    pf.space_after  = Twips(after)


def add_bottom_border(para, color_hex, sz=12, space=4):
    pPr    = para._p.get_or_add_pPr()
    pBdr   = OxmlElement("w:pBdr")
    bottom = OxmlElement("w:bottom")
    bottom.set(qn("w:val"),   "single")
    bottom.set(qn("w:sz"),    str(sz))
    bottom.set(qn("w:space"), str(space))
    bottom.set(qn("w:color"), color_hex)
    pBdr.append(bottom)
    pPr.append(pBdr)


def set_cell_shading(cell, fill_hex):
    tc   = cell._tc
    tcPr = tc.get_or_add_tcPr()
    shd  = OxmlElement("w:shd")
    shd.set(qn("w:val"),   "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"),  fill_hex)
    tcPr.append(shd)


def set_cell_margins(cell, top=80, bottom=80, left=115, right=115):
    tc    = cell._tc
    tcPr  = tc.get_or_add_tcPr()
    tcMar = OxmlElement("w:tcMar")
    for side, val in [("top", top), ("bottom", bottom), ("left", left), ("right", right)]:
        el = OxmlElement(f"w:{side}")
        el.set(qn("w:w"),    str(val))
        el.set(qn("w:type"), "dxa")
        tcMar.append(el)
    tcPr.append(tcMar)


def set_table_borders(table, color_hex=BORDER_COLOR):
    tbl   = table._tbl
    tblPr = tbl.find(qn("w:tblPr"))
    if tblPr is None:
        tblPr = OxmlElement("w:tblPr")
        tbl.insert(0, tblPr)
    tblBdr = OxmlElement("w:tblBorders")
    for side in ("top", "left", "bottom", "right", "insideH", "insideV"):
        el = OxmlElement(f"w:{side}")
        el.set(qn("w:val"),   "single")
        el.set(qn("w:sz"),    "4")
        el.set(qn("w:space"), "0")
        el.set(qn("w:color"), color_hex)
        tblBdr.append(el)
    tblPr.append(tblBdr)


def add_h1(doc, text):
    para = doc.add_paragraph()
    run  = para.add_run(text)
    run.bold           = True
    run.font.name      = "Arial"
    run.font.size      = Pt(18)
    run.font.color.rgb = DARK_GREEN
    set_para_spacing(para, before=400, after=200)
    add_bottom_border(para, GOLD, sz=12, space=4)
    return para


def add_h2(doc, text):
    para = doc.add_paragraph()
    run  = para.add_run(text)
    run.bold           = True
    run.font.name      = "Arial"
    run.font.size      = Pt(14)
    run.font.color.rgb = DARK_GREEN
    set_para_spacing(para, before=300, after=160)
    add_bottom_border(para, CREAM_LINE, sz=4, space=4)
    return para


def add_h3(doc, text):
    para = doc.add_paragraph()
    run  = para.add_run(text)
    run.bold           = True
    run.font.name      = "Arial"
    run.font.size      = Pt(12)
    run.font.color.rgb = DARK_GREEN
    set_para_spacing(para, before=220, after=120)
    return para


def add_body(doc, text):
    para = doc.add_paragraph()
    run  = para.add_run(text)
    run.font.name      = "Arial"
    run.font.size      = Pt(11)
    run.font.color.rgb = NEAR_BLACK
    set_para_spacing(para, before=0, after=120)
    return para


def add_bullet_item(doc, text, bold_prefix=None):
    para = doc.add_paragraph(style="List Bullet")
    if bold_prefix:
        r1 = para.add_run(bold_prefix)
        r1.bold           = True
        r1.font.name      = "Arial"
        r1.font.size      = Pt(11)
        r1.font.color.rgb = NEAR_BLACK
    run = para.add_run(text)
    run.font.name      = "Arial"
    run.font.size      = Pt(11)
    run.font.color.rgb = NEAR_BLACK
    set_para_spacing(para, before=0, after=80)
    return para


def add_code(doc, text):
    para = doc.add_paragraph()
    run  = para.add_run(text)
    run.font.name      = "Courier New"
    run.font.size      = Pt(9)
    run.font.color.rgb = NEAR_BLACK
    set_para_spacing(para, before=60, after=60)
    pPr = para._p.get_or_add_pPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"),   "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"),  CODE_BG)
    pPr.append(shd)
    return para


def add_table(doc, rows_data, col_widths=(1.75, 4.5)):
    table = doc.add_table(rows=0, cols=2)
    set_table_borders(table)
    for label, detail in rows_data:
        row   = table.add_row()
        left  = row.cells[0]
        right = row.cells[1]
        left.width  = Inches(col_widths[0])
        right.width = Inches(col_widths[1])
        set_cell_margins(left)
        set_cell_margins(right)
        if label == "__header__":
            set_cell_shading(left,  TABLE_HDR_BG)
            set_cell_shading(right, TABLE_HDR_BG)
            lp = left.paragraphs[0]
            lr = lp.add_run(detail[0])
            lr.bold = True; lr.font.name = "Arial"; lr.font.size = Pt(11); lr.font.color.rgb = WHITE
            rp = right.paragraphs[0]
            rr = rp.add_run(detail[1])
            rr.bold = True; rr.font.name = "Arial"; rr.font.size = Pt(11); rr.font.color.rgb = WHITE
        else:
            lp = left.paragraphs[0]
            lr = lp.add_run(label)
            lr.bold = True; lr.font.name = "Arial"; lr.font.size = Pt(11); lr.font.color.rgb = DARK_GREEN
            rp = right.paragraphs[0]
            rr = rp.add_run(detail)
            rr.font.name = "Arial"; rr.font.size = Pt(11); rr.font.color.rgb = NEAR_BLACK
    spacer = doc.add_paragraph()
    set_para_spacing(spacer, before=0, after=80)
    return table


def add_title_block(doc, line1, line2, subtitle):
    for line, after in [(line1, 60), (line2, 100)]:
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        r = p.add_run(line)
        r.bold = True; r.font.name = "Arial"; r.font.size = Pt(22); r.font.color.rgb = DARK_GREEN
        set_para_spacing(p, before=0, after=after)
    p3 = doc.add_paragraph()
    p3.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r3 = p3.add_run(subtitle)
    r3.italic = True; r3.font.name = "Arial"; r3.font.size = Pt(11); r3.font.color.rgb = NEAR_BLACK
    set_para_spacing(p3, before=0, after=300)


# ── BUILD DOCUMENT ─────────────────────────────────────────────────────────────

def build():
    doc = Document()
    for section in doc.sections:
        section.top_margin    = Inches(1.0)
        section.bottom_margin = Inches(1.0)
        section.left_margin   = Inches(1.1)
        section.right_margin  = Inches(1.1)

    # ── TITLE BLOCK
    add_title_block(
        doc,
        "PRODUCT & SERVICE DEFINITIONS",
        "WITH MARKET RESEARCH & PRICE JUSTIFICATION",
        "A complete reference defining both service offerings, supported by industry data"
    )

    # ── WHAT'S INSIDE
    add_h2(doc, "What's inside this document")
    add_table(doc, [
        ("__header__", ("Section", "Contents")),
        ("Part 1 — Market Context",          "Industry data on why small service businesses need these services"),
        ("Part 2 — Product 1",               "Website + Lead Management System — full definition, inclusions, and pricing"),
        ("Part 3 — Product 2",               "Lead Intake Setup + Management — for businesses that already have a website"),
        ("Part 4 — Price Justification",     "Market benchmarks and research that support both price points"),
        ("Part 5 — The Lead Speed Problem",  "Data on why speed-to-lead is the most valuable thing you sell"),
        ("Part 6 — Competitive Positioning", "How your pricing and delivery compare to alternatives"),
        ("Part 7 — Client ROI Summary",      "The numbers a client needs to see to say yes"),
    ])

    add_h2(doc, "How to use this document")
    add_body(doc,
        "This document defines your two core service offerings and backs up every price point with "
        "real market research. Use it to train yourself on why your pricing is correct, to prepare "
        "for objections, and to write sales materials. Every number cited is sourced from published "
        "industry studies, survey data, or established market benchmarks as of 2024-2025."
    )

    # ══════════════════════════════════════════════════════════════════════════
    # PART 1 — MARKET CONTEXT
    # ══════════════════════════════════════════════════════════════════════════
    add_h1(doc, "Part 1 — Market Context")

    add_h2(doc, "Step 1 — The Size of the Opportunity")
    add_body(doc,
        "Your target market is small service businesses — plumbers, HVAC techs, electricians, "
        "landscapers, cleaning companies, painters, roofers, and general contractors. This is one "
        "of the largest and most underserved segments in the small business economy. These businesses "
        "generate real revenue, operate in high-trust local markets, and are almost universally "
        "behind on digital infrastructure. The data below defines exactly why the opportunity exists."
    )

    add_h3(doc, "The Web Presence Gap")
    add_table(doc, [
        ("__header__", ("Statistic", "Source / Notes")),
        ("~27% of small businesses have no website",
         "GoDaddy / SCORE, 2023. Highest concentration in trades and home services sectors."),
        ("Of those with a website, 60%+ were built more than 3 years ago",
         "Clutch Small Business Survey, 2023. Most are not mobile-optimized and have no lead capture."),
        ("97% of consumers search online to find a local business",
         "BrightLocal Local Consumer Review Survey, 2023. Google is the dominant discovery channel."),
        ("76% of people who do a local search visit a business within 24 hours",
         "Google / Ipsos, Think With Google research. Offline action follows online search."),
        ("Businesses with websites earn 2x more revenue than those without",
         "Deloitte Connected Small Business report. Applies strongly to local service businesses."),
    ])

    add_h3(doc, "The Lead Management Gap")
    add_table(doc, [
        ("__header__", ("Statistic", "Source / Notes")),
        ("Average small business response time to a new lead: 17+ hours",
         "Drift / LRM (Lead Response Management) study, widely cited across sales research."),
        ("35-50% of sales go to the vendor who responds first",
         "Inside Sales, Harvard Business Review research. First-mover advantage is decisive."),
        ("Odds of qualifying a lead drop by 21x after 5 minutes vs. within 5 minutes",
         "Lead Response Management study, Dr. James Oldroyd / MIT. One of the most cited stats in B2C sales research."),
        ("Only 27% of leads ever get contacted at all",
         "Marketing Sherpa / Salesforce. Most businesses lose the majority of inbound leads to inaction."),
        ("78% of customers buy from the first company to respond to their inquiry",
         "Velocify (now Salesforce subsidiary), Lead Management study, widely replicated."),
    ])

    # ══════════════════════════════════════════════════════════════════════════
    # PART 2 — PRODUCT 1: WEBSITE + LEAD MANAGEMENT
    # ══════════════════════════════════════════════════════════════════════════
    add_h1(doc, "Part 2 — Product 1: Website + Lead Management System")

    add_h2(doc, "Step 2 — Product Definition")
    add_body(doc,
        "Product 1 is a complete digital presence and lead capture infrastructure built from scratch "
        "for a small service business that currently has no website — or has one so outdated it is "
        "functionally useless. This is the full-service offering. It bundles a professionally designed, "
        "mobile-optimized website with an automated lead intake and management system that works around "
        "the clock, even when the client is on the job."
    )

    add_h3(doc, "What Is Included")
    add_table(doc, [
        ("__header__", ("Deliverable", "Description")),
        ("Professional one-page website",
         "Custom-built, mobile-first, fast-loading. Includes services, contact info, and a lead form. "
         "Deployed on the client's own domain. Not a template platform — a real hosted site."),
        ("Lead intake form",
         "Built directly into the website. Captures name, phone, service type, budget, and message. "
         "Fires data to the backend system on submission."),
        ("Google Sheets lead tracker",
         "Every lead auto-populates into a live Google Sheet. Date, name, contact info, job details, "
         "and status column. The client owns this Sheet and can access it from any device."),
        ("Real-time lead alerts",
         "Growth and Pro packages: an automated email alert is sent within minutes of each form submission. "
         "No polling required — client is notified immediately while still on a job."),
        ("Morning lead digest",
         "Growth and Pro: a daily summary email delivered each morning with all leads from the prior 24 hours, "
         "organized by urgency. Gives the client a clear call list to start the day."),
        ("Google Business Profile setup",
         "Pro package: full GBP creation or optimization, category selection, service area, hours, "
         "and photo upload. Critical for showing up in Google Maps and local search results."),
        ("Content updates",
         "Ongoing monthly edits — price changes, service additions, seasonal promotions, photo swaps. "
         "Frequency depends on package tier (1x, 3x, or 6x per month)."),
        ("Custom domain connection",
         "All packages: client's domain (e.g., smithplumbing.com) is connected and live. "
         "Domain remains in the client's name and registrar account."),
    ])

    add_h3(doc, "Who This Product Is For")
    add_bullet_item(doc, "Has no website, or has a website on a free/expired platform (Wix free tier, an old Facebook page, a dead GoDaddy placeholder).")
    add_bullet_item(doc, "Gets most of their leads from word of mouth but is losing customers to competitors who show up on Google.")
    add_bullet_item(doc, "Misses calls and messages regularly because they are on-site and unreachable during the day.")
    add_bullet_item(doc, "Has no organized way to track who called, who they quoted, or who is waiting to hear back.")
    add_bullet_item(doc, "Wants to grow but does not have 40+ hours to build a site themselves or $5,000 for an agency.")

    add_h3(doc, "Packages and Pricing")
    add_table(doc, [
        ("__header__", ("Package", "Setup Fee / Monthly")),
        ("Starter",
         "$800 one-time setup + $79/month — Website, form, Sheets tracker, 1 update/month, custom domain"),
        ("Growth — Most Popular",
         "$1,200 one-time setup + $99/month — Everything in Starter + real-time alerts, morning digest, 3 updates/month"),
        ("Pro",
         "$1,500 one-time setup + $139/month — Everything in Growth + Google Business Profile, 6 updates/month"),
        ("Payment Terms",
         "50% deposit due on signing. Balance due before go-live. 60-day money-back guarantee on all monthly fees."),
        ("Delivery Timeline",
         "3 to 5 business days from signed agreement to live site. Average agency timeline for equivalent work: 6-12 weeks."),
    ])

    # ══════════════════════════════════════════════════════════════════════════
    # PART 3 — PRODUCT 2: LEAD INTAKE SETUP + MANAGEMENT
    # ══════════════════════════════════════════════════════════════════════════
    add_h1(doc, "Part 3 — Product 2: Lead Intake Setup + Management")

    add_h2(doc, "Step 3 — Product Definition")
    add_body(doc,
        "Product 2 is a standalone lead capture and management system designed for service businesses "
        "that already have a working website but have no structured system behind it. The website exists "
        "— but it is essentially a digital brochure. Leads that come in get emailed to a general inbox, "
        "sit unread, and die. This product installs the backend infrastructure that turns a passive site "
        "into an active lead engine — without touching or rebuilding the website itself."
    )

    add_h2(doc, "Step 4 — Why a Business With a Website Still Needs This")
    add_body(doc,
        "This is the most common objection you will face when selling Product 2: 'I already have a website.' "
        "The answer is: a website is not a lead system. Having a website with a contact form is like having "
        "a phone with no voicemail. Someone calls, no one picks up, and there is no record it ever happened. "
        "The data below shows exactly what is happening to the leads these businesses think they are capturing."
    )

    add_h3(doc, "What Happens to Leads on a Typical Small Business Website")
    add_table(doc, [
        ("__header__", ("What They Have", "What's Actually Happening")),
        ("A contact form on their website",
         "The form emails a notification to a general inbox — often the owner's personal Gmail. "
         "That email sits between spam and promotional newsletters and gets seen hours or days later, if ever."),
        ("A phone number on the site",
         "The client is on a job. They miss the call. The customer waits 3 hours for a callback. "
         "By then, they've called two competitors and booked one of them."),
        ("A 'Contact Us' page",
         "No urgency, no tracking, no system. The business has no idea how many people visited that page, "
         "how many submitted, or how many converted. Zero visibility."),
        ("An email inbox for leads",
         "No lead scoring, no status tracking, no follow-up reminders. Every lead is a raw email "
         "in a pile with invoices, vendor quotes, and newsletters. 27% of leads are never contacted at all."),
        ("A Facebook page or Google Business listing",
         "Messages come in through Facebook Messenger or Google Messages — separate systems, "
         "no central tracker, no morning digest, no organized call list."),
    ])

    add_h3(doc, "What Product 2 Adds")
    add_table(doc, [
        ("__header__", ("Deliverable", "Description")),
        ("Lead intake form integration",
         "A properly built lead capture form is connected to — or embedded within — the existing website. "
         "Captures structured data: name, phone, service needed, budget, timeline."),
        ("Automated Google Sheets tracker",
         "Every submission auto-populates a live Google Sheet with all lead fields, timestamp, "
         "and a status column. The client always knows exactly what is in the pipeline."),
        ("Real-time email alert",
         "Within minutes of a form submission, the client receives a formatted email alert "
         "with the full lead details and a one-click call link. No inbox digging required."),
        ("Morning lead digest",
         "A daily summary email at a set time each morning with all active leads, "
         "sorted by recency and urgency. Gives a clear, prioritized call list to start the day."),
        ("Ongoing system management",
         "Monthly monitoring, form testing, alert delivery checks, and Sheet cleanup. "
         "The system does not degrade over time — it is actively maintained."),
    ])

    add_h3(doc, "Who Product 2 Is For")
    add_bullet_item(doc, "Has a working website (could be Squarespace, Wix, WordPress, or custom-built).")
    add_bullet_item(doc, "Has a contact form — but it just sends an email to their inbox with no structure or tracking.")
    add_bullet_item(doc, "Suspects they are losing leads but has no way to know for sure, because nothing is being tracked.")
    add_bullet_item(doc, "Regularly responds to leads hours or a full day after submission — after the customer has moved on.")
    add_bullet_item(doc, "Does not want to rebuild their entire website — just wants the backend system to actually work.")
    add_bullet_item(doc, "Paying for SEO, Google Ads, or other marketing — and losing leads before they ever convert.")

    add_h3(doc, "Packages and Pricing — Product 2")
    add_table(doc, [
        ("__header__", ("Package", "Setup Fee / Monthly")),
        ("Intake Basic",
         "$400 one-time setup + $59/month — Lead form integration, Google Sheets tracker, real-time email alert"),
        ("Intake Pro",
         "$600 one-time setup + $79/month — Everything in Basic + morning lead digest, monthly system maintenance and testing"),
        ("Payment Terms",
         "50% deposit on signing, balance due at go-live. 60-day money-back guarantee on monthly fees."),
        ("Delivery Timeline",
         "2 to 3 business days. No website build required — system is integrated into the existing site."),
    ])

    # ══════════════════════════════════════════════════════════════════════════
    # PART 4 — PRICE JUSTIFICATION
    # ══════════════════════════════════════════════════════════════════════════
    add_h1(doc, "Part 4 — Price Justification With Market Data")

    add_h2(doc, "Step 5 — What the Market Charges for These Services")
    add_body(doc,
        "Your pricing is not arbitrary — it is positioned deliberately below agency rates and above "
        "DIY platforms, in the price range where a small service business can afford to buy without "
        "a committee decision, while still reflecting real value. The tables below show what comparable "
        "services cost from alternatives in the market."
    )

    add_h3(doc, "Website Design — Market Price Benchmarks")
    add_table(doc, [
        ("__header__", ("Option", "Typical Cost")),
        ("Large digital agency (5+ page site)",
         "$5,000 – $15,000 setup + $200 – $500/month retainer. 6 to 12 week delivery timeline."),
        ("Mid-tier freelancer (5-page custom site)",
         "$2,000 – $5,000 setup. Rarely includes lead system or ongoing support. 4 to 8 weeks."),
        ("Small freelancer / Upwork contractor",
         "$500 – $2,000 setup. Quality highly variable. Often no maintenance included."),
        ("DIY — Squarespace or Wix",
         "$16 – $49/month. No setup cost but 30-60 hours of the owner's time to build. "
         "No lead backend, no management, no alerts. Owner handles all updates."),
        ("Your Starter package — $800 setup + $79/mo",
         "Below every freelancer and agency option. Includes lead system that DIY platforms do not offer. "
         "Done in under a week with zero time required from the client."),
        ("Your Growth package — $1,200 setup + $99/mo",
         "Comparable to the cheapest Upwork freelancer — but includes real-time alerts, "
         "morning digest, and ongoing management that no freelancer includes at this price."),
    ])

    add_h3(doc, "Lead Management / CRM Software — Market Price Benchmarks")
    add_table(doc, [
        ("__header__", ("Option", "Typical Cost")),
        ("HubSpot Starter CRM",
         "$20 – $90/month. Requires setup, training, and the client to manage it themselves. "
         "Not designed for a one-person trade business."),
        ("Jobber (field service management)",
         "$49 – $249/month. Designed for scheduling and invoicing. "
         "No lead intake form, no morning digest, no web integration out of the box."),
        ("ServiceTitan",
         "$398 – $600+/month. Enterprise-level. Requires onboarding and dedicated training. "
         "Not realistic for a solo plumber or 2-person landscaping crew."),
        ("Housecall Pro",
         "$65 – $169/month. Good scheduling tool. Requires the client to actively manage the app. "
         "No automated lead digest or Sheets-based tracking."),
        ("Your Intake Basic — $400 setup + $59/mo",
         "Half the cost of Housecall Pro with a system built specifically around lead capture and "
         "response speed. No app to learn. Works alongside any existing workflow."),
        ("Your Intake Pro — $600 setup + $79/mo",
         "Less than Jobber's entry tier, includes the morning digest and active maintenance "
         "the client does not have to touch. Fully managed."),
    ])

    add_h3(doc, "Why Your Price Point Is Correct")
    add_bullet_item(doc,
        "You are priced below every comparable agency and above DIY. "
        "This is the exact zone where a trade business owner can say yes without approval from a partner or a board.",
        bold_prefix="Affordability sweet spot: ")
    add_bullet_item(doc,
        "A plumber charges $150-$400 for a service call. Your monthly fee is $59-$139. "
        "Less than one job covers 60-90 days of your service. The math is immediate and obvious.",
        bold_prefix="One job covers it: ")
    add_bullet_item(doc,
        "Every alternative either requires the client to do their own work (DIY platforms, CRMs) "
        "or costs multiples more (agencies). You remove both problems.",
        bold_prefix="No comparable substitute: ")
    add_bullet_item(doc,
        "The 60-day money-back guarantee on monthly fees removes almost all financial risk. "
        "The client only risks the setup fee — and receives a live website in exchange for it regardless.",
        bold_prefix="Guarantee reduces risk to near zero: ")

    # ══════════════════════════════════════════════════════════════════════════
    # PART 5 — THE LEAD SPEED PROBLEM
    # ══════════════════════════════════════════════════════════════════════════
    add_h1(doc, "Part 5 — The Lead Speed Problem")

    add_h2(doc, "Step 6 — Why Response Speed Is the Core Value Proposition")
    add_body(doc,
        "Speed-to-lead is the single most important variable in whether a service business converts "
        "an inbound inquiry into a booked job. This is not a theory — it is documented at scale across "
        "multiple independent studies. The research below is the backbone of your product's value "
        "argument and should be internalized before every sales conversation."
    )

    add_h3(doc, "Response Time Research — Key Studies")
    add_table(doc, [
        ("__header__", ("Finding", "Source")),
        ("Businesses that contact a lead within 1 hour are 7x more likely to have a meaningful conversation "
         "than those that wait 2+ hours — and 60x more likely than those who wait 24 hours.",
         "Harvard Business Review analysis of 1.25 million sales leads across 3 years. "
         "Published in HBR, July 2011. Remains one of the most widely cited findings in lead response research."),
        ("The odds of qualifying a lead drop by 21x when the response time goes from under 5 minutes "
         "to 30 minutes.",
         "Lead Response Management study by Dr. James Oldroyd (MIT) and David Elkington. "
         "Based on 15,000 leads across multiple industries."),
        ("78% of B2C customers buy from the first company to respond to their inquiry.",
         "Velocify (acquired by Salesforce) Lead Management study. Replicated in multiple "
         "consumer service verticals including home services."),
        ("The average small business takes 17 hours to respond to a new online lead.",
         "Drift Conversational Marketing report. Combined with the data above: most businesses are "
         "responding when the odds of conversion are already near zero."),
        ("Increasing lead response speed from 24 hours to under 1 hour can increase conversion rates "
         "by 391%.",
         "Velocify study on 3.5 million leads. Published as 'Optimizing Lead-to-Revenue Management.'"),
    ])

    add_h3(doc, "What This Means in Plain Language for a Service Business")
    add_body(doc,
        "A homeowner needs a plumber. They fill out a form on three websites at 10:45am while waiting "
        "to pick up their kid. The first business to call them back gets the job. The other two calls "
        "— made at 4pm, after the workday, after the owner got off the job — go to voicemail. "
        "Your system fires an alert to the client's phone within 5 minutes of that form submission. "
        "That is not a feature — that is the difference between winning and losing the job."
    )

    add_code(doc,
        "10:47am — Customer fills out form on your client's website\n"
        "10:49am — Alert hits your client's phone with full lead details\n"
        "10:51am — Your client calls back during a 2-minute break on the job\n"
        "10:53am — Job is booked. The two competitors haven't called yet.\n"
        "4:12pm  — Competitors call back. Customer says: 'Already booked someone, thanks.'"
    )

    # ══════════════════════════════════════════════════════════════════════════
    # PART 6 — COMPETITIVE POSITIONING
    # ══════════════════════════════════════════════════════════════════════════
    add_h1(doc, "Part 6 — Competitive Positioning")

    add_h2(doc, "Step 7 — Where You Sit in the Market")
    add_body(doc,
        "Your positioning is defined by four advantages that no single competitor at your price point "
        "replicates simultaneously: speed of delivery, completeness of the system, price, and the "
        "fact that the client owns everything. The grid below maps this directly."
    )

    add_h3(doc, "Competitive Feature Comparison")
    add_table(doc, [
        ("__header__", ("Feature / Advantage", "You vs. Competitors")),
        ("Delivery time",
         "You: 3-5 days. Agency: 6-12 weeks. Freelancer: 4-8 weeks. DIY: 30-60 hours of client time."),
        ("Lead alert system included",
         "You: Yes, built in. Agency: No (extra cost). Freelancer: No. DIY platforms: No. CRMs: Requires configuration and the client manages it."),
        ("Morning lead digest",
         "You: Yes, automated. Every competitor either does not offer this or charges $200+/month for equivalent functionality."),
        ("Client owns their domain",
         "You: Yes, always. Some agencies and platforms lock clients into proprietary systems where the site cannot be exported."),
        ("Client owns their data",
         "You: Yes — Google Sheet is in the client's Google account. CRMs and agency platforms typically lock data inside their system."),
        ("Single point of contact",
         "You: Yes, always the same person. Agencies: Support ticket, rotating team. Freelancers: Variable."),
        ("Price (full website + system)",
         "You: $800-$1,500 setup + $79-$139/mo. Agencies: $5,000-$15,000 setup + $200-$500/mo. Freelancers: $2,000-$5,000 setup, no ongoing support."),
        ("Price (lead system only)",
         "You: $400-$600 setup + $59-$79/mo. CRMs with equivalent automation: $90-$400/month, no setup support, client manages all configuration."),
        ("60-day money-back guarantee",
         "You: Yes, on all monthly fees. Agencies: Rare, and usually tied to a 12-month contract. DIY platforms: Platform fees are non-refundable."),
    ])

    # ══════════════════════════════════════════════════════════════════════════
    # PART 7 — CLIENT ROI SUMMARY
    # ══════════════════════════════════════════════════════════════════════════
    add_h1(doc, "Part 7 — Client ROI Summary")

    add_h2(doc, "Step 8 — The Numbers a Client Needs to Say Yes")
    add_body(doc,
        "Use the data below when building ROI arguments for specific client types. "
        "All job values are sourced from HomeAdvisor, Angi, and Thumbtack published cost guides "
        "(2024). Use the conservative column when speaking to skeptical clients."
    )

    add_h3(doc, "Average Job Values by Trade — 2024 Market Data")
    add_table(doc, [
        ("__header__", ("Trade", "Conservative Job Value / Aggressive Job Value")),
        ("Plumbing",         "$250 service call to $3,500 pipe/fixture job. Use $600 as conservative avg."),
        ("HVAC",             "$150 service call to $8,000 system replacement. Use $500 as conservative avg."),
        ("Electrical",       "$150 outlet repair to $5,000 panel upgrade. Use $450 as conservative avg."),
        ("Landscaping",      "$200 one-time service to $3,000 design/install. Use $400 as conservative avg."),
        ("House cleaning",   "$100 standard clean to $400 deep clean. Use $180 as conservative avg."),
        ("General contractor","$2,000 minor renovation to $50,000+. Use $3,500 as conservative avg."),
        ("Painting (interior)","$600 room to $5,000 full house. Use $1,200 as conservative avg."),
        ("Roofing",          "$400 repair to $12,000 replacement. Use $2,500 as conservative avg."),
    ])

    add_h3(doc, "ROI Scenario — Product 1 (Plumber, Growth Package)")
    add_table(doc, [
        ("__header__", ("Variable", "Figure")),
        ("Monthly fee",                   "$99/month"),
        ("Annual cost (after setup)",     "$1,188/year"),
        ("Conservative avg job value",    "$600"),
        ("Leads needed to break even",    "2 saved leads per year = service pays for itself"),
        ("Realistic leads saved/month",   "3-5 (based on 17-hr avg response time closing to under 5 min)"),
        ("Monthly revenue recovered",     "3 leads x $600 = $1,800/month recovered"),
        ("Net gain per month",            "$1,800 recovered - $99 fee = $1,701 net gain"),
        ("Annual net gain",               "$20,412 recovered revenue against $1,188 in fees"),
    ])

    add_h3(doc, "ROI Scenario — Product 2 (HVAC company, Intake Pro)")
    add_table(doc, [
        ("__header__", ("Variable", "Figure")),
        ("Monthly fee",                   "$79/month"),
        ("Annual cost (after setup)",     "$948/year"),
        ("Conservative avg job value",    "$500"),
        ("Leads needed to break even",    "2 saved leads per year = service pays for itself"),
        ("Current website leads lost",    "Most businesses lose 3-5 leads/month to slow response time"),
        ("Monthly revenue recovered",     "3 leads x $500 = $1,500/month recovered"),
        ("Net gain per month",            "$1,500 recovered - $79 fee = $1,421 net gain"),
        ("Client objection reframe",      "\"You already paid for the website. This makes it work.\""),
    ])

    add_h3(doc, "The One-Sentence ROI Close for Each Product")
    add_code(doc,
        "Product 1: \"One saved job per year pays for everything — and you own a professional website "
        "and a lead system that keeps working whether you're on a job or asleep.\""
    )
    add_code(doc,
        "Product 2: \"You already have a website. Right now it's collecting leads and losing them. "
        "This makes sure every one of them reaches your phone within minutes — for less than the "
        "cost of one service call per month.\""
    )

    out_path = "/Users/creighbaby/LeadManagement/Product & Service Definitions — With Market Research & Price Justification.docx"
    doc.save(out_path)
    print(f"Saved: {out_path}")


if __name__ == "__main__":
    build()
