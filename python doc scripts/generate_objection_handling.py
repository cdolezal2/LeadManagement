"""
Cold Call Objection Handling — Complete Reference Guide
Generates a branded .docx following the LeadManagement Source-of-Truth style.
"""

from docx import Document
from docx.shared import Pt, Inches, RGBColor, Twips
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

DARK_GREEN   = RGBColor(0x1a, 0x3a, 0x2a)
NEAR_BLACK   = RGBColor(0x1c, 0x1c, 0x1a)
GOLD         = "c8a96e"
CREAM_LINE   = "e2e0d8"
TABLE_HDR_BG = "2c1f0e"
WHITE        = RGBColor(0xFF, 0xFF, 0xFF)
CODE_BG      = "F2F2F2"
BORDER_COLOR = "D0CEC6"


def set_para_spacing(para, before=0, after=120):
    pf = para.paragraph_format
    pf.space_before = Twips(before)
    pf.space_after  = Twips(after)


def add_bottom_border(para, color_hex, sz=12, space=4):
    pPr    = para._p.get_or_add_pPr()
    pBdr   = OxmlElement("w:pBdr")
    bottom = OxmlElement("w:bottom")
    bottom.set(qn("w:val"),   "single")
    bottom.set(qn("w:sz"),    str(sz))
    bottom.set(qn("w:space"), str(space))
    bottom.set(qn("w:color"), color_hex)
    pBdr.append(bottom)
    pPr.append(pBdr)


def set_cell_shading(cell, fill_hex):
    tc   = cell._tc
    tcPr = tc.get_or_add_tcPr()
    shd  = OxmlElement("w:shd")
    shd.set(qn("w:val"),   "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"),  fill_hex)
    tcPr.append(shd)


def set_cell_margins(cell, top=80, bottom=80, left=115, right=115):
    tc    = cell._tc
    tcPr  = tc.get_or_add_tcPr()
    tcMar = OxmlElement("w:tcMar")
    for side, val in [("top", top), ("bottom", bottom), ("left", left), ("right", right)]:
        el = OxmlElement(f"w:{side}")
        el.set(qn("w:w"),    str(val))
        el.set(qn("w:type"), "dxa")
        tcMar.append(el)
    tcPr.append(tcMar)


def set_table_borders(table, color_hex=BORDER_COLOR):
    tbl   = table._tbl
    tblPr = tbl.find(qn("w:tblPr"))
    if tblPr is None:
        tblPr = OxmlElement("w:tblPr")
        tbl.insert(0, tblPr)
    tblBdr = OxmlElement("w:tblBorders")
    for side in ("top", "left", "bottom", "right", "insideH", "insideV"):
        el = OxmlElement(f"w:{side}")
        el.set(qn("w:val"),   "single")
        el.set(qn("w:sz"),    "4")
        el.set(qn("w:space"), "0")
        el.set(qn("w:color"), color_hex)
        tblBdr.append(el)
    tblPr.append(tblBdr)


def add_h1(doc, text):
    para = doc.add_paragraph()
    run  = para.add_run(text)
    run.bold = True; run.font.name = "Arial"; run.font.size = Pt(18); run.font.color.rgb = DARK_GREEN
    set_para_spacing(para, before=400, after=200)
    add_bottom_border(para, GOLD, sz=12, space=4)
    return para


def add_h2(doc, text):
    para = doc.add_paragraph()
    run  = para.add_run(text)
    run.bold = True; run.font.name = "Arial"; run.font.size = Pt(14); run.font.color.rgb = DARK_GREEN
    set_para_spacing(para, before=300, after=160)
    add_bottom_border(para, CREAM_LINE, sz=4, space=4)
    return para


def add_h3(doc, text):
    para = doc.add_paragraph()
    run  = para.add_run(text)
    run.bold = True; run.font.name = "Arial"; run.font.size = Pt(12); run.font.color.rgb = DARK_GREEN
    set_para_spacing(para, before=220, after=120)
    return para


def add_body(doc, text):
    para = doc.add_paragraph()
    run  = para.add_run(text)
    run.font.name = "Arial"; run.font.size = Pt(11); run.font.color.rgb = NEAR_BLACK
    set_para_spacing(para, before=0, after=120)
    return para


def add_bullet_item(doc, text, bold_prefix=None):
    para = doc.add_paragraph(style="List Bullet")
    if bold_prefix:
        r1 = para.add_run(bold_prefix)
        r1.bold = True; r1.font.name = "Arial"; r1.font.size = Pt(11); r1.font.color.rgb = NEAR_BLACK
    run = para.add_run(text)
    run.font.name = "Arial"; run.font.size = Pt(11); run.font.color.rgb = NEAR_BLACK
    set_para_spacing(para, before=0, after=80)
    return para


def add_code(doc, text):
    para = doc.add_paragraph()
    run  = para.add_run(text)
    run.font.name = "Courier New"; run.font.size = Pt(9); run.font.color.rgb = NEAR_BLACK
    set_para_spacing(para, before=60, after=60)
    pPr = para._p.get_or_add_pPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"), "clear"); shd.set(qn("w:color"), "auto"); shd.set(qn("w:fill"), CODE_BG)
    pPr.append(shd)
    return para


def add_table(doc, rows_data, col_widths=(1.75, 4.5)):
    table = doc.add_table(rows=0, cols=2)
    set_table_borders(table)
    for label, detail in rows_data:
        row   = table.add_row()
        left  = row.cells[0]
        right = row.cells[1]
        left.width  = Inches(col_widths[0])
        right.width = Inches(col_widths[1])
        set_cell_margins(left); set_cell_margins(right)
        if label == "__header__":
            set_cell_shading(left, TABLE_HDR_BG); set_cell_shading(right, TABLE_HDR_BG)
            lp = left.paragraphs[0]; lr = lp.add_run(detail[0])
            lr.bold = True; lr.font.name = "Arial"; lr.font.size = Pt(11); lr.font.color.rgb = WHITE
            rp = right.paragraphs[0]; rr = rp.add_run(detail[1])
            rr.bold = True; rr.font.name = "Arial"; rr.font.size = Pt(11); rr.font.color.rgb = WHITE
        else:
            lp = left.paragraphs[0]; lr = lp.add_run(label)
            lr.bold = True; lr.font.name = "Arial"; lr.font.size = Pt(11); lr.font.color.rgb = DARK_GREEN
            rp = right.paragraphs[0]; rr = rp.add_run(detail)
            rr.font.name = "Arial"; rr.font.size = Pt(11); rr.font.color.rgb = NEAR_BLACK
    spacer = doc.add_paragraph()
    set_para_spacing(spacer, before=0, after=80)
    return table


def add_title_block(doc, line1, line2, subtitle):
    for line, after in [(line1, 60), (line2, 100)]:
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        r = p.add_run(line)
        r.bold = True; r.font.name = "Arial"; r.font.size = Pt(22); r.font.color.rgb = DARK_GREEN
        set_para_spacing(p, before=0, after=after)
    p3 = doc.add_paragraph()
    p3.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r3 = p3.add_run(subtitle)
    r3.italic = True; r3.font.name = "Arial"; r3.font.size = Pt(11); r3.font.color.rgb = NEAR_BLACK
    set_para_spacing(p3, before=0, after=300)


# ── Objection block helper: label + what they really mean + response + follow-up question
def add_objection(doc, objection, what_it_means, response, followup):
    add_h3(doc, f'Objection: "{objection}"')

    p = doc.add_paragraph()
    r1 = p.add_run("What they really mean: ")
    r1.bold = True; r1.font.name = "Arial"; r1.font.size = Pt(11); r1.font.color.rgb = DARK_GREEN
    r2 = p.add_run(what_it_means)
    r2.font.name = "Arial"; r2.font.size = Pt(11); r2.font.color.rgb = NEAR_BLACK
    r2.italic = True
    set_para_spacing(p, before=0, after=100)

    add_code(doc, response)

    p2 = doc.add_paragraph()
    r3 = p2.add_run("Follow-up question: ")
    r3.bold = True; r3.font.name = "Arial"; r3.font.size = Pt(11); r3.font.color.rgb = DARK_GREEN
    r4 = p2.add_run(followup)
    r4.font.name = "Arial"; r4.font.size = Pt(11); r4.font.color.rgb = NEAR_BLACK
    set_para_spacing(p2, before=0, after=200)


# ── BUILD DOCUMENT ─────────────────────────────────────────────────────────────

def build():
    doc = Document()
    for section in doc.sections:
        section.top_margin    = Inches(1.0)
        section.bottom_margin = Inches(1.0)
        section.left_margin   = Inches(1.1)
        section.right_margin  = Inches(1.1)

    add_title_block(
        doc,
        "COLD CALL OBJECTION HANDLING",
        "COMPLETE REFERENCE GUIDE",
        "Word-for-word responses to every pushback — for selling websites and lead management to small service businesses"
    )

    # ── WHAT'S INSIDE
    add_h2(doc, "What's inside this document")
    add_table(doc, [
        ("__header__", ("Section", "Contents")),
        ("Part 1 — The Framework",        "The four-step method behind every response: Acknowledge, Clarify, Reframe, Ask"),
        ("Part 2 — Website Objections",   "They already have a site / they use social media / their nephew will build it"),
        ("Part 3 — Lead Source Objections","Referrals only / they use Angi or HomeAdvisor / already too busy"),
        ("Part 4 — Money Objections",     "Can't afford it / not worth the price / want to think about it"),
        ("Part 5 — Trust Objections",     "Bad experience before / how do I know it works / never heard of you"),
        ("Part 6 — Time Objections",      "Too busy / call me next month / I'll do it myself"),
        ("Part 7 — Stall Tactics",        "Send me information / let me talk to my wife / I'll look at your website"),
        ("Part 8 — Quick Reference",      "One-line responses to every objection — for when you need a fast answer"),
    ])

    add_h2(doc, "How to use this document")
    add_body(doc,
        "Each objection entry includes three parts: what the client actually means beneath the surface, "
        "the word-for-word response to say out loud, and a follow-up question to keep the conversation "
        "moving. Read the responses out loud before calls. The goal of every response is not to win an "
        "argument — it is to lower perceived risk, address the real concern, and earn the next 30 seconds."
    )

    # ══════════════════════════════════════════════════════════════════════════
    # PART 1 — THE FRAMEWORK
    # ══════════════════════════════════════════════════════════════════════════
    add_h1(doc, "Part 1 — The Objection Handling Framework")

    add_h2(doc, "Step 1 — The Four-Step Method: A.C.R.A.")
    add_body(doc,
        "Every objection response in this guide follows the same four-step structure. "
        "Internalize this before reading the individual responses. When you know the structure, "
        "you can handle objections that are not in this guide."
    )

    add_table(doc, [
        ("__header__", ("Step", "What It Means and Why It Works")),
        ("A — Acknowledge",
         "Repeat back or validate what the client said. Never argue immediately. "
         "People need to feel heard before they can be moved. Example: 'Totally understand.' / 'That makes sense.' / 'I hear that a lot.'"),
        ("C — Clarify",
         "Ask one short question to understand what is really behind the objection. "
         "Most objections are surface-level. The real concern is usually risk, trust, or time. "
         "Example: 'Can I ask what the main concern is?' / 'Is it the price, or more the timing?'"),
        ("R — Reframe",
         "Reposition the objection as something that actually points toward your solution. "
         "Do not fight the objection — use it. Example: 'That's actually exactly why I'm calling — "
         "most businesses in your position are losing leads they don't even know they're losing.'"),
        ("A — Ask",
         "End every response with a question. Questions keep the call alive. "
         "A statement ends the conversation. A question continues it. "
         "Always ask for a specific small commitment: a 15-minute call, a callback time, a yes to seeing examples."),
    ])

    add_h3(doc, "What to Avoid")
    add_bullet_item(doc, "Do not argue facts. If they say 'I have plenty of leads,' don't say 'no you don't.' Say 'that's great — this makes sure it stays that way.'", bold_prefix="Never debate: ")
    add_bullet_item(doc, "Long explanations signal desperation. Every response should be 3-5 sentences. Then ask a question.", bold_prefix="Don't over-explain: ")
    add_bullet_item(doc, "'That's a great question!' reads as fake and scripted. Just answer.", bold_prefix="Don't use filler phrases: ")
    add_bullet_item(doc, "After a response, go silent. Silence is pressure. Let them fill it.", bold_prefix="Silence is a tool: ")
    add_bullet_item(doc, "If they are clearly not interested after two responses, give a clean exit. Respect their time — they will remember it.", bold_prefix="Know when to let go: ")

    # ══════════════════════════════════════════════════════════════════════════
    # PART 2 — WEBSITE OBJECTIONS
    # ══════════════════════════════════════════════════════════════════════════
    add_h1(doc, "Part 2 — Website Objections")

    add_h2(doc, "Step 2 — Objections About Their Existing Web Presence")

    add_objection(
        doc,
        objection   = "I already have a website.",
        what_it_means = "They may have a site, but almost certainly have no lead system behind it. This is a Product 2 opportunity.",
        response    = (
            '"That\'s great — a lot of the people I work with have websites too. The question is '
            'what happens when someone fills out your contact form at 10am while you\'re on a job. '
            'Does it text your phone within a few minutes with their name, number, and what they need? '
            'Or does it send an email to your inbox that you see four hours later? Because by then, '
            'they\'ve usually already called someone else. That\'s the gap I close."'
        ),
        followup    = '"Is your form connected to any kind of alert system, or does it just email you?"'
    )

    add_objection(
        doc,
        objection   = "I use Facebook / Instagram for my business.",
        what_it_means = "They have an online presence but it is not working for lead capture. Social platforms don't own search results.",
        response    = (
            '"Facebook is solid for staying in front of people who already know you. But when '
            'someone searches \'plumber near me\' on Google at 9pm, Facebook doesn\'t show up — '
            'a website does. And even if they find your Facebook page, there\'s no form, no system, '
            'nothing that alerts you within minutes. You\'re also at the mercy of the algorithm '
            'for reach. A website you own — it always shows up when someone\'s actively looking."'
        ),
        followup    = '"When someone finds you on Google right now, what do they see — a Facebook page, a website, or nothing?"'
    )

    add_objection(
        doc,
        objection   = "My nephew / friend / kid is going to build me one.",
        what_it_means = "They want to avoid spending money. The nephew almost never delivers. This is a stall, not a plan.",
        response    = (
            '"I hear that one fairly often, and it usually goes one of two ways — either it takes '
            'six months and comes out looking like a school project, or it never happens at all. '
            'And even if it does get built, it almost certainly won\'t have the lead system behind it '
            'that fires an alert to your phone when someone submits. I can have you live in three to '
            'five days, it\'ll look professional, and the whole system works from day one."'
        ),
        followup    = '"When is your nephew planning to have it done? Because I can have something live before that — and if his version comes out great, you can always switch."'
    )

    add_objection(
        doc,
        objection   = "I already pay for Wix / Squarespace / GoDaddy.",
        what_it_means = "They have a DIY site. These platforms have no lead management backend. This is a Product 2 sale.",
        response    = (
            '"Those platforms are fine for putting something online, but they don\'t have the '
            'backend system that actually manages your leads. When someone fills out your form, '
            'it sends you an email — and that email goes into the same inbox as everything else. '
            'No real-time alert, no Google Sheet that tracks every lead with their status, '
            'no morning digest with your call list for the day. I can add that system to what '
            'you already have without touching your site — you keep Wix, you just get the engine behind it."'
        ),
        followup    = '"How are you currently keeping track of who\'s reached out and whether you\'ve called them back?"'
    )

    # ══════════════════════════════════════════════════════════════════════════
    # PART 3 — LEAD SOURCE OBJECTIONS
    # ══════════════════════════════════════════════════════════════════════════
    add_h1(doc, "Part 3 — Lead Source Objections")

    add_h2(doc, "Step 3 — Objections About Where They Get Their Work")

    add_objection(
        doc,
        objection   = "I get all my work from referrals.",
        what_it_means = "They are comfortable and not feeling pain right now. The referral channel is not permanent or scalable.",
        response    = (
            '"Referrals are the best kind of customer — no argument there. But here\'s the thing: '
            'even your referral customers Google you before they call. If they find nothing, '
            'or an outdated site, you\'ve already lost a percentage of them without knowing it. '
            'And referrals dry up — slow seasons, someone moves, word of mouth slows down. '
            'This gives you a second channel that works while you\'re sleeping, so you\'re not '
            'entirely dependent on who talked to who this week."'
        ),
        followup    = '"Has there ever been a slow stretch where the referrals just weren\'t coming in fast enough?"'
    )

    add_objection(
        doc,
        objection   = "I already use Angi / HomeAdvisor / Thumbtack.",
        what_it_means = "They are paying for leads elsewhere — often expensively and for shared leads. This is a positioning opportunity.",
        response    = (
            '"Those services can work, but you\'re paying $20 to $80 per lead — and that same lead '
            'goes to three or four other contractors at the same time. You\'re in a race to call '
            'first every time, and you\'re paying whether you win the job or not. What I build '
            'sends you exclusive leads from your own site — people who specifically found you and '
            'chose to reach out. No competition, no cost per lead. One saved job from your own site '
            'pays for my service for months."'
        ),
        followup    = '"What\'s your close rate on the Angi leads — are you winning most of them or fighting for them?"'
    )

    add_objection(
        doc,
        objection   = "I'm already too busy — I don't need more leads.",
        what_it_means = "Business is good right now. They don't see a future gap. The risk is they have no safety net when it slows down.",
        response    = (
            '"That\'s a great position to be in. A lot of my best clients came to me exactly when '
            'things were good — because that\'s when you have the space to set up the system '
            'without pressure. What you don\'t want is to be scrambling for leads when work '
            'slows down in the off-season and have nothing built. This takes three to five days '
            'to set up, you barely notice it\'s happening, and it\'s just there working for you '
            'when you need it."'
        ),
        followup    = '"Is it consistently this busy year-round, or is there a slower stretch where you\'re looking for more work?"'
    )

    # ══════════════════════════════════════════════════════════════════════════
    # PART 4 — MONEY OBJECTIONS
    # ══════════════════════════════════════════════════════════════════════════
    add_h1(doc, "Part 4 — Money Objections")

    add_h2(doc, "Step 4 — Objections About Price and Value")

    add_objection(
        doc,
        objection   = "I can't afford it right now.",
        what_it_means = "Cash is tight, or they don't see the ROI yet. This is almost always a value problem, not a budget problem.",
        response    = (
            '"I understand — and that\'s exactly why the math matters here. If your average job '
            'is worth $500 and you recover even one lead a month you would have otherwise missed, '
            'that\'s $500 coming back in against $79 going out. The service pays for itself the '
            'first time your phone buzzes with a lead while you\'re under a sink. And if it '
            'doesn\'t work in the first 60 days, I refund every monthly payment — no questions. '
            'You\'re not spending money. You\'re recovering money you\'re already losing."'
        ),
        followup    = '"What would it take for you to feel like this was worth it — like, what would you need to see in the first month?"'
    )

    add_objection(
        doc,
        objection   = "That's too expensive.",
        what_it_means = "They are comparing to zero (doing nothing) or don't yet see what they are getting.",
        response    = (
            '"Compared to what, though? An agency charges $5,000 to $10,000 for a site with no '
            'lead system and takes three months to deliver. Squarespace is $20 a month but you '
            'spend 40 hours building it yourself and there\'s no alert system, no tracker, '
            'nothing automated. I\'m under $100 a month, it\'s live in a week, and I manage '
            'everything. The value isn\'t the website — it\'s the leads it catches and the '
            'time it saves you."'
        ),
        followup    = '"What were you expecting it to cost — just so I understand where you\'re coming from?"'
    )

    add_objection(
        doc,
        objection   = "I need to think about it.",
        what_it_means = "They are not sold yet — usually one unanswered concern is still sitting there. Find it.",
        response    = (
            '"Completely fair — I\'d never push you into a decision. Can I ask: is there '
            'a specific part you\'re unsure about? Sometimes there\'s one question I can answer '
            'that makes the rest clear. Is it the price, the setup process, whether it\'ll '
            'actually work — what\'s the thing that\'s making you pause?"'
        ),
        followup    = '"If that concern wasn\'t there, would this be something you\'d move forward with?"'
    )

    add_objection(
        doc,
        objection   = "Can you just do it cheaper?",
        what_it_means = "They want the service but are testing whether the price is real. Hold the price — offer the Starter tier instead.",
        response    = (
            '"I don\'t discount the packages because the price reflects the actual work that '
            'goes into them. But what I can do is start you on the Starter package — that\'s '
            '$800 to set up and $79 a month. You get the professional website, the lead form, '
            'and the Google Sheet tracker. If you want to add the real-time alerts and morning '
            'digest later, we can move up to Growth. A lot of clients start there and upgrade '
            'within the first couple months once they see how it works."'
        ),
        followup    = '"Would the Starter package be something you could move forward with, or is the concern more about the monthly fee?"'
    )

    # ══════════════════════════════════════════════════════════════════════════
    # PART 5 — TRUST OBJECTIONS
    # ══════════════════════════════════════════════════════════════════════════
    add_h1(doc, "Part 5 — Trust Objections")

    add_h2(doc, "Step 5 — Objections About Credibility and Risk")

    add_objection(
        doc,
        objection   = "I had a bad experience with a web company before.",
        what_it_means = "They were burned — slow delivery, no results, locked into a contract, or couldn't reach anyone. Fear of repeating it.",
        response    = (
            '"That happens a lot and it\'s completely understandable. Most web agencies are slow, '
            'expensive, and once the site is built you\'re dealing with a support ticket and '
            'a rotating team. Here\'s how I\'m different: I\'m one person, you deal with me '
            'directly, you can call or text me when something needs to change. Your domain is '
            'in your name. Your leads go to your own Google Sheet. If you leave, you take '
            'everything. And there\'s a 60-day money-back guarantee on the monthly fee — '
            'so you\'re not locked into anything."'
        ),
        followup    = '"What went wrong with the last company — was it the quality, the timeline, or something else? I want to make sure I\'m not repeating whatever that was."'
    )

    add_objection(
        doc,
        objection   = "How do I know this will actually work?",
        what_it_means = "They want proof. They are risk-averse and need evidence before committing.",
        response    = (
            '"That\'s the right question to ask. I\'ll show you examples of exactly what '
            'the sites and the lead system look like — real working examples, not mockups. '
            'And the 60-day guarantee means if you sign up and in two months it hasn\'t '
            'generated a single lead or you\'re not happy for any reason, I refund every '
            'monthly payment. You\'re risking the setup fee, which you get a live website '
            'and working lead system for regardless. That\'s the deal."'
        ),
        followup    = '"Can I send you two or three examples of what the site and lead system look like right now — just so you can see what you\'d actually be getting?"'
    )

    add_objection(
        doc,
        objection   = "I've never heard of you.",
        what_it_means = "They want social proof. They are asking: are you real and can I trust you?",
        response    = (
            '"Fair enough — I\'m not a big agency with billboards. I work directly with '
            'small service businesses in [city/area], one client at a time. That\'s actually '
            'the point — you get one person who knows your business, not a call center. '
            'I can show you what I\'ve built for other clients, walk you through the system '
            'live, and if you\'re not satisfied in 60 days, you get your monthly fees back. '
            'You\'re not trusting me on faith — you\'re trusting a guarantee."'
        ),
        followup    = '"Would it help to see a couple of live examples of sites I\'ve built before we go any further?"'
    )

    add_objection(
        doc,
        objection   = "What if I want to cancel?",
        what_it_means = "They are worried about being locked in. This is a buying signal — they are considering it.",
        response    = (
            '"You can cancel anytime. There\'s no contract on the monthly service — '
            'it\'s month to month. The only commitment is the one-time setup fee, '
            'which is what covers the work to actually build everything. Your domain '
            'stays in your name. Your Google Sheet stays in your Google account. '
            'If you walk away, you take everything with you. I don\'t lock anyone in '
            'because I want clients who are staying because it\'s working, not because they\'re trapped."'
        ),
        followup    = '"Does the fact that it\'s month-to-month make this feel more manageable?"'
    )

    # ══════════════════════════════════════════════════════════════════════════
    # PART 6 — TIME OBJECTIONS
    # ══════════════════════════════════════════════════════════════════════════
    add_h1(doc, "Part 6 — Time Objections")

    add_h2(doc, "Step 6 — Objections About Time and Availability")

    add_objection(
        doc,
        objection   = "I'm too busy right now.",
        what_it_means = "They are genuinely busy, or this is an easy exit. Either way, make the ask tiny.",
        response    = (
            '"I figured — that\'s why I built the onboarding to take almost nothing from you. '
            'You answer a few questions about your business, I handle everything else. '
            'The whole setup conversation takes about 15 minutes. You don\'t manage the system, '
            'you don\'t update the site — that\'s all me. The only thing you do is call leads '
            'back when they come in. That\'s it."'
        ),
        followup    = '"When\'s a 15-minute window that works for you — even tomorrow morning before your first job?"'
    )

    add_objection(
        doc,
        objection   = "Call me back next month / after the season.",
        what_it_means = "They are pushing you off. Sometimes genuine (genuinely slammed), often a soft no. Pin down a specific date.",
        response    = (
            '"No problem at all. Can I ask — is it more that the timing is bad right now, '
            'or is there something specific you want to think through before then? '
            'Because if it\'s timing, I\'m happy to circle back. I just want to make sure '
            'we\'re not leaving leads on the table for another month while the slow season '
            'starts to creep in."'
        ),
        followup    = '"What\'s a specific date I can put in my calendar to follow up — I\'ll call you on that morning."'
    )

    add_objection(
        doc,
        objection   = "I don't have time to manage a website.",
        what_it_means = "They think this requires ongoing effort from them. They do not understand the fully managed model.",
        response    = (
            '"You don\'t manage anything — that\'s the entire point. I manage the site, '
            'I make updates when things change, I monitor the lead system. You get a '
            'notification when a lead comes in and you decide whether to call them. '
            'That\'s the only thing you do. It\'s less work than checking your Facebook messages."'
        ),
        followup    = '"What part of it did you think would fall on you? I want to make sure I\'m being clear about what\'s actually on my plate vs. yours."'
    )

    add_objection(
        doc,
        objection   = "I'll do it myself when I have time.",
        what_it_means = "Classic DIY intention that almost never happens. Building a site and a lead system takes 40-80 hours.",
        response    = (
            '"I\'ve heard that one a lot — and honestly, most people who say that never get '
            'around to it, not because they don\'t want to but because the job always comes first. '
            'And even if you do build the site, the lead alert system and the Google Sheets '
            'tracker and the morning digest — that\'s a separate technical build on top of the '
            'site. I can have the whole thing live in less than a week. What would it take '
            'for you to do it yourself — realistically?"'
        ),
        followup    = '"What\'s your timeline for getting to it — like a real date?"'
    )

    # ══════════════════════════════════════════════════════════════════════════
    # PART 7 — STALL TACTICS
    # ══════════════════════════════════════════════════════════════════════════
    add_h1(doc, "Part 7 — Stall Tactics")

    add_h2(doc, "Step 7 — Polite Exits That Are Not Real Objections")

    add_objection(
        doc,
        objection   = "Just send me some information.",
        what_it_means = "They want to end the call without saying no. Information sent cold almost never converts. Use it as a reason to re-engage.",
        response    = (
            '"Absolutely — I\'ll send you a couple of examples of sites I\'ve built and '
            'a quick breakdown of how the lead system works. Can I ask: is there a specific '
            'part of it you\'re most curious about so I can make sure what I send is actually '
            'useful? And when I follow up, should I call you or text you — which do you prefer?"'
        ),
        followup    = '"I\'ll get that over to you today. What\'s the best day to follow up — would [specific day] work to do a quick 10-minute call once you\'ve looked it over?"'
    )

    add_objection(
        doc,
        objection   = "I need to talk to my wife / partner / business partner.",
        what_it_means = "They are not the only decision maker, or this is a soft stall. Either way, offer to include the other person.",
        response    = (
            '"That makes total sense — I\'d want to do the same. Would it be easier to '
            'get both of you on a quick call at the same time? That way I can answer '
            'any questions directly and you\'re not playing telephone. It doesn\'t have '
            'to be long — 15 minutes and you\'ll have everything you need to make a decision."'
        ),
        followup    = '"When would a time work that you\'re both available — even just briefly?"'
    )

    add_objection(
        doc,
        objection   = "I'll look at your website and get back to you.",
        what_it_means = "Another exit tactic. They will not look. Offer to walk them through it live instead.",
        response    = (
            '"I appreciate that. Honestly, the website gives you a general idea but it\'s '
            'not the same as seeing the lead system actually work — like watching a lead '
            'come in and hit the Google Sheet and trigger an email in real time. That\'s '
            'the part that usually makes it click for people. Can I do a 10-minute screen '
            'share with you this week so you can see exactly how it works live?"'
        ),
        followup    = '"What day works — I can do mornings or evenings, whatever fits around your schedule."'
    )

    add_objection(
        doc,
        objection   = "I'm not interested.",
        what_it_means = "Hard no, or they were never qualified to begin with. One soft reframe, then a clean exit.",
        response    = (
            '"Completely understand — I won\'t take more of your time. Can I ask one quick '
            'question before I let you go: is it that the timing is wrong, or is it more '
            'that this isn\'t something that feels relevant to your business right now? '
            'Just helps me understand where I\'m missing the mark."'
        ),
        followup    = '"Either way, appreciate you picking up. If anything changes down the road, feel free to reach out. Have a good one."'
    )

    # ══════════════════════════════════════════════════════════════════════════
    # PART 8 — QUICK REFERENCE
    # ══════════════════════════════════════════════════════════════════════════
    add_h1(doc, "Part 8 — Quick Reference Cheat Sheet")

    add_h2(doc, "Step 8 — One-Line Reframes for Every Objection")
    add_body(doc,
        "Use this table when you need a fast response and can't recall the full script. "
        "These are the core reframe for each objection — the single sentence that turns it around. "
        "Follow every reframe with a question."
    )

    add_table(doc, [
        ("__header__", ("Objection", "One-Line Reframe")),
        ("I already have a website.",
         "A website without a lead system is a brochure. This makes it work."),
        ("I use Facebook for my business.",
         "Facebook keeps you in front of people who know you. Google gets you in front of people who don't — yet."),
        ("My nephew is building one.",
         "Most of those never happen. I'll have you live in 5 days — if his version comes out great, you can switch."),
        ("I already pay for Wix / Squarespace.",
         "Keep it — I just add the backend system that makes your form actually capture and track leads."),
        ("I get all my work from referrals.",
         "Even referral customers Google you first. And referrals dry up — this is your backup."),
        ("I use Angi / HomeAdvisor.",
         "You're paying $20-80 per shared lead. My system sends you exclusive leads — no per-lead cost."),
        ("I'm already too busy.",
         "Set it up now while things are good. You don't want to be scrambling when it slows down."),
        ("I can't afford it.",
         "One recovered job pays for 6-12 months of service. You're not spending money — you're recovering it."),
        ("That's too expensive.",
         "Agencies charge $5,000-$15,000. This is under $100 a month, fully managed, live in a week."),
        ("I need to think about it.",
         "What's the one thing making you pause? Usually I can answer that in under a minute."),
        ("Can you do it cheaper?",
         "The packages are fixed, but the Starter tier at $79/month is the right place to start."),
        ("I had a bad experience before.",
         "You deal with me directly, you own everything, and you can cancel anytime. That's not how agencies work."),
        ("How do I know it'll work?",
         "60-day money-back guarantee on monthly fees. If it doesn't work, you get every payment back."),
        ("I've never heard of you.",
         "I'm one person — that's the point. You always know who you're calling."),
        ("What if I want to cancel?",
         "Month-to-month, no contract. You own your domain and your data. Walk away anytime."),
        ("I'm too busy right now.",
         "The setup takes 15 minutes from you — I do everything else."),
        ("Call me next month.",
         "Absolutely — what specific date should I put in my calendar?"),
        ("I don't have time to manage a website.",
         "You don't manage anything. That's entirely on me."),
        ("I'll do it myself.",
         "It takes 40-80 hours to build a site and lead system. I do it in 5 days."),
        ("Send me some information.",
         "I'll send it today — and I'll follow up [day] to walk through it. Call or text?"),
        ("I need to talk to my partner.",
         "Let's get both of you on a 15-minute call — easier than explaining it secondhand."),
        ("I'll look at your website.",
         "The site doesn't show the system working live. Can I do a 10-minute screen share instead?"),
        ("I'm not interested.",
         "Understood — is it timing or is this just not relevant to your business right now?"),
    ])

    add_h3(doc, "The Clean Exit — Use This When It's a Hard No")
    add_code(doc,
        '"No problem at all — I appreciate you taking the time. If anything changes, feel free '
        'to reach out. Have a good one."'
    )
    add_body(doc,
        "Do not follow up more than once after a hard no. Move on. A short, respectful exit "
        "leaves a positive impression — and sometimes that person calls back six months later "
        "when they're ready."
    )

    out_path = "/Users/creighbaby/LeadManagement/Cold Call Objection Handling — Complete Reference Guide.docx"
    doc.save(out_path)
    print(f"Saved: {out_path}")


if __name__ == "__main__":
    build()
