"""
Lead Intake — Apps Script Setup Guide
Generates a branded .docx following the LeadManagement Source-of-Truth style.
Covers: what to gather, Google Sheet setup, webhook paste + deploy, Twilio SMS,
voicemail transcription via Twilio Voice, morning digest trigger, and testing.
"""

from docx import Document
from docx.shared import Pt, Inches, RGBColor, Twips
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

DARK_GREEN   = RGBColor(0x1a, 0x3a, 0x2a)
NEAR_BLACK   = RGBColor(0x1c, 0x1c, 0x1a)
GOLD         = "c8a96e"
CREAM_LINE   = "e2e0d8"
TABLE_HDR_BG = "2c1f0e"
WHITE        = RGBColor(0xFF, 0xFF, 0xFF)
CODE_BG      = "F2F2F2"
BORDER_COLOR = "D0CEC6"


def set_para_spacing(para, before=0, after=120):
    pf = para.paragraph_format
    pf.space_before = Twips(before)
    pf.space_after  = Twips(after)

def add_bottom_border(para, color_hex, sz=12, space=4):
    pPr = para._p.get_or_add_pPr()
    pBdr = OxmlElement("w:pBdr")
    bottom = OxmlElement("w:bottom")
    bottom.set(qn("w:val"), "single"); bottom.set(qn("w:sz"), str(sz))
    bottom.set(qn("w:space"), str(space)); bottom.set(qn("w:color"), color_hex)
    pBdr.append(bottom); pPr.append(pBdr)

def set_cell_shading(cell, fill_hex):
    tc = cell._tc; tcPr = tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"), "clear"); shd.set(qn("w:color"), "auto"); shd.set(qn("w:fill"), fill_hex)
    tcPr.append(shd)

def set_cell_margins(cell, top=80, bottom=80, left=115, right=115):
    tc = cell._tc; tcPr = tc.get_or_add_tcPr()
    tcMar = OxmlElement("w:tcMar")
    for side, val in [("top", top), ("bottom", bottom), ("left", left), ("right", right)]:
        el = OxmlElement(f"w:{side}"); el.set(qn("w:w"), str(val)); el.set(qn("w:type"), "dxa")
        tcMar.append(el)
    tcPr.append(tcMar)

def set_table_borders(table, color_hex=BORDER_COLOR):
    tbl = table._tbl; tblPr = tbl.find(qn("w:tblPr"))
    if tblPr is None: tblPr = OxmlElement("w:tblPr"); tbl.insert(0, tblPr)
    tblBdr = OxmlElement("w:tblBorders")
    for side in ("top", "left", "bottom", "right", "insideH", "insideV"):
        el = OxmlElement(f"w:{side}")
        el.set(qn("w:val"), "single"); el.set(qn("w:sz"), "4")
        el.set(qn("w:space"), "0"); el.set(qn("w:color"), color_hex)
        tblBdr.append(el)
    tblPr.append(tblBdr)

def add_h1(doc, text):
    para = doc.add_paragraph()
    run = para.add_run(text)
    run.bold = True; run.font.name = "Arial"; run.font.size = Pt(18); run.font.color.rgb = DARK_GREEN
    set_para_spacing(para, before=400, after=200); add_bottom_border(para, GOLD, sz=12, space=4)
    return para

def add_h2(doc, text):
    para = doc.add_paragraph()
    run = para.add_run(text)
    run.bold = True; run.font.name = "Arial"; run.font.size = Pt(14); run.font.color.rgb = DARK_GREEN
    set_para_spacing(para, before=300, after=160); add_bottom_border(para, CREAM_LINE, sz=4, space=4)
    return para

def add_h3(doc, text):
    para = doc.add_paragraph()
    run = para.add_run(text)
    run.bold = True; run.font.name = "Arial"; run.font.size = Pt(12); run.font.color.rgb = DARK_GREEN
    set_para_spacing(para, before=220, after=120)
    return para

def add_body(doc, text):
    para = doc.add_paragraph()
    run = para.add_run(text)
    run.font.name = "Arial"; run.font.size = Pt(11); run.font.color.rgb = NEAR_BLACK
    set_para_spacing(para, before=0, after=120)
    return para

def add_bullet_item(doc, text, bold_prefix=None):
    para = doc.add_paragraph(style="List Bullet")
    if bold_prefix:
        r1 = para.add_run(bold_prefix); r1.bold = True
        r1.font.name = "Arial"; r1.font.size = Pt(11); r1.font.color.rgb = NEAR_BLACK
    run = para.add_run(text)
    run.font.name = "Arial"; run.font.size = Pt(11); run.font.color.rgb = NEAR_BLACK
    set_para_spacing(para, before=0, after=80)
    return para

def add_code(doc, text):
    para = doc.add_paragraph()
    run = para.add_run(text)
    run.font.name = "Courier New"; run.font.size = Pt(9); run.font.color.rgb = NEAR_BLACK
    set_para_spacing(para, before=60, after=60)
    pPr = para._p.get_or_add_pPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"), "clear"); shd.set(qn("w:color"), "auto"); shd.set(qn("w:fill"), CODE_BG)
    pPr.append(shd)
    return para

def add_table(doc, rows_data, col_widths=(1.75, 4.5)):
    table = doc.add_table(rows=0, cols=2)
    set_table_borders(table)
    for label, detail in rows_data:
        row = table.add_row(); left = row.cells[0]; right = row.cells[1]
        left.width = Inches(col_widths[0]); right.width = Inches(col_widths[1])
        set_cell_margins(left); set_cell_margins(right)
        if label == "__header__":
            set_cell_shading(left, TABLE_HDR_BG); set_cell_shading(right, TABLE_HDR_BG)
            lp = left.paragraphs[0]; lr = lp.add_run(detail[0])
            lr.bold = True; lr.font.name = "Arial"; lr.font.size = Pt(11); lr.font.color.rgb = WHITE
            rp = right.paragraphs[0]; rr = rp.add_run(detail[1])
            rr.bold = True; rr.font.name = "Arial"; rr.font.size = Pt(11); rr.font.color.rgb = WHITE
        else:
            lp = left.paragraphs[0]; lr = lp.add_run(label)
            lr.bold = True; lr.font.name = "Arial"; lr.font.size = Pt(11); lr.font.color.rgb = DARK_GREEN
            rp = right.paragraphs[0]; rr = rp.add_run(detail)
            rr.font.name = "Arial"; rr.font.size = Pt(11); rr.font.color.rgb = NEAR_BLACK
    spacer = doc.add_paragraph(); set_para_spacing(spacer, before=0, after=80)
    return table

def add_title_block(doc, line1, line2, subtitle):
    p1 = doc.add_paragraph(); p1.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r1 = p1.add_run(line1); r1.bold = True; r1.font.name = "Arial"
    r1.font.size = Pt(22); r1.font.color.rgb = DARK_GREEN
    set_para_spacing(p1, before=0, after=60)
    p2 = doc.add_paragraph(); p2.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r2 = p2.add_run(line2); r2.bold = True; r2.font.name = "Arial"
    r2.font.size = Pt(22); r2.font.color.rgb = DARK_GREEN
    set_para_spacing(p2, before=0, after=100)
    p3 = doc.add_paragraph(); p3.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r3 = p3.add_run(subtitle); r3.italic = True; r3.font.name = "Arial"
    r3.font.size = Pt(11); r3.font.color.rgb = NEAR_BLACK
    set_para_spacing(p3, before=0, after=300)


def build():
    doc = Document()
    for section in doc.sections:
        section.top_margin    = Inches(1.0)
        section.bottom_margin = Inches(1.0)
        section.left_margin   = Inches(1.1)
        section.right_margin  = Inches(1.1)

    add_title_block(doc,
        "LEAD INTAKE",
        "APPS SCRIPT SETUP GUIDE",
        "What to gather, how to deploy the webhook, set up voicemail transcription, and test everything"
    )

    add_h2(doc, "What's inside this document")
    add_table(doc, [
        ("__header__", ("Section", "Contents")),
        ("Part 1 — What to Gather",            "Line-by-line list of everything to collect before starting any client setup"),
        ("Part 2 — How It Works",              "The full lead flow — form submissions and voicemail, end to end"),
        ("Part 3 — Create the Google Sheet",   "Setting up the client's lead log"),
        ("Part 4 — Paste and Deploy the Script","Opening Apps Script, pasting the webhook, setting properties, deploying"),
        ("Part 5 — Set Up Voicemail",          "Twilio phone number, TwiML Bin, and connecting it to the webhook"),
        ("Part 6 — Connect to the Website",    "Pasting the webhook URL into index.html"),
        ("Part 7 — Set Up Morning Digest",     "Daily trigger for the morning lead summary email"),
        ("Part 8 — Test Everything",           "Running both test functions and confirming all outputs"),
        ("Quick Reference",                    "Full checklist, Script Properties table, and troubleshooting"),
    ])

    add_h2(doc, "How to use this document")
    add_body(doc,
        "Run through this guide once per client. Each client gets their own Google Sheet, "
        "their own Apps Script deployment, and their own Twilio phone number. "
        "Start with Part 1 — gather everything before opening a single browser tab."
    )

    # ══════════════════════════════════════════════════════════════════════════
    # PART 1 — WHAT TO GATHER
    # ══════════════════════════════════════════════════════════════════════════
    add_h1(doc, "Part 1 — What to Gather Before You Start")

    add_h2(doc, "Step 1 — Master Gather List")
    add_body(doc,
        "Collect everything below before touching any setup. Items marked with an asterisk (*) "
        "change per client. Everything else is from your own accounts and stays the same."
    )

    add_h3(doc, "From the Client — Collect Before or During Onboarding Call")
    add_bullet_item(doc, "* Client's full name")
    add_bullet_item(doc, "* Business name (exact spelling they want on the site and voicemail greeting)")
    add_bullet_item(doc, "* Business email address — this is where lead alerts and the monthly report go")
    add_bullet_item(doc, "* Cell phone number — this is where SMS alerts go. Write digits only with country code (e.g. 13125550100)")
    add_bullet_item(doc, "* List of services they offer — for the intake form dropdown and the website services section")
    add_bullet_item(doc, "* City and service area — for the website and form (e.g. 'Chicago and surrounding suburbs')")
    add_bullet_item(doc, "* Google account email — the Gmail tied to their Google Calendar (needed if setting up Calendly)")
    add_bullet_item(doc, "* Their availability hours — days and times they accept bookings (needed if setting up Calendly)")
    add_bullet_item(doc, "* Existing website URL — if they have one (determines which package applies)")
    add_bullet_item(doc, "* Domain name — if they have one, or if you need to help them buy one")
    add_bullet_item(doc, "* Any existing branding — logo file, brand colors, photos (for Website + Lead Management package only)")

    add_h3(doc, "From Your Twilio Account — twilio.com Console")
    add_bullet_item(doc, "Your Account SID — Console dashboard, top left. Starts with AC. Same for every client.")
    add_bullet_item(doc, "Your Auth Token — Console dashboard, click the eye icon to reveal. Same for every client.")
    add_bullet_item(doc, "Your SMS sending number — Phone Numbers → Manage → Active Numbers. Format: +15551234567. Same for every client.")
    add_bullet_item(doc, "* A new Twilio phone number for this client — buy one per client for voicemail. Phone Numbers → Buy a number. Cost: $1.00/month.")

    add_h3(doc, "From Your Google Account — sheets.google.com")
    add_bullet_item(doc, "* Create a new blank Google Sheet named '[Business Name] — Lead Log' before starting the Apps Script setup.")

    add_h3(doc, "From Your GitHub Account — github.com (Website + Lead Management package only)")
    add_bullet_item(doc, "* Create a new repository named client-[business-slug]-site before uploading the HTML.")
    add_bullet_item(doc, "The client's completed index.html file with all placeholders filled in.")

    add_h3(doc, "From Calendly — calendly.com (if client is getting scheduling)")
    add_bullet_item(doc, "* Create the client's Calendly account using their business email before the onboarding call.")
    add_bullet_item(doc, "* The booking event URL — from Event Types → Share (e.g. calendly.com/smith-plumbing/free-estimate).")
    add_bullet_item(doc, "* Their accent hex color — from the --accent CSS variable in their HTML file (no # symbol).")

    add_h3(doc, "Generated During Setup — Note These as You Go")
    add_bullet_item(doc, "* Apps Script deployment URL — copy this immediately after deploying. Goes into index.html AND the TwiML Bin.")
    add_bullet_item(doc, "* Client's Twilio phone number — the one you just bought. Goes on the website and in the TwiML Bin.")
    add_bullet_item(doc, "* GitHub Pages URL — copy from repo Settings → Pages after enabling. Goes into DNS setup.")

    add_h3(doc, "What Changes Per Client vs. What Stays the Same")
    add_table(doc, [
        ("__header__", ("Item", "Per Client or Fixed?")),
        ("Client email (NOTIFY_EMAIL)",        "Per client — changes every time"),
        ("Client phone (NOTIFY_PHONE)",        "Per client — changes every time"),
        ("Twilio Account SID",                 "Fixed — always your account"),
        ("Twilio Auth Token",                  "Fixed — always your account"),
        ("Twilio SMS sending number",          "Fixed — always your number"),
        ("Twilio voicemail number",            "Per client — one new number per client ($1/mo)"),
        ("Apps Script deployment URL",         "Per client — new deployment for each client's Sheet"),
        ("Google Sheet",                       "Per client — one Sheet per client"),
        ("GitHub repository",                  "Per client — one repo per client"),
    ])

    # ══════════════════════════════════════════════════════════════════════════
    # PART 2 — HOW IT WORKS
    # ══════════════════════════════════════════════════════════════════════════
    add_h1(doc, "Part 2 — How It Works")

    add_h2(doc, "Step 2 — The Two Lead Flows")

    add_h3(doc, "Flow A — Form Submission")
    add_table(doc, [
        ("__header__", ("Event", "What Happens")),
        ("1. Form submitted",       "Visitor fills out the lead intake form and clicks Send."),
        ("2. Browser POSTs JSON",   "JavaScript sends the lead data to the Apps Script webhook URL."),
        ("3. doPost() receives it", "Script detects JSON body — routes to handleFormSubmission()."),
        ("4. Sheet logged",         "Lead appended to Leads tab with timestamp and Source = Form."),
        ("5. Email sent",           "Formatted email alert delivered to client's inbox immediately."),
        ("6. SMS sent",             "Twilio API called — text arrives on client's phone within seconds."),
    ])

    add_h3(doc, "Flow B — Voicemail Transcription")
    add_table(doc, [
        ("__header__", ("Event", "What Happens")),
        ("1. Customer calls",          "Customer dials the client's Twilio phone number."),
        ("2. TwiML plays greeting",    "Twilio reads the greeting text and records the voicemail (up to 2 min)."),
        ("3. Twilio transcribes",      "After the call ends, Twilio transcribes the audio automatically."),
        ("4. Twilio POSTs callback",   "Twilio sends the transcription to the Apps Script webhook URL (same URL as the form)."),
        ("5. doPost() detects Twilio", "Script sees TranscriptionText parameter — routes to handleVoicemailTranscription()."),
        ("6. Sheet logged",            "Lead appended to Leads tab with caller's phone number and Source = Voicemail."),
        ("7. Email sent",              "Voicemail-specific email with transcription text delivered to client."),
        ("8. SMS sent",                "First 100 characters of transcription + caller number sent to client's phone."),
    ])

    # ══════════════════════════════════════════════════════════════════════════
    # PART 3 — CREATE THE GOOGLE SHEET
    # ══════════════════════════════════════════════════════════════════════════
    add_h1(doc, "Part 3 — Create the Google Sheet")

    add_h2(doc, "Step 3 — Set Up the Client's Lead Log")
    add_body(doc,
        "Each client gets their own Google Sheet. The script auto-creates the Leads tab "
        "and column headers on the first submission — you do not need to build the columns manually."
    )
    add_bullet_item(doc, "Go to sheets.google.com → click + to create a blank spreadsheet")
    add_bullet_item(doc, "Rename it: [Business Name] — Lead Log (e.g. Smith Plumbing — Lead Log)")
    add_bullet_item(doc, "Share it with the client's Google account — Viewer access is sufficient")
    add_bullet_item(doc, "Leave this tab open — you will open Apps Script from it in the next step")

    # ══════════════════════════════════════════════════════════════════════════
    # PART 4 — PASTE AND DEPLOY THE SCRIPT
    # ══════════════════════════════════════════════════════════════════════════
    add_h1(doc, "Part 4 — Paste and Deploy the Script")

    add_h2(doc, "Step 4 — Open Apps Script and Load the Webhook")
    add_bullet_item(doc, "From inside the Google Sheet: Extensions → Apps Script")
    add_bullet_item(doc, "Delete all existing code in the editor (select all → delete)")
    add_bullet_item(doc, "Open the file: clients/templates/lead-intake-webhook.gs")
    add_bullet_item(doc, "Copy the entire contents and paste into the Apps Script editor")
    add_bullet_item(doc, "Save (Cmd+S) — rename the project to the client's business name if desired")

    add_h2(doc, "Step 5 — Set Script Properties")
    add_body(doc,
        "Script Properties store credentials outside the code. Never hardcode them in the script."
    )
    add_bullet_item(doc, "Click the gear icon (Project Settings) in the left sidebar")
    add_bullet_item(doc, "Scroll to Script Properties → click Add script property for each of the five below")

    add_h3(doc, "The Five Script Properties")
    add_table(doc, [
        ("__header__", ("Property Name", "Value")),
        ("NOTIFY_EMAIL",
         "Client's email address — e.g. mike@smithplumbing.com\n"
         "Lead alert emails and the monthly digest go here."),
        ("NOTIFY_PHONE",
         "Client's cell phone — digits + country code only, no symbols.\n"
         "e.g. 13125550100  (for +1 312 555 0100)"),
        ("TWILIO_ACCOUNT_SID",
         "Your Twilio Account SID — from the Twilio Console dashboard.\n"
         "Starts with AC. Same for every client."),
        ("TWILIO_AUTH_TOKEN",
         "Your Twilio Auth Token — from the Twilio Console dashboard.\n"
         "Same for every client. Treat like a password."),
        ("TWILIO_FROM_NUMBER",
         "Your Twilio SMS sending number — e.g. +15551234567\n"
         "Include the + and country code. Same for every client."),
    ])

    add_body(doc, "Click Save after entering all five. Property names are case-sensitive.")

    add_h2(doc, "Step 6 — Deploy as a Web App")
    add_bullet_item(doc, "Click Deploy (top right) → New deployment")
    add_bullet_item(doc, "Click the gear icon next to Select type → Web app")
    add_bullet_item(doc, "Execute as: Me  |  Who has access: Anyone")
    add_bullet_item(doc, "Click Deploy → authorize if prompted (Review permissions → Allow)")
    add_bullet_item(doc, "Copy the Web App URL — it looks like:")
    add_code(doc, "https://script.google.com/macros/s/AKfycb.../exec")
    add_bullet_item(doc, "Save this URL — you will paste it into index.html AND the TwiML Bin")

    add_h3(doc, "When Redeployment Is Needed")
    add_table(doc, [
        ("__header__", ("Change Type", "Redeploy Required?")),
        ("Updating Script Properties (email, phone, Twilio keys)", "No — read at runtime, effective immediately"),
        ("Editing the script code",                                "Yes — Deploy → New deployment. URL changes, update index.html."),
        ("Updating the form fields in HTML only",                  "No — HTML change only, script unchanged"),
    ])

    # ══════════════════════════════════════════════════════════════════════════
    # PART 5 — SET UP VOICEMAIL
    # ══════════════════════════════════════════════════════════════════════════
    add_h1(doc, "Part 5 — Set Up Voicemail Transcription")

    add_h2(doc, "Step 7 — Buy a Twilio Phone Number for This Client")
    add_body(doc,
        "Each client gets their own dedicated Twilio number. This is the number that goes on "
        "their website and that customers call. When a call is not answered, Twilio plays a "
        "greeting, records the voicemail, transcribes it, and sends the transcription to the "
        "Apps Script webhook — where it is logged and alerted exactly like a form submission."
    )
    add_bullet_item(doc, "Log in to twilio.com → Console → Phone Numbers → Manage → Buy a number")
    add_bullet_item(doc, "Search for a local number in the client's area code")
    add_bullet_item(doc, "Make sure Voice capability is checked")
    add_bullet_item(doc, "Click Buy — cost is $1.00/month, charged to your Twilio account")
    add_bullet_item(doc, "Note the number in E.164 format (e.g. +13125550199) — this goes on the client's website")

    add_h2(doc, "Step 8 — Create a TwiML Bin")
    add_body(doc,
        "A TwiML Bin is a free Twilio feature that stores static XML instructions. "
        "It tells Twilio what to do when someone calls the number — play a greeting and record the voicemail."
    )
    add_bullet_item(doc, "In the Twilio Console: Explore Products → TwiML Bins (or search 'TwiML Bins')")
    add_bullet_item(doc, "Click Create new TwiML Bin")
    add_bullet_item(doc, "Name it: [Business Name] Voicemail (e.g. Smith Plumbing Voicemail)")
    add_bullet_item(doc, "Open the file: clients/templates/twiml-voicemail.xml")
    add_bullet_item(doc, "Replace [BUSINESS_NAME] with the client's business name")
    add_bullet_item(doc, "Replace [APPS_SCRIPT_WEBHOOK_URL] with the deployment URL from Step 6")
    add_bullet_item(doc, "Paste the full XML into the TwiML Bin editor")
    add_bullet_item(doc, "Click Save")

    add_h3(doc, "The TwiML — What It Does")
    add_code(doc,
        "<Response>\n"
        "  <Say>You've reached [Business Name]. We're with a customer right now —\n"
        "  please leave your name, number, and what you need, and we will call\n"
        "  you back as soon as possible. Please speak after the beep.</Say>\n"
        "  <Record maxLength=\"120\" transcribe=\"true\"\n"
        "          transcribeCallback=\"[APPS_SCRIPT_WEBHOOK_URL]\"\n"
        "          playBeep=\"true\" timeout=\"5\" />\n"
        "  <Say>Thank you — we will be in touch shortly.</Say>\n"
        "</Response>"
    )

    add_h2(doc, "Step 9 — Point the Phone Number to the TwiML Bin")
    add_bullet_item(doc, "In the Twilio Console: Phone Numbers → Manage → Active Numbers")
    add_bullet_item(doc, "Click the client's new Twilio number")
    add_bullet_item(doc, "Under Voice Configuration → A call comes in:")
    add_bullet_item(doc, "  — Set the dropdown to: TwiML Bin")
    add_bullet_item(doc, "  — Select the bin you just created from the list")
    add_bullet_item(doc, "Click Save configuration")
    add_bullet_item(doc, "Test it: call the number from your own phone — you should hear the greeting and be prompted to record")

    add_h3(doc, "Where the Voicemail Number Goes")
    add_body(doc,
        "This Twilio number is the client's public contact number — it goes on their website "
        "in the footer, the contact form, and anywhere else a phone number appears. "
        "Customers call this number. If the client wants to also receive live calls, "
        "add a Dial verb in the TwiML before the Record to forward to their real number first."
    )

    # ══════════════════════════════════════════════════════════════════════════
    # PART 6 — CONNECT TO THE WEBSITE
    # ══════════════════════════════════════════════════════════════════════════
    add_h1(doc, "Part 6 — Connect to the Website")

    add_h2(doc, "Step 10 — Paste the Webhook URL Into index.html")
    add_code(doc,
        "Find this line near the top of the <script> block:\n"
        "  const WEBHOOK_URL = 'YOUR_APPS_SCRIPT_WEBHOOK_URL_HERE';\n\n"
        "Replace with:\n"
        "  const WEBHOOK_URL = 'https://script.google.com/macros/s/AKfycb.../exec';"
    )
    add_bullet_item(doc, "Also update the phone number in the footer — replace [YOUR PHONE] with the client's Twilio number")
    add_bullet_item(doc, "Save the file → commit to GitHub → site updates within 60–90 seconds")

    # ══════════════════════════════════════════════════════════════════════════
    # PART 7 — MORNING DIGEST TRIGGER
    # ══════════════════════════════════════════════════════════════════════════
    add_h1(doc, "Part 7 — Set Up the Morning Digest Trigger")

    add_h2(doc, "Step 11 — Schedule the Daily Lead Summary")
    add_body(doc,
        "The sendMorningDigest() function emails the client a summary of new leads from the "
        "past 24 hours every morning — both form leads and voicemail leads. "
        "It only sends if there are new leads to report."
    )
    add_bullet_item(doc, "In Apps Script: click the clock icon (Triggers) in the left sidebar")
    add_bullet_item(doc, "Click Add Trigger (bottom right)")
    add_bullet_item(doc, "Function to run: sendMorningDigest")
    add_bullet_item(doc, "Event source: Time-driven")
    add_bullet_item(doc, "Type: Day timer")
    add_bullet_item(doc, "Time of day: 7am to 8am")
    add_bullet_item(doc, "Click Save — trigger is active immediately, no redeployment needed")

    # ══════════════════════════════════════════════════════════════════════════
    # PART 8 — TEST EVERYTHING
    # ══════════════════════════════════════════════════════════════════════════
    add_h1(doc, "Part 8 — Test Everything")

    add_h2(doc, "Step 12 — Run the Built-In Test Functions")
    add_body(doc,
        "The script has two test functions — one for form submissions and one for voicemail. "
        "Run each from the Apps Script editor before testing from the live site."
    )

    add_h3(doc, "Test 1 — Form Submission")
    add_bullet_item(doc, "In the function dropdown, select: testFormSubmission")
    add_bullet_item(doc, "Click Run — authorize if prompted")
    add_bullet_item(doc, "Confirm: new row in the Leads Sheet with Source = Form")
    add_bullet_item(doc, "Confirm: email arrives at NOTIFY_EMAIL with subject 'New lead: Test Lead'")
    add_bullet_item(doc, "Confirm: SMS arrives at NOTIFY_PHONE from your Twilio number")

    add_h3(doc, "Test 2 — Voicemail Transcription")
    add_bullet_item(doc, "In the function dropdown, select: testVoicemailTranscription")
    add_bullet_item(doc, "Click Run")
    add_bullet_item(doc, "Confirm: new row in the Leads Sheet with Source = Voicemail and the transcription text in Notes")
    add_bullet_item(doc, "Confirm: email arrives with subject 'New voicemail from +16305550182'")
    add_bullet_item(doc, "Confirm: SMS arrives with the truncated transcription text")

    add_h3(doc, "Test 3 — Live Voicemail")
    add_bullet_item(doc, "Call the client's Twilio number from your own phone")
    add_bullet_item(doc, "Wait for the greeting, leave a short test message")
    add_bullet_item(doc, "Wait 30–60 seconds for Twilio to transcribe and fire the callback")
    add_bullet_item(doc, "Confirm the lead appears in the Sheet, email arrives, and SMS arrives")
    add_bullet_item(doc, "Delete the test rows from the Sheet before handing off to the client")

    add_h3(doc, "Troubleshooting")
    add_table(doc, [
        ("__header__", ("Problem", "Fix")),
        ("Authorization error on first run",
         "Click Review permissions → Allow. Required once per script."),
        ("No Sheet row created",
         "Confirm you opened Apps Script from inside the correct Google Sheet via Extensions → Apps Script."),
        ("Email not received",
         "Check NOTIFY_EMAIL spelling in Script Properties. Check spam folder."),
        ("SMS not received",
         "Open Executions tab in Apps Script for errors. Most common: NOTIFY_PHONE not digits-only "
         "with country code (should be 13125550100, not +1-312-555-0100)."),
        ("Voicemail transcription never arrives",
         "Confirm the transcribeCallback in the TwiML Bin matches the Apps Script deployment URL exactly. "
         "Confirm the TwiML Bin is assigned to the phone number under Voice Configuration."),
        ("Transcription says 'unavailable'",
         "Twilio transcription failed — this happens occasionally with poor audio quality. "
         "The lead still logs with the caller's number. Upgrade to Whisper (Option 3) for better accuracy."),
        ("Twilio error: invalid phone number",
         "TWILIO_FROM_NUMBER must be +15551234567 format. NOTIFY_PHONE must be 13125550100 format."),
    ])

    # ══════════════════════════════════════════════════════════════════════════
    # QUICK REFERENCE
    # ══════════════════════════════════════════════════════════════════════════
    add_h1(doc, "Quick Reference")

    add_h2(doc, "Full Setup Checklist — New Client")
    add_table(doc, [
        ("__header__", ("Step", "Action")),
        ("Gather",   "Collect all items from Part 1 before opening any browser tab."),
        ("Sheet",    "Create '[Business Name] — Lead Log' in Google Sheets. Share with client."),
        ("Script",   "Extensions → Apps Script → paste lead-intake-webhook.gs → Save."),
        ("Props",    "Project Settings → Script Properties → add all 5 properties → Save."),
        ("Deploy",   "Deploy → New deployment → Web App → Execute as Me → Anyone → Deploy → copy URL."),
        ("Trigger",  "Triggers → Add Trigger → sendMorningDigest → Day timer → 7am–8am → Save."),
        ("TwiML",    "Twilio Console → TwiML Bins → Create → paste twiml-voicemail.xml with placeholders filled."),
        ("Number",   "Phone Numbers → Buy a number → point Voice URL to the TwiML Bin → Save."),
        ("HTML",     "Paste webhook URL into index.html as WEBHOOK_URL. Add Twilio number as phone in footer."),
        ("GitHub",   "Commit index.html → GitHub Pages deploys in 60–90 seconds."),
        ("Test",     "Run testFormSubmission() and testVoicemailTranscription() → confirm all 3 outputs each."),
        ("Live",     "Call the Twilio number → leave test message → confirm Sheet row + email + SMS."),
        ("Cleanup",  "Delete test rows from Sheet. Client starts with a clean lead log."),
    ])

    add_h2(doc, "Script Properties — Quick Reference")
    add_table(doc, [
        ("__header__", ("Property", "Value Format / Source")),
        ("NOTIFY_EMAIL",       "client@business.com — from the client"),
        ("NOTIFY_PHONE",       "13125550100 — digits + country code only, no symbols"),
        ("TWILIO_ACCOUNT_SID", "ACxxxxxxxxxxxxxxxxxxxxxxxxxxxxxxxx — Twilio Console dashboard"),
        ("TWILIO_AUTH_TOKEN",  "your auth token — Twilio Console dashboard (click eye to reveal)"),
        ("TWILIO_FROM_NUMBER", "+15551234567 — your Twilio SMS number, with + and country code"),
    ])

    add_h2(doc, "Twilio Cost Reference")
    add_table(doc, [
        ("__header__", ("Item", "Cost")),
        ("SMS sending number (yours, one-time)",  "$1.00/month — one number covers all client SMS alerts"),
        ("Voicemail number (per client)",          "$1.00/month per client"),
        ("Outbound SMS (US)",                      "$0.0079 per text"),
        ("Twilio transcription",                   "$0.05/minute of audio — most voicemails are under 1 minute"),
        ("10 clients, 30 leads + 10 voicemails/mo","~$12/mo total: $11 numbers + $3.75 SMS + $5 transcription"),
    ])

    out_path = "/Users/creighbaby/LeadManagement/docs/Lead Intake — Apps Script Setup Guide.docx"
    doc.save(out_path)
    print(f"Saved: {out_path}")


if __name__ == "__main__":
    build()
