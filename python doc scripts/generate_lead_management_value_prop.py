"""
Generates: Lead Management Subscription — Value Proposition & Service Overview.docx
A client-facing document explaining what the subscription includes,
why it matters, and how it compares to alternatives.
"""

from docx import Document
from docx.shared import Pt, Inches, RGBColor, Twips
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
import copy

# ── Color palette ──────────────────────────────────────────────
DARK_GREEN  = RGBColor(0x1a, 0x3a, 0x2a)
NEAR_BLACK  = RGBColor(0x1c, 0x1c, 0x1a)
GOLD        = "c8a96e"
CREAM_LINE  = "e2e0d8"
TABLE_HDR   = "2c1f0e"
WHITE       = RGBColor(0xFF, 0xFF, 0xFF)
CODE_BG     = "F2F2F2"

# ── Spacing helpers ─────────────────────────────────────────────
def set_para_spacing(para, before=0, after=120):
    pPr = para._p.get_or_add_pPr()
    spacing = OxmlElement('w:spacing')
    spacing.set(qn('w:before'), str(before))
    spacing.set(qn('w:after'),  str(after))
    pPr.append(spacing)

def add_bottom_border(para, color, sz, space, val='single'):
    pPr  = para._p.get_or_add_pPr()
    pBdr = OxmlElement('w:pBdr')
    bot  = OxmlElement('w:bottom')
    bot.set(qn('w:val'),   val)
    bot.set(qn('w:sz'),    str(sz))
    bot.set(qn('w:space'), str(space))
    bot.set(qn('w:color'), color)
    pBdr.append(bot)
    pPr.append(pBdr)

def set_cell_shading(cell, fill_hex):
    tc   = cell._tc
    tcPr = tc.get_or_add_tcPr()
    shd  = OxmlElement('w:shd')
    shd.set(qn('w:val'),   'clear')
    shd.set(qn('w:color'), 'auto')
    shd.set(qn('w:fill'),  fill_hex)
    tcPr.append(shd)

def set_cell_margins(cell, top=80, bottom=80, left=115, right=115):
    tc   = cell._tc
    tcPr = tc.get_or_add_tcPr()
    tcMar = OxmlElement('w:tcMar')
    for side, val in [('top', top), ('bottom', bottom), ('left', left), ('right', right)]:
        el = OxmlElement(f'w:{side}')
        el.set(qn('w:w'),    str(val))
        el.set(qn('w:type'), 'dxa')
        tcMar.append(el)
    tcPr.append(tcMar)

def set_table_borders(table, color='D0CEC6'):
    for row in table.rows:
        for cell in row.cells:
            tc   = cell._tc
            tcPr = tc.get_or_add_tcPr()
            tcBorders = OxmlElement('w:tcBorders')
            for side in ['top', 'left', 'bottom', 'right', 'insideH', 'insideV']:
                el = OxmlElement(f'w:{side}')
                el.set(qn('w:val'),   'single')
                el.set(qn('w:sz'),    '4')
                el.set(qn('w:space'), '0')
                el.set(qn('w:color'), color)
                tcBorders.append(el)
            tcPr.append(tcBorders)

# ── Content helpers ─────────────────────────────────────────────
def add_h1(doc, text):
    p   = doc.add_paragraph()
    run = p.add_run(text)
    run.bold      = True
    run.font.name = 'Arial'
    run.font.size = Pt(18)
    run.font.color.rgb = DARK_GREEN
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    set_para_spacing(p, before=400, after=200)
    add_bottom_border(p, GOLD, sz=12, space=4)
    return p

def add_h2(doc, text):
    p   = doc.add_paragraph()
    run = p.add_run(text)
    run.bold      = True
    run.font.name = 'Arial'
    run.font.size = Pt(14)
    run.font.color.rgb = DARK_GREEN
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    set_para_spacing(p, before=300, after=160)
    add_bottom_border(p, CREAM_LINE, sz=4, space=4)
    return p

def add_h3(doc, text):
    p   = doc.add_paragraph()
    run = p.add_run(text)
    run.bold      = True
    run.font.name = 'Arial'
    run.font.size = Pt(12)
    run.font.color.rgb = DARK_GREEN
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    set_para_spacing(p, before=220, after=120)
    return p

def add_body(doc, text):
    p   = doc.add_paragraph()
    run = p.add_run(text)
    run.font.name  = 'Arial'
    run.font.size  = Pt(11)
    run.font.color.rgb = NEAR_BLACK
    set_para_spacing(p, before=0, after=120)
    return p

def add_bullet(doc, text, bold_prefix=None):
    p   = doc.add_paragraph(style='List Bullet')
    if bold_prefix:
        br = p.add_run(bold_prefix)
        br.bold = True
        br.font.name = 'Arial'
        br.font.size = Pt(11)
        br.font.color.rgb = NEAR_BLACK
    run = p.add_run(text)
    run.font.name  = 'Arial'
    run.font.size  = Pt(11)
    run.font.color.rgb = NEAR_BLACK
    set_para_spacing(p, before=0, after=80)
    return p

def add_code(doc, text):
    p   = doc.add_paragraph()
    run = p.add_run(text)
    run.font.name  = 'Courier New'
    run.font.size  = Pt(9)
    run.font.color.rgb = NEAR_BLACK
    pPr = p._p.get_or_add_pPr()
    shd = OxmlElement('w:shd')
    shd.set(qn('w:val'),   'clear')
    shd.set(qn('w:color'), 'auto')
    shd.set(qn('w:fill'),  CODE_BG)
    pPr.append(shd)
    set_para_spacing(p, before=60, after=60)
    return p

def add_table(doc, rows_data, col_widths=(1.75, 4.5)):
    table = doc.add_table(rows=len(rows_data), cols=2)
    table.style = 'Table Grid'
    for i, (label, detail) in enumerate(rows_data):
        row = table.rows[i]
        lc  = row.cells[0]
        dc  = row.cells[1]
        lc.width = Inches(col_widths[0])
        dc.width = Inches(col_widths[1])
        set_cell_margins(lc)
        set_cell_margins(dc)
        if i == 0:
            set_cell_shading(lc, TABLE_HDR)
            set_cell_shading(dc, TABLE_HDR)
            lr = lc.paragraphs[0].add_run(label)
            dr = dc.paragraphs[0].add_run(detail)
            for r in [lr, dr]:
                r.bold = True
                r.font.name = 'Arial'
                r.font.size = Pt(11)
                r.font.color.rgb = WHITE
        else:
            lr = lc.paragraphs[0].add_run(label)
            lr.bold = True
            lr.font.name = 'Arial'
            lr.font.size = Pt(11)
            lr.font.color.rgb = DARK_GREEN
            dr = dc.paragraphs[0].add_run(detail)
            dr.font.name = 'Arial'
            dr.font.size = Pt(11)
            dr.font.color.rgb = NEAR_BLACK
    set_table_borders(table)
    doc.add_paragraph()
    return table

def add_wide_table(doc, rows_data, col_widths):
    """Variable-column table for comparison grids."""
    n_cols = len(col_widths)
    table  = doc.add_table(rows=len(rows_data), cols=n_cols)
    table.style = 'Table Grid'
    for i, row_data in enumerate(rows_data):
        row = table.rows[i]
        for j, cell_text in enumerate(row_data):
            cell = row.cells[j]
            cell.width = Inches(col_widths[j])
            set_cell_margins(cell, top=60, bottom=60, left=80, right=80)
            if i == 0:
                set_cell_shading(cell, TABLE_HDR)
                r = cell.paragraphs[0].add_run(cell_text)
                r.bold = True
                r.font.name = 'Arial'
                r.font.size = Pt(10)
                r.font.color.rgb = WHITE
            elif j == 0:
                r = cell.paragraphs[0].add_run(cell_text)
                r.bold = True
                r.font.name = 'Arial'
                r.font.size = Pt(10)
                r.font.color.rgb = DARK_GREEN
            else:
                r = cell.paragraphs[0].add_run(cell_text)
                r.font.name = 'Arial'
                r.font.size = Pt(10)
                r.font.color.rgb = NEAR_BLACK
    set_table_borders(table)
    doc.add_paragraph()
    return table

def add_spacer(doc):
    p = doc.add_paragraph()
    set_para_spacing(p, before=0, after=60)
    return p

# ── Build document ──────────────────────────────────────────────
def build():
    doc = Document()

    # Page margins
    for section in doc.sections:
        section.top_margin    = Inches(1)
        section.bottom_margin = Inches(1)
        section.left_margin   = Inches(1.1)
        section.right_margin  = Inches(1.1)

    # ── TITLE BLOCK ──────────────────────────────────────────────
    title = doc.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    t1 = title.add_run('LEAD MANAGEMENT SUBSCRIPTION\n')
    t1.bold = True
    t1.font.name = 'Arial'
    t1.font.size = Pt(22)
    t1.font.color.rgb = DARK_GREEN
    t2 = title.add_run('What You Get — and Why It Matters')
    t2.bold = True
    t2.font.name = 'Arial'
    t2.font.size = Pt(22)
    t2.font.color.rgb = DARK_GREEN
    set_para_spacing(title, before=0, after=160)

    sub = doc.add_paragraph()
    sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
    sr  = sub.add_run('A clear breakdown of your subscription — what runs every month, why it works, and how it compares.')
    sr.italic = True
    sr.font.name = 'Arial'
    sr.font.size = Pt(11)
    sr.font.color.rgb = NEAR_BLACK
    set_para_spacing(sub, before=0, after=200)

    # ── CONTENTS TABLE ───────────────────────────────────────────
    add_table(doc, [
        ('Section', 'Contents'),
        ('Part 1 — The Problem', 'Why most home service businesses lose leads before they even know they had one'),
        ('Part 2 — What the Subscription Includes', 'Every deliverable — all seven, running every month, on your behalf'),
        ('Part 3 — Speed to Lead: The Market Reality', 'The research behind response time and what it means for your close rate'),
        ('Part 4 — How We Compare', 'Side-by-side with GoHighLevel, Jobber, Thryv, and DIY'),
        ('Part 5 — The Numbers', 'Data that shows exactly what a missed lead costs your business'),
        ('Part 6 — Why This Works Long-Term', 'How the subscription pays for itself — and then some'),
    ])

    add_body(doc,
        'This document explains the full value of your Lead Management subscription. '
        'It is designed to give you a clear picture of what is working for your business every single month — '
        'even when you are on a job and not thinking about it.'
    )

    # ════════════════════════════════════════════════════════════
    # PART 1 — THE PROBLEM
    # ════════════════════════════════════════════════════════════
    add_h1(doc, 'Part 1 — The Problem Every Home Service Business Faces')

    add_body(doc,
        'When a homeowner needs a plumber, a landscaper, a cleaner, or a contractor, they do not call one '
        'company and wait. They fill out a form — or make a call — to two or three businesses, and they hire '
        'whoever gets back to them first.'
    )
    add_body(doc,
        'The data on this is not subtle.'
    )

    add_table(doc, [
        ('The Statistic', 'What It Means for Your Business'),
        ('78% of customers hire the first contractor who responds',
         'Coming in second place does not get you the job. Speed is the deciding factor — not price, not reviews.'),
        ('62% of home service businesses miss inbound web leads entirely',
         'Most of your competitors never even see the lead come in. They have no system.'),
        ('74% of businesses miss the 5-minute response window',
         'The conversion rate on a lead drops by up to 100x if you wait more than 5 minutes to follow up.'),
        ('The average response time in home services is 42+ hours',
         'Your competition is responding — on average — almost two days after the lead submitted. That is your window.'),
        ('Every missed lead = $300–$1,200 in lost revenue',
         'Even one missed lead per month costs more than a full year of this subscription.'),
    ])

    add_body(doc,
        'The problem is not a lack of demand. Home service businesses get inquiries. '
        'The problem is that most businesses have no reliable system to capture those inquiries and '
        'get notified instantly. That is exactly what this subscription fixes.'
    )

    # ════════════════════════════════════════════════════════════
    # PART 2 — WHAT THE SUBSCRIPTION INCLUDES
    # ════════════════════════════════════════════════════════════
    add_h1(doc, 'Part 2 — What the Subscription Includes')

    add_body(doc,
        'Everything below runs on your behalf, every month. You do not log into anything. '
        'You do not manage any software. When a lead comes in, your phone buzzes — that is the whole experience for you.'
    )

    add_h2(doc, 'Step 1 — Lead Intake Form (Always Running)')

    add_body(doc,
        'A professionally designed lead intake form lives on your website and captures every inquiry — '
        'name, phone, email, service needed, location, and project details. '
        'The form is set up and maintained by us. If you add a service, change your phone number, '
        'or want to update a question, we handle it. No charge.'
    )

    add_h2(doc, 'Step 2 — Instant Notification (Every Lead, Every Time)')

    add_body(doc,
        'The moment someone submits the form, two things happen simultaneously:'
    )
    add_bullet(doc, 'You receive an email with the full lead details — name, number, service, notes — ready to act on.', 'Email alert:  ')
    add_bullet(doc, 'You receive a text message to your phone with the same information. No app. No login. Just a text.', 'SMS alert:  ')
    add_body(doc,
        'This puts you in a position to be the first contractor to respond — every time.'
    )

    add_h2(doc, 'Step 3 — Lead Log (Your Permanent Record)')

    add_body(doc,
        'Every lead is automatically saved to a private Google Sheet — your Lead Log. '
        'This gives you a permanent, searchable history of every inquiry your business has ever received: '
        'who submitted, when, what they needed, and where they are located. '
        'If your phone dies, if an email gets lost, if you need to look back at a lead from three months ago — '
        'it is all there.'
    )

    add_h2(doc, 'Step 4 — Speed-to-Lead Facts Sheet (Included at Onboarding)')

    add_body(doc,
        'At the start of your subscription, we send you a one-page Speed-to-Lead Facts Sheet. '
        'It is not a how-to guide — it is a breakdown of the market research on lead response behavior '
        'in home services: what the data says about how quickly homeowners make hiring decisions, '
        'what your competitors are doing (and not doing), and what the numbers show about timing and close rates. '
        'Consider it background intelligence on the market you are competing in.'
    )

    add_h2(doc, 'Step 5 — Monthly Lead Summary')

    add_body(doc,
        'On the first of every month, you receive a summary from us covering:'
    )
    add_bullet(doc, 'How many leads came in during the month')
    add_bullet(doc, 'Which days and times had the most inquiries')
    add_bullet(doc, 'Any notes or observations from our side')
    add_body(doc,
        'This is your proof that the system is working. '
        'You will always know exactly what your subscription is doing for you — no guessing, no logging in anywhere.'
    )

    add_h2(doc, 'Step 6 — Google Business Profile Audit (Quarterly)')

    add_body(doc,
        'Google Business Profile is the single biggest driver of free local leads for home service businesses — '
        'it determines whether you appear when someone nearby searches for what you do. '
        'Most business owners set it up once and never revisit it. '
        'Every quarter, we review yours and flag anything that is incomplete, outdated, or costing you visibility: '
        'category accuracy, service areas, photos, hours, and Q&A. '
        'No extra charge.'
    )

    add_h2(doc, 'Step 7 — Ongoing Maintenance & Form Updates (Always Included)')

    add_body(doc,
        'Your lead intake system is not set-and-forgotten. We actively maintain it:'
    )
    add_bullet(doc, 'Service or price changes on your form — we update it, no charge.')
    add_bullet(doc, 'Technical issues — we handle them before you know there was a problem.')
    add_bullet(doc, 'Form adjustments — if we see an improvement, we make it.')

    add_spacer(doc)
    add_table(doc, [
        ('What Is Included', 'Details'),
        ('Lead intake form',            'Professional form, live on your site 24/7'),
        ('Instant email alert',         'Full lead details to your inbox the moment it submits'),
        ('Instant SMS alert',           'Text to your phone — no app, no login needed'),
        ('Google Sheet lead log',       'Permanent record of every lead, always accessible'),
        ('Speed-to-Lead Facts Sheet',   'Market research on lead response sent at onboarding'),
        ('Monthly lead summary',        'Simple report sent to you on the 1st of each month'),
        ('GBP audit (quarterly)',       'Your Google Business Profile reviewed every 3 months'),
        ('Form updates & maintenance',  'Any changes to your form or system — included, no extra charge'),
    ])

    # ════════════════════════════════════════════════════════════
    # PART 3 — SPEED TO LEAD: THE MARKET REALITY
    # ════════════════════════════════════════════════════════════
    add_h1(doc, 'Part 3 — Speed to Lead: The Market Reality')

    add_body(doc,
        'This section is not about tactics — it is about understanding the market your business operates in. '
        'The research on how homeowners hire contractors tells a consistent story, '
        'and it is worth knowing the numbers.'
    )

    add_h2(doc, 'How Homeowners Actually Make Hiring Decisions')

    add_body(doc,
        'When a homeowner needs a home service — a repair, an install, a cleanup — '
        'the typical behavior is not to research multiple companies carefully and choose the best one. '
        'It is to submit one or two inquiries and respond to whoever contacts them first. '
        'The decision often happens within minutes of that first contact, not hours.'
    )
    add_body(doc,
        'This is well-documented in the research:'
    )

    add_table(doc, [
        ('Finding', 'Source Data'),
        ('78% of customers hire the first contractor to respond',
         'Industry-wide finding across home services. Being second rarely results in a booking.'),
        ('Lead conversion drops by up to 100x after 5 minutes',
         'A lead that gets a response in under 5 minutes is dramatically more likely to convert '
         'than one that waits even 30 minutes. The window closes fast.'),
        ('The average home service business takes 42+ hours to respond to web leads',
         'Most competitors are not responding quickly. The bar is low — and that is an advantage '
         'for any business with a reliable notification system.'),
        ('74% of businesses miss the 5-minute window entirely',
         'Fewer than 1 in 4 businesses respond within the timeframe that has the highest conversion rate.'),
        ('62% of home service businesses have no web lead capture system at all',
         'The majority of competitors are not capturing web leads in any systematic way. '
         'They are relying on calls, referrals, or manual follow-up.'),
    ])

    add_h2(doc, 'What This Means in Practice')

    add_body(doc,
        'The market is set up to reward the business that is simply aware and available when a lead comes in. '
        'It is less about being the cheapest or having the most reviews — '
        'and more about being the one that shows up in the customer\'s inbox or voicemail before anyone else does.'
    )
    add_body(doc,
        'Our system is designed around this reality. '
        'The instant email and text notification exists specifically because the research shows '
        'that the time between a lead submitting a form and a business responding is where most revenue is lost. '
        'You receive the lead the moment it is submitted — the rest is in your hands.'
    )
    add_body(doc,
        'For a deeper breakdown of the data, see the Speed-to-Lead Facts Sheet included with your subscription.'
    )

    # ════════════════════════════════════════════════════════════
    # PART 4 — HOW WE COMPARE
    # ════════════════════════════════════════════════════════════
    add_h1(doc, 'Part 4 — How We Compare')

    add_body(doc,
        'There are other tools and services that claim to handle lead management for small businesses. '
        'Here is how they actually stack up.'
    )

    add_wide_table(doc, [
        ('',                       'Our Service',        'GoHighLevel',      'Jobber',           'Thryv',            'DIY (no system)'),
        ('Price / month',          '$59 – $79',          '$97+',             '$59 – $349',       '$250+',            '$0'),
        ('Who manages it',         'We do — for you',    'You do',           'You do',           'You do',           'Nobody'),
        ('Setup required',         'We handle it all',   'Extensive',        'Moderate',         'Extensive',        'You figure it out'),
        ('Lead form capture',      'Yes',                'Yes',              'Limited',          'Yes',              'Maybe'),
        ('Instant SMS alert',      'Yes',                'With setup',       'No',               'With setup',       'No'),
        ('Lead log / history',     'Yes (Google Sheet)', 'Yes (CRM)',        'Yes (CRM)',        'Yes (CRM)',        'No'),
        ('Monthly report',         'Yes — we send it',   'You pull it',      'You pull it',      'You pull it',      'No'),
        ('GBP audit (quarterly)',  'Yes — we do it',     'No',               'No',               'No',               'No'),
        ('Speed-to-lead intel',    'Yes — included',     'No',               'No',               'No',               'No'),
        ('Form updates included',  'Yes — always',       'You do it',        'You do it',        'You do it',        'N/A'),
        ('Learning curve',         'None',               'High',             'Medium',           'High',             'None (nothing works)'),
        ('Ongoing support',        'Direct — us',        'Chat/ticket',      'Chat/ticket',      'Chat/ticket',      'None'),
    ], col_widths=[1.5, 1.05, 1.05, 0.95, 0.95, 1.1])

    add_body(doc,
        'The most important column in that table is "Who manages it." '
        'Every platform except ours requires you — the business owner — to learn it, configure it, and run it. '
        'Most small business owners do not have time for that. '
        'Our service is the only one where you do nothing except receive leads.'
    )

    # ════════════════════════════════════════════════════════════
    # PART 5 — THE NUMBERS
    # ════════════════════════════════════════════════════════════
    add_h1(doc, 'Part 5 — The Numbers')

    add_body(doc,
        'The subscription pays for itself the first time it catches a lead you would have otherwise missed. '
        'Here is what the math looks like for a typical home service business.'
    )

    add_table(doc, [
        ('Scenario', 'What It Means'),
        ('You receive 8 web leads per month',
         'A realistic volume for a small home service business with an active website.'),
        ('You close 40% of qualified leads',
         'Industry average for home services when responding within 5 minutes.'),
        ('Average job value: $500',
         'Conservative for most trades — many average $800–$2,000+.'),
        ('Without this system: you miss ~5 of those 8 leads',
         'Based on industry data: 62% of businesses have no reliable capture system.'),
        ('With this system: you see all 8',
         'Every submission goes directly to your email and phone, instantly.'),
        ('Revenue difference: 3 additional jobs × $500',
         '$1,500 in additional revenue per month from leads you would have missed.'),
        ('Subscription cost',
         '$79/month — less than 5% of the revenue recovered from one missed lead.'),
    ])

    add_body(doc,
        'This is a conservative scenario. If your average job value is higher, or your response time '
        'is faster than your competitors, the return compounds. The subscription is not an expense — '
        'it is a system that pays you back every month.'
    )

    # ════════════════════════════════════════════════════════════
    # PART 6 — WHY THIS WORKS LONG-TERM
    # ════════════════════════════════════════════════════════════
    add_h1(doc, 'Part 6 — Why This Works Long-Term')

    add_body(doc,
        'A one-time website build gets your business online. '
        'The subscription keeps it actively generating and capturing business — month after month.'
    )

    add_h2(doc, 'The Compounding Effect')

    add_body(doc,
        'Every lead captured goes into your Lead Log. Over months, you build a picture of your peak '
        'inquiry days, your most-requested services, and your busiest seasons. '
        'That data helps you make smarter decisions about your business — when to hire help, '
        'when to run a promotion, which services to push.'
    )

    add_h2(doc, 'You Are Always Ahead of the Competition')

    add_body(doc,
        'Most of your local competitors have no lead capture system at all. '
        'Every month that your system runs, you are building an advantage that is hard to reverse — '
        'faster response, more jobs won, more reviews, more referrals. '
        'The businesses that move on this early tend to dominate their local market.'
    )

    add_h2(doc, 'We Grow With You')

    add_body(doc,
        'As your business grows, your lead system grows with it. '
        'New services, updated pricing, seasonal changes — we handle all of it. '
        'You focus on doing the work. We make sure the leads keep coming to you.'
    )

    add_spacer(doc)

    # Closing box (using a 1-col table as a callout)
    closing_table = doc.add_table(rows=1, cols=1)
    closing_table.style = 'Table Grid'
    cell = closing_table.rows[0].cells[0]
    set_cell_shading(cell, '1a3a2a')
    set_cell_margins(cell, top=200, bottom=200, left=200, right=200)
    set_table_borders(closing_table, color='1a3a2a')

    p1 = cell.add_paragraph()
    r1 = p1.add_run('The bottom line')
    r1.bold = True
    r1.font.name = 'Arial'
    r1.font.size = Pt(13)
    r1.font.color.rgb = RGBColor(0xc8, 0xa9, 0x6e)
    set_para_spacing(p1, before=0, after=100)

    p2 = cell.add_paragraph()
    r2 = p2.add_run(
        'Most lead management tools are software you pay for and manage yourself. '
        'We are a service. You pay us, we set everything up, we keep it running, '
        'and when someone fills out your form, your phone buzzes. '
        'No dashboard to learn. No software to maintain. '
        'Just leads in your inbox — at a fraction of what any alternative costs.'
    )
    r2.font.name = 'Arial'
    r2.font.size = Pt(11)
    r2.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
    r2.italic = True
    set_para_spacing(p2, before=0, after=0)

    doc.add_paragraph()

    # Save
    out = '/Users/creighbaby/LeadManagement/Lead Management Subscription — Value Proposition.docx'
    doc.save(out)
    print(f'Saved: {out}')

if __name__ == '__main__':
    build()
