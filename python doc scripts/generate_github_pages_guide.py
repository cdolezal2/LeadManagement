"""
GitHub Pages — Website Hosting Setup Guide
Generates a branded .docx following the LeadManagement Source-of-Truth style.
Covers: initial account setup, per-client site deployment, custom domains, and site updates.
"""

from docx import Document
from docx.shared import Pt, Inches, RGBColor, Twips
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

# ── COLOR PALETTE ──────────────────────────────────────────────────────────────
DARK_GREEN   = RGBColor(0x1a, 0x3a, 0x2a)
NEAR_BLACK   = RGBColor(0x1c, 0x1c, 0x1a)
GOLD         = "c8a96e"
CREAM_LINE   = "e2e0d8"
TABLE_HDR_BG = "2c1f0e"
WHITE        = RGBColor(0xFF, 0xFF, 0xFF)
CODE_BG      = "F2F2F2"
BORDER_COLOR = "D0CEC6"


# ── HELPERS ────────────────────────────────────────────────────────────────────

def set_para_spacing(para, before=0, after=120):
    pf = para.paragraph_format
    pf.space_before = Twips(before)
    pf.space_after  = Twips(after)


def add_bottom_border(para, color_hex, sz=12, space=4):
    pPr = para._p.get_or_add_pPr()
    pBdr = OxmlElement("w:pBdr")
    bottom = OxmlElement("w:bottom")
    bottom.set(qn("w:val"),   "single")
    bottom.set(qn("w:sz"),    str(sz))
    bottom.set(qn("w:space"), str(space))
    bottom.set(qn("w:color"), color_hex)
    pBdr.append(bottom)
    pPr.append(pBdr)


def set_cell_shading(cell, fill_hex):
    tc   = cell._tc
    tcPr = tc.get_or_add_tcPr()
    shd  = OxmlElement("w:shd")
    shd.set(qn("w:val"),   "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"),  fill_hex)
    tcPr.append(shd)


def set_cell_margins(cell, top=80, bottom=80, left=115, right=115):
    tc   = cell._tc
    tcPr = tc.get_or_add_tcPr()
    tcMar = OxmlElement("w:tcMar")
    for side, val in [("top", top), ("bottom", bottom), ("left", left), ("right", right)]:
        el = OxmlElement(f"w:{side}")
        el.set(qn("w:w"),    str(val))
        el.set(qn("w:type"), "dxa")
        tcMar.append(el)
    tcPr.append(tcMar)


def set_table_borders(table, color_hex=BORDER_COLOR):
    tbl   = table._tbl
    tblPr = tbl.find(qn("w:tblPr"))
    if tblPr is None:
        tblPr = OxmlElement("w:tblPr")
        tbl.insert(0, tblPr)
    tblBdr = OxmlElement("w:tblBorders")
    for side in ("top", "left", "bottom", "right", "insideH", "insideV"):
        el = OxmlElement(f"w:{side}")
        el.set(qn("w:val"),   "single")
        el.set(qn("w:sz"),    "4")
        el.set(qn("w:space"), "0")
        el.set(qn("w:color"), color_hex)
        tblBdr.append(el)
    tblPr.append(tblBdr)


def add_h1(doc, text):
    para = doc.add_paragraph()
    run  = para.add_run(text)
    run.bold      = True
    run.font.name = "Arial"
    run.font.size = Pt(18)
    run.font.color.rgb = DARK_GREEN
    set_para_spacing(para, before=400, after=200)
    add_bottom_border(para, GOLD, sz=12, space=4)
    return para


def add_h2(doc, text):
    para = doc.add_paragraph()
    run  = para.add_run(text)
    run.bold      = True
    run.font.name = "Arial"
    run.font.size = Pt(14)
    run.font.color.rgb = DARK_GREEN
    set_para_spacing(para, before=300, after=160)
    add_bottom_border(para, CREAM_LINE, sz=4, space=4)
    return para


def add_h3(doc, text):
    para = doc.add_paragraph()
    run  = para.add_run(text)
    run.bold      = True
    run.font.name = "Arial"
    run.font.size = Pt(12)
    run.font.color.rgb = DARK_GREEN
    set_para_spacing(para, before=220, after=120)
    return para


def add_body(doc, text):
    para = doc.add_paragraph()
    run  = para.add_run(text)
    run.font.name = "Arial"
    run.font.size = Pt(11)
    run.font.color.rgb = NEAR_BLACK
    set_para_spacing(para, before=0, after=120)
    return para


def add_bullet_item(doc, text, bold_prefix=None):
    para = doc.add_paragraph(style="List Bullet")
    if bold_prefix:
        r1 = para.add_run(bold_prefix)
        r1.bold = True
        r1.font.name = "Arial"
        r1.font.size = Pt(11)
        r1.font.color.rgb = NEAR_BLACK
    run = para.add_run(text)
    run.font.name = "Arial"
    run.font.size = Pt(11)
    run.font.color.rgb = NEAR_BLACK
    set_para_spacing(para, before=0, after=80)
    return para


def add_code(doc, text):
    para = doc.add_paragraph()
    run  = para.add_run(text)
    run.font.name = "Courier New"
    run.font.size = Pt(9)
    run.font.color.rgb = NEAR_BLACK
    set_para_spacing(para, before=60, after=60)
    pPr = para._p.get_or_add_pPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"),   "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"),  CODE_BG)
    pPr.append(shd)
    return para


def add_table(doc, rows_data, col_widths=(1.75, 4.5)):
    table = doc.add_table(rows=0, cols=2)
    set_table_borders(table)

    for i, (label, detail) in enumerate(rows_data):
        row   = table.add_row()
        left  = row.cells[0]
        right = row.cells[1]

        left.width  = Inches(col_widths[0])
        right.width = Inches(col_widths[1])

        set_cell_margins(left)
        set_cell_margins(right)

        if label == "__header__":
            set_cell_shading(left,  TABLE_HDR_BG)
            set_cell_shading(right, TABLE_HDR_BG)
            lp = left.paragraphs[0]
            lr = lp.add_run(detail[0])
            lr.bold = True; lr.font.name = "Arial"; lr.font.size = Pt(11); lr.font.color.rgb = WHITE
            rp = right.paragraphs[0]
            rr = rp.add_run(detail[1])
            rr.bold = True; rr.font.name = "Arial"; rr.font.size = Pt(11); rr.font.color.rgb = WHITE
        else:
            lp = left.paragraphs[0]
            lr = lp.add_run(label)
            lr.bold = True; lr.font.name = "Arial"; lr.font.size = Pt(11); lr.font.color.rgb = DARK_GREEN
            rp = right.paragraphs[0]
            rr = rp.add_run(detail)
            rr.font.name = "Arial"; rr.font.size = Pt(11); rr.font.color.rgb = NEAR_BLACK

    spacer = doc.add_paragraph()
    set_para_spacing(spacer, before=0, after=80)
    return table


def add_title_block(doc, line1, line2, subtitle):
    p1 = doc.add_paragraph()
    p1.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r1 = p1.add_run(line1)
    r1.bold = True; r1.font.name = "Arial"; r1.font.size = Pt(22); r1.font.color.rgb = DARK_GREEN
    set_para_spacing(p1, before=0, after=60)

    p2 = doc.add_paragraph()
    p2.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r2 = p2.add_run(line2)
    r2.bold = True; r2.font.name = "Arial"; r2.font.size = Pt(22); r2.font.color.rgb = DARK_GREEN
    set_para_spacing(p2, before=0, after=100)

    p3 = doc.add_paragraph()
    p3.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r3 = p3.add_run(subtitle)
    r3.italic = True; r3.font.name = "Arial"; r3.font.size = Pt(11); r3.font.color.rgb = NEAR_BLACK
    set_para_spacing(p3, before=0, after=300)


# ── BUILD DOCUMENT ─────────────────────────────────────────────────────────────

def build():
    doc = Document()

    for section in doc.sections:
        section.top_margin    = Inches(1.0)
        section.bottom_margin = Inches(1.0)
        section.left_margin   = Inches(1.1)
        section.right_margin  = Inches(1.1)

    # ── TITLE BLOCK
    add_title_block(
        doc,
        "GITHUB PAGES",
        "WEBSITE HOSTING GUIDE",
        "How to set up, deploy, and manage client websites using GitHub Pages — free, fast, and fully compatible with your existing setup"
    )

    # ── WHAT'S INSIDE
    add_h2(doc, "What's inside this document")
    add_table(doc, [
        ("__header__", ("Section", "Contents")),
        ("Part 1 — Why GitHub Pages",          "What it is, what it costs, and why it works for this business"),
        ("Part 2 — One-Time Account Setup",    "Creating your GitHub account and installing GitHub Desktop"),
        ("Part 3 — Deploying a Client Site",   "Step-by-step: create a repo, upload the HTML, go live"),
        ("Part 4 — Connecting a Custom Domain","How to point the client's domain to GitHub Pages"),
        ("Part 5 — Updating a Client Site",    "How to make changes and push them live"),
        ("Part 6 — Managing Multiple Clients", "Naming conventions and how to stay organized at scale"),
        ("Part 7 — Quick Reference",           "Cheat sheet of the full workflow from new client to live site"),
    ])

    add_h2(doc, "How to use this document")
    add_body(doc,
        "This guide covers everything needed to host client websites on GitHub Pages. "
        "Part 2 is done once when you first set up. Parts 3 through 5 repeat for every client. "
        "All existing site functionality — the Apps Script webhook, Google Sheets lead log, "
        "email and SMS alerts, and Calendly embed — works identically on GitHub Pages. "
        "No code changes are required when migrating from any other host."
    )

    # ══════════════════════════════════════════════════════════════════════════
    # PART 1 — WHY GITHUB PAGES
    # ══════════════════════════════════════════════════════════════════════════
    add_h1(doc, "Part 1 — Why GitHub Pages")

    add_h2(doc, "Step 1 — What GitHub Pages Is")
    add_body(doc,
        "GitHub Pages is a free static website hosting service built into GitHub. You upload an "
        "HTML file to a GitHub repository, enable Pages in the settings, and the site is live — "
        "on a public URL, with SSL/HTTPS, on GitHub's global infrastructure. There is no server "
        "to manage, no build process, and no monthly hosting fee."
    )

    add_h3(doc, "Cost Comparison")
    add_table(doc, [
        ("__header__", ("Option", "Cost")),
        ("GitHub Pages",        "Free — unlimited sites, custom domains, SSL, no credit card required"),
        ("Netlify (Pro)",       "$19/month per team member — not needed for static HTML sites"),
        ("Netlify (free tier)", "Free but with bandwidth limits and one concurrent build — unnecessary constraints"),
        ("Vercel (free)",       "Free alternative — also works well, same tradeoffs"),
        ("Client's domain",     "$10–$15/year from Namecheap or Google Domains — the only cost involved"),
    ])

    add_h3(doc, "What Works Without Any Changes")
    add_body(doc,
        "GitHub Pages only serves the HTML file. All external integrations run in the browser "
        "or on Google's servers — completely independent of the host. The following all work "
        "identically on GitHub Pages:"
    )
    add_bullet_item(doc, "Google Apps Script webhook — the form submits directly to the Apps Script URL")
    add_bullet_item(doc, "Google Sheets lead log — Apps Script writes to it on the backend")
    add_bullet_item(doc, "Email and SMS alerts — Apps Script sends these, not the host")
    add_bullet_item(doc, "Calendly embed — loads from Calendly's servers via a URL in the JS")
    add_bullet_item(doc, "Custom domain — GitHub Pages supports it natively with free SSL")

    # ══════════════════════════════════════════════════════════════════════════
    # PART 2 — ONE-TIME ACCOUNT SETUP
    # ══════════════════════════════════════════════════════════════════════════
    add_h1(doc, "Part 2 — One-Time Account Setup")

    add_h2(doc, "Step 2 — Create a GitHub Account")
    add_body(doc,
        "You only do this once. All client sites will live under your single GitHub account "
        "as separate repositories. You do not need a separate account per client."
    )
    add_bullet_item(doc, "Go to github.com and click Sign up")
    add_bullet_item(doc, "Use your business email address")
    add_bullet_item(doc, "Choose a username that represents your business — clients will never see this, but keep it professional")
    add_bullet_item(doc, "Select the Free plan — it includes everything needed")
    add_bullet_item(doc, "Verify your email address when prompted")

    add_h2(doc, "Step 3 — Install GitHub Desktop (Recommended)")
    add_body(doc,
        "GitHub Desktop is a free app that lets you manage GitHub repositories without using "
        "the command line. For uploading and updating HTML files, it is the simplest workflow. "
        "Download it at desktop.github.com and sign in with your GitHub account."
    )
    add_body(doc,
        "Alternative: You can do everything through the GitHub website (github.com) without "
        "installing anything. The web interface works well for single-file uploads. "
        "GitHub Desktop becomes more useful when you are updating sites frequently."
    )

    # ══════════════════════════════════════════════════════════════════════════
    # PART 3 — DEPLOYING A CLIENT SITE
    # ══════════════════════════════════════════════════════════════════════════
    add_h1(doc, "Part 3 — Deploying a Client Site")

    add_h2(doc, "Step 4 — Prepare the HTML File")
    add_body(doc,
        "Before creating the repository, make sure the client's HTML file is finalized and "
        "named exactly index.html — this is required. GitHub Pages looks for a file with "
        "this exact name to serve as the homepage. Any other filename will result in a "
        "download prompt instead of a rendered website."
    )

    add_h3(doc, "Pre-Launch Checklist")
    add_bullet_item(doc, "File is named index.html (lowercase, no spaces)")
    add_bullet_item(doc, "Client name, phone, and email are filled in — no placeholder text remaining")
    add_bullet_item(doc, "Apps Script webhook URL is inserted in the JavaScript")
    add_bullet_item(doc, "Calendly URL is inserted (if using the scheduling package)")
    add_bullet_item(doc, "Test the form submission once locally before uploading")

    add_h2(doc, "Step 5 — Create a GitHub Repository")
    add_body(doc,
        "A repository (repo) is GitHub's term for a project folder. Each client gets their own repo. "
        "You create a new one for each new client site."
    )
    add_bullet_item(doc, "Go to github.com and click the + icon in the top-right corner, then New repository")
    add_bullet_item(doc, "Name the repo using the client naming convention (see Part 6)")
    add_bullet_item(doc, "Set visibility to Public — GitHub Pages requires this on the free plan")
    add_bullet_item(doc, "Leave all other settings at their defaults")
    add_bullet_item(doc, "Click Create repository")

    add_h2(doc, "Step 6 — Upload the HTML File")

    add_h3(doc, "Option A — Upload via the GitHub Website (Simplest)")
    add_bullet_item(doc, "After creating the repo, click Add file → Upload files")
    add_bullet_item(doc, "Drag your index.html file into the upload area")
    add_bullet_item(doc, "Scroll down and click Commit changes")
    add_bullet_item(doc, "The file is now in the repository")

    add_h3(doc, "Option B — Upload via GitHub Desktop")
    add_bullet_item(doc, "In GitHub Desktop, click File → Clone Repository and select the new repo")
    add_bullet_item(doc, "Choose a local folder on your computer — this becomes a synced copy")
    add_bullet_item(doc, "Copy index.html into that folder")
    add_bullet_item(doc, "GitHub Desktop will show the file as a new change — add a summary and click Commit to main")
    add_bullet_item(doc, "Click Push origin to send it to GitHub")

    add_h2(doc, "Step 7 — Enable GitHub Pages")
    add_body(doc,
        "Once the HTML file is in the repo, you turn on GitHub Pages. This takes about 60 seconds."
    )
    add_bullet_item(doc, "Go to the repository on github.com")
    add_bullet_item(doc, "Click the Settings tab (top of the repo page)")
    add_bullet_item(doc, "In the left sidebar, click Pages")
    add_bullet_item(doc, "Under Branch, select main from the dropdown — leave the folder set to / (root)")
    add_bullet_item(doc, "Click Save")
    add_bullet_item(doc, "Wait 60–90 seconds, then refresh the page")
    add_bullet_item(doc, "A green banner will appear with the live URL — typically: your-username.github.io/repo-name")

    add_h3(doc, "Confirm the Site Is Live")
    add_bullet_item(doc, "Click the URL in the green banner — the site should load in the browser")
    add_bullet_item(doc, "Submit a test lead through the form to confirm the webhook is firing")
    add_bullet_item(doc, "Check the Google Sheet to confirm the lead logged correctly")
    add_bullet_item(doc, "Confirm the email and SMS alert arrived")

    # ══════════════════════════════════════════════════════════════════════════
    # PART 4 — CONNECTING A CUSTOM DOMAIN
    # ══════════════════════════════════════════════════════════════════════════
    add_h1(doc, "Part 4 — Connecting a Custom Domain")

    add_h2(doc, "Step 8 — What You Need Before Starting")
    add_body(doc,
        "The client's domain must already be purchased and accessible. The domain registrar "
        "(Namecheap, GoDaddy, Google Domains, etc.) is where the DNS settings live. "
        "You will need login access to the registrar account, or the client will need to "
        "make the DNS change themselves following your instructions."
    )

    add_h3(doc, "Domain Setup Overview")
    add_table(doc, [
        ("__header__", ("Item", "Detail")),
        ("Who buys the domain",  "The client — it stays in their name and account, always"),
        ("Recommended registrar","Namecheap ($10–$12/year) or Squarespace Domains ($12–$20/year)"),
        ("DNS change required",  "One CNAME record pointing to your GitHub Pages URL"),
        ("Time to propagate",    "Usually 5–30 minutes, up to 48 hours in rare cases"),
        ("SSL/HTTPS",            "GitHub Pages provisions this automatically — no action needed"),
    ])

    add_h2(doc, "Step 9 — Add the DNS Record at the Registrar")
    add_body(doc,
        "Log in to the client's domain registrar. Find the DNS settings for the domain "
        "(sometimes called DNS Management, Zone Editor, or Advanced DNS). "
        "Add the following record:"
    )

    add_h3(doc, "DNS Record to Add")
    add_code(doc,
        "Type:  CNAME\n"
        "Host:  www   (or @ for root domain — see note below)\n"
        "Value: your-github-username.github.io\n"
        "TTL:   Automatic (or 3600)"
    )

    add_h3(doc, "Root Domain vs. www — Which to Use")
    add_body(doc,
        "If the client wants smithplumbing.com (no www), add an A record instead of a CNAME, "
        "pointing to GitHub's IP addresses. If they want www.smithplumbing.com, a CNAME to "
        "your-username.github.io is all that is needed. Most clients are fine with www — "
        "it is simpler and propagates faster."
    )

    add_h3(doc, "GitHub's A Record IPs (for root domain only)")
    add_code(doc,
        "185.199.108.153\n"
        "185.199.109.153\n"
        "185.199.110.153\n"
        "185.199.111.153"
    )

    add_h2(doc, "Step 10 — Add the Domain in GitHub Pages Settings")
    add_bullet_item(doc, "Go to the repo → Settings → Pages")
    add_bullet_item(doc, "Under Custom domain, type the client's domain (e.g., www.smithplumbing.com)")
    add_bullet_item(doc, "Click Save")
    add_bullet_item(doc, "GitHub will check that the DNS record is in place — this may take a few minutes")
    add_bullet_item(doc, "Once verified, check Enforce HTTPS — GitHub provisions the SSL certificate automatically")
    add_bullet_item(doc, "Wait 5–30 minutes, then visit the domain to confirm the site loads")

    # ══════════════════════════════════════════════════════════════════════════
    # PART 5 — UPDATING A CLIENT SITE
    # ══════════════════════════════════════════════════════════════════════════
    add_h1(doc, "Part 5 — Updating a Client Site")

    add_h2(doc, "Step 11 — Making Changes to a Live Site")
    add_body(doc,
        "To update a client's site, you replace the index.html file in their repository "
        "with the new version. GitHub Pages will automatically redeploy within 60–90 seconds "
        "of the new file being committed. No other steps required."
    )

    add_h3(doc, "Option A — Update via the GitHub Website")
    add_bullet_item(doc, "Go to the client's repository on github.com")
    add_bullet_item(doc, "Click on index.html in the file list")
    add_bullet_item(doc, "Click the pencil (edit) icon in the top right of the file view")
    add_bullet_item(doc, "Make your changes directly in the browser editor, OR")
    add_bullet_item(doc, "Click the three-dot menu → Delete file, then re-upload the new version via Add file → Upload files")
    add_bullet_item(doc, "Scroll down and click Commit changes")
    add_bullet_item(doc, "The site will update within 60–90 seconds")

    add_h3(doc, "Option B — Update via GitHub Desktop")
    add_bullet_item(doc, "Open GitHub Desktop and select the client's repository")
    add_bullet_item(doc, "Click Fetch origin to make sure your local copy is current")
    add_bullet_item(doc, "Replace the index.html in the local folder with the updated version")
    add_bullet_item(doc, "GitHub Desktop shows the file as changed — add a brief summary (e.g., 'Update service list')")
    add_bullet_item(doc, "Click Commit to main, then Push origin")
    add_bullet_item(doc, "Site updates within 60–90 seconds")

    add_h3(doc, "Confirming the Update Deployed")
    add_body(doc,
        "Go to the repo → Actions tab. You will see a yellow circle (deploying) that turns "
        "green (complete) within about 60 seconds. Once green, the live site reflects your changes."
    )

    # ══════════════════════════════════════════════════════════════════════════
    # PART 6 — MANAGING MULTIPLE CLIENTS
    # ══════════════════════════════════════════════════════════════════════════
    add_h1(doc, "Part 6 — Managing Multiple Clients")

    add_h2(doc, "Step 12 — Naming Convention for Repositories")
    add_body(doc,
        "Every client gets one repository. Keep names consistent so you can find them instantly "
        "as the client list grows. Use this format:"
    )

    add_h3(doc, "Repository Naming Format")
    add_code(doc,
        "client-[lastname or business slug]-site\n\n"
        "Examples:\n"
        "  client-smith-plumbing-site\n"
        "  client-johnson-hvac-site\n"
        "  client-green-landscaping-site"
    )

    add_h3(doc, "What to Track Per Client")
    add_table(doc, [
        ("__header__", ("Item", "Where It Lives")),
        ("Client's HTML file",       "GitHub repo: client-[name]-site / index.html"),
        ("Client's domain",          "Their registrar account — they own it"),
        ("Apps Script webhook URL",  "Inside index.html in the JS section — also note it in your client tracker"),
        ("Google Sheet lead log",    "Shared with the client via Google Drive"),
        ("Live site URL",            "their-domain.com — confirm this is working after every update"),
    ])

    add_h2(doc, "Step 13 — Keeping Sites Organized at Scale")
    add_bullet_item(doc, "One repo per client — never put two clients' files in the same repository")
    add_bullet_item(doc, "Always test after every update — submit a test lead and confirm the Sheet logs it")
    add_bullet_item(doc, "Note the date of every change in the GitHub commit message — this is your change history")
    add_bullet_item(doc, "If a client cancels, disable GitHub Pages in the repo settings rather than deleting the repo — preserves history")

    # ══════════════════════════════════════════════════════════════════════════
    # PART 7 — QUICK REFERENCE
    # ══════════════════════════════════════════════════════════════════════════
    add_h1(doc, "Part 7 — Quick Reference")

    add_h2(doc, "Step 14 — Full Workflow: New Client to Live Site")
    add_table(doc, [
        ("__header__", ("Step", "Action")),
        ("1 — Finalize HTML",       "Name the file index.html. Fill in all placeholders. Insert webhook and Calendly URLs."),
        ("2 — Create repo",         "github.com → + → New repository → name it client-[name]-site → Public → Create"),
        ("3 — Upload file",         "Add file → Upload files → drag index.html → Commit changes"),
        ("4 — Enable Pages",        "Settings → Pages → Branch: main → / (root) → Save"),
        ("5 — Test on GitHub URL",  "Visit your-username.github.io/client-name-site → submit test lead → confirm Sheet + alerts"),
        ("6 — Add DNS at registrar","CNAME: www → your-username.github.io (or A records for root domain)"),
        ("7 — Add domain in GitHub","Settings → Pages → Custom domain → enter domain → Save → Enforce HTTPS"),
        ("8 — Final confirmation",  "Visit live domain → submit test lead → confirm everything fires correctly"),
    ])

    add_h2(doc, "Step 15 — Troubleshooting")
    add_table(doc, [
        ("__header__", ("Problem", "Fix")),
        ("Site shows 404 after enabling Pages",
         "Make sure the file is named exactly index.html (lowercase). Check Settings → Pages to confirm branch is set to main."),
        ("Custom domain shows 'not secure'",
         "Wait up to 24 hours for GitHub to provision the SSL certificate. Make sure Enforce HTTPS is checked."),
        ("Custom domain shows GitHub's 404 page",
         "DNS has not propagated yet — wait 30 minutes and try again. Use dnschecker.org to confirm the CNAME is resolving."),
        ("Form submission does nothing",
         "Open browser DevTools → Console tab and look for errors. The most common cause is a missing or incorrect webhook URL in the JS."),
        ("Site did not update after committing",
         "Go to the repo → Actions tab — look for a failed deploy (red X). If green, hard refresh the browser (Cmd+Shift+R)."),
        ("Cannot enable Pages (option is grayed out)",
         "The repository must be Public. Go to Settings → General → Danger Zone → Change visibility → Public."),
    ])

    out_path = "/Users/creighbaby/LeadManagement/docs/GitHub Pages — Website Hosting Guide.docx"
    doc.save(out_path)
    print(f"Saved: {out_path}")


if __name__ == "__main__":
    build()
