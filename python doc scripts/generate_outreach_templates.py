"""
Cold Outreach — Email & Text Templates
Generates a branded .docx with ready-to-use outreach templates for cold email,
follow-up email, text message, and voicemail drop scripts.
"""

from docx import Document
from docx.shared import Pt, Inches, RGBColor, Twips
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

# ── COLOR PALETTE ──────────────────────────────────────────────────────────────
DARK_GREEN   = RGBColor(0x1a, 0x3a, 0x2a)
NEAR_BLACK   = RGBColor(0x1c, 0x1c, 0x1a)
GOLD         = "c8a96e"
CREAM_LINE   = "e2e0d8"
TABLE_HDR_BG = "2c1f0e"
WHITE        = RGBColor(0xFF, 0xFF, 0xFF)
CODE_BG      = "F2F2F2"
BORDER_COLOR = "D0CEC6"


# ── HELPERS ────────────────────────────────────────────────────────────────────

def set_para_spacing(para, before=0, after=120):
    pf = para.paragraph_format
    pf.space_before = Twips(before)
    pf.space_after  = Twips(after)


def add_bottom_border(para, color_hex, sz=12, space=4):
    pPr = para._p.get_or_add_pPr()
    pBdr = OxmlElement("w:pBdr")
    bottom = OxmlElement("w:bottom")
    bottom.set(qn("w:val"),   "single")
    bottom.set(qn("w:sz"),    str(sz))
    bottom.set(qn("w:space"), str(space))
    bottom.set(qn("w:color"), color_hex)
    pBdr.append(bottom)
    pPr.append(pBdr)


def set_cell_shading(cell, fill_hex):
    tc   = cell._tc
    tcPr = tc.get_or_add_tcPr()
    shd  = OxmlElement("w:shd")
    shd.set(qn("w:val"),   "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"),  fill_hex)
    tcPr.append(shd)


def set_cell_margins(cell, top=80, bottom=80, left=115, right=115):
    tc   = cell._tc
    tcPr = tc.get_or_add_tcPr()
    tcMar = OxmlElement("w:tcMar")
    for side, val in [("top", top), ("bottom", bottom), ("left", left), ("right", right)]:
        el = OxmlElement(f"w:{side}")
        el.set(qn("w:w"),    str(val))
        el.set(qn("w:type"), "dxa")
        tcMar.append(el)
    tcPr.append(tcMar)


def set_table_borders(table, color_hex=BORDER_COLOR):
    tbl   = table._tbl
    tblPr = tbl.find(qn("w:tblPr"))
    if tblPr is None:
        tblPr = OxmlElement("w:tblPr")
        tbl.insert(0, tblPr)
    tblBdr = OxmlElement("w:tblBorders")
    for side in ("top", "left", "bottom", "right", "insideH", "insideV"):
        el = OxmlElement(f"w:{side}")
        el.set(qn("w:val"),   "single")
        el.set(qn("w:sz"),    "4")
        el.set(qn("w:space"), "0")
        el.set(qn("w:color"), color_hex)
        tblBdr.append(el)
    tblPr.append(tblBdr)


def add_h1(doc, text):
    para = doc.add_paragraph()
    run  = para.add_run(text)
    run.bold      = True
    run.font.name = "Arial"
    run.font.size = Pt(18)
    run.font.color.rgb = DARK_GREEN
    set_para_spacing(para, before=400, after=200)
    add_bottom_border(para, GOLD, sz=12, space=4)
    return para


def add_h2(doc, text):
    para = doc.add_paragraph()
    run  = para.add_run(text)
    run.bold      = True
    run.font.name = "Arial"
    run.font.size = Pt(14)
    run.font.color.rgb = DARK_GREEN
    set_para_spacing(para, before=300, after=160)
    add_bottom_border(para, CREAM_LINE, sz=4, space=4)
    return para


def add_h3(doc, text):
    para = doc.add_paragraph()
    run  = para.add_run(text)
    run.bold      = True
    run.font.name = "Arial"
    run.font.size = Pt(12)
    run.font.color.rgb = DARK_GREEN
    set_para_spacing(para, before=220, after=120)
    return para


def add_body(doc, text):
    para = doc.add_paragraph()
    run  = para.add_run(text)
    run.font.name = "Arial"
    run.font.size = Pt(11)
    run.font.color.rgb = NEAR_BLACK
    set_para_spacing(para, before=0, after=120)
    return para


def add_bullet_item(doc, text, bold_prefix=None):
    para = doc.add_paragraph(style="List Bullet")
    if bold_prefix:
        r1 = para.add_run(bold_prefix)
        r1.bold = True
        r1.font.name = "Arial"
        r1.font.size = Pt(11)
        r1.font.color.rgb = NEAR_BLACK
    run = para.add_run(text)
    run.font.name = "Arial"
    run.font.size = Pt(11)
    run.font.color.rgb = NEAR_BLACK
    set_para_spacing(para, before=0, after=80)
    return para


def add_code(doc, text):
    para = doc.add_paragraph()
    run  = para.add_run(text)
    run.font.name = "Courier New"
    run.font.size = Pt(9)
    run.font.color.rgb = NEAR_BLACK
    set_para_spacing(para, before=60, after=60)
    pPr = para._p.get_or_add_pPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"),   "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"),  CODE_BG)
    pPr.append(shd)
    return para


def add_table(doc, rows_data, col_widths=(1.75, 4.5)):
    table = doc.add_table(rows=0, cols=2)
    set_table_borders(table)

    for i, (label, detail) in enumerate(rows_data):
        row   = table.add_row()
        left  = row.cells[0]
        right = row.cells[1]

        left.width  = Inches(col_widths[0])
        right.width = Inches(col_widths[1])

        set_cell_margins(left)
        set_cell_margins(right)

        if label == "__header__":
            set_cell_shading(left,  TABLE_HDR_BG)
            set_cell_shading(right, TABLE_HDR_BG)

            lp = left.paragraphs[0]
            lr = lp.add_run(detail[0])
            lr.bold = True
            lr.font.name = "Arial"
            lr.font.size = Pt(11)
            lr.font.color.rgb = WHITE

            rp = right.paragraphs[0]
            rr = rp.add_run(detail[1])
            rr.bold = True
            rr.font.name = "Arial"
            rr.font.size = Pt(11)
            rr.font.color.rgb = WHITE
        else:
            lp = left.paragraphs[0]
            lr = lp.add_run(label)
            lr.bold = True
            lr.font.name = "Arial"
            lr.font.size = Pt(11)
            lr.font.color.rgb = DARK_GREEN

            rp = right.paragraphs[0]
            rr = rp.add_run(detail)
            rr.font.name = "Arial"
            rr.font.size = Pt(11)
            rr.font.color.rgb = NEAR_BLACK

    spacer = doc.add_paragraph()
    set_para_spacing(spacer, before=0, after=80)
    return table


def add_title_block(doc, line1, line2, subtitle):
    p1 = doc.add_paragraph()
    p1.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r1 = p1.add_run(line1)
    r1.bold = True
    r1.font.name = "Arial"
    r1.font.size = Pt(22)
    r1.font.color.rgb = DARK_GREEN
    set_para_spacing(p1, before=0, after=60)

    p2 = doc.add_paragraph()
    p2.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r2 = p2.add_run(line2)
    r2.bold = True
    r2.font.name = "Arial"
    r2.font.size = Pt(22)
    r2.font.color.rgb = DARK_GREEN
    set_para_spacing(p2, before=0, after=100)

    p3 = doc.add_paragraph()
    p3.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r3 = p3.add_run(subtitle)
    r3.italic = True
    r3.font.name = "Arial"
    r3.font.size = Pt(11)
    r3.font.color.rgb = NEAR_BLACK
    set_para_spacing(p3, before=0, after=300)


# ── BUILD DOCUMENT ─────────────────────────────────────────────────────────────

def build():
    doc = Document()

    for section in doc.sections:
        section.top_margin    = Inches(1.0)
        section.bottom_margin = Inches(1.0)
        section.left_margin   = Inches(1.1)
        section.right_margin  = Inches(1.1)

    # ── TITLE BLOCK
    add_title_block(
        doc,
        "COLD OUTREACH",
        "EMAIL & TEXT TEMPLATES",
        "Ready-to-use scripts for reaching prospects via email, text, and voicemail"
    )

    # ── WHAT'S INSIDE
    add_h2(doc, "What's inside this document")
    add_table(doc, [
        ("__header__", ("Section", "Contents")),
        ("Part 1 — What You're Offering",       "Two-sentence summary of each service — exactly what to say when someone asks"),
        ("Part 2 — Cold Email (Initial)",        "First outreach email — short, specific, one ask"),
        ("Part 3 — Follow-Up Email 1",           "Three days later if no response — different angle, shorter"),
        ("Part 4 — Follow-Up Email 2",           "One week later — final touch, value-add, clear exit"),
        ("Part 5 — Text Message Template",       "Ultra-short text for warm leads, referrals, or post-voicemail follow-up"),
        ("Part 6 — Voicemail Drop Script",       "What to say when they don't pick up — 20 seconds, leaves curiosity"),
        ("Part 7 — Outreach Sequence Overview",  "Day-by-day sequence: when to send what, and when to move on"),
    ])

    add_h2(doc, "How to use this document")
    add_body(doc,
        "Every template in this document is written the way a busy tradesperson reads — fast, concrete, "
        "and skimmable. Avoid the urge to add more. Shorter always outperforms longer in cold outreach. "
        "Everything in [brackets] is a variable you fill in. Personalize the business name and trade "
        "whenever possible. One clear ask per message. Never pitch more than one service at a time."
    )

    # ══════════════════════════════════════════════════════════════════════════
    # PART 1 — WHAT YOU'RE OFFERING
    # ══════════════════════════════════════════════════════════════════════════
    add_h1(doc, "Part 1 — What You're Offering")

    add_h2(doc, "Step 1 — Know How to Say It in Two Sentences")
    add_body(doc,
        "When a prospect asks 'what do you do?' you need to say it in plain language — no jargon, "
        "no tech talk, no features list. The description below is what you say. Practice it until "
        "it sounds natural, not rehearsed."
    )

    add_h3(doc, "If They Have No Website (or a Bad One) — Website + Lead Management")
    add_code(doc,
        "\"I build a professional website for your business and wire it up so every time someone "
        "fills out your contact form or calls and misses you, you get a text on your phone within "
        "minutes — their name, number, and what they need. Every lead is automatically saved so "
        "nothing falls through the cracks. You're live in under a week.\""
    )

    add_h3(doc, "If They Already Have a Website — Lead Intake + Lead Management")
    add_code(doc,
        "\"I add a proper lead capture form to your existing site and wire up a system so every "
        "inquiry — form submission or missed call — sends a text to your phone within minutes with "
        "the person's name, number, and what they need. No new website. No app to check. Just a "
        "text every time someone wants to hire you.\""
    )

    add_h3(doc, "Pricing — Say It Simply")
    add_table(doc, [
        ("__header__", ("Service", "How to Say the Price")),
        ("Website + Lead Management",
         "\"It's a $750 one-time setup and $79 a month after that. 60-day money-back guarantee on "
         "the monthly fee — if it's not working, I refund every cent of it.\""),
        ("Lead Intake + Lead Management",
         "\"It's $400 to set up and $59 a month. Same 60-day money-back guarantee. You keep "
         "everything — your leads, your Google Sheet, your data — no matter what.\""),
    ])

    # ══════════════════════════════════════════════════════════════════════════
    # PART 2 — COLD EMAIL (INITIAL)
    # ══════════════════════════════════════════════════════════════════════════
    add_h1(doc, "Part 2 — Cold Email (Initial Outreach)")

    add_h2(doc, "Step 2 — First Contact Email")
    add_body(doc,
        "This email is for a cold prospect — someone who has never heard of you. "
        "The goal is not to sell them. The goal is to get a reply or a call booked. "
        "Keep it under 100 words. One question. One ask."
    )

    add_h3(doc, "Version A — For Businesses With No Website or a Weak One")
    add_code(doc,
        "Subject: Missed calls turning into missed jobs?\n"
        "\n"
        "Hi [Name],\n"
        "\n"
        "Quick question — when someone calls [Business Name] and can't reach you, what happens to that lead?\n"
        "\n"
        "Most [plumbers / electricians / contractors] I talk to find out hours later. By then, the customer "
        "already booked someone else.\n"
        "\n"
        "I build websites for [trade] businesses in [City] that automatically text you every new lead — "
        "form submission or missed call — within minutes. Name, number, what they need. No app, no login.\n"
        "\n"
        "Worth a 15-minute call this week?\n"
        "\n"
        "[Your Name]"
    )

    add_h3(doc, "Version B — For Businesses That Already Have a Website")
    add_code(doc,
        "Subject: Is your website actually sending you leads?\n"
        "\n"
        "Hi [Name],\n"
        "\n"
        "I came across [Business Name] — you've got a solid site. Quick question: when someone fills out "
        "your contact form or calls and gets voicemail, does it automatically text you with their info?\n"
        "\n"
        "If not, you're probably losing leads without knowing it.\n"
        "\n"
        "I add a lead alert system to existing sites for [trade] businesses in [City] — every inquiry "
        "goes straight to your phone as a text, within minutes.\n"
        "\n"
        "Would a 15-minute call this week make sense?\n"
        "\n"
        "[Your Name]"
    )

    add_h3(doc, "Why These Work")
    add_bullet_item(doc, "Opens with a question about their business, not a pitch about yours.", bold_prefix="Self-referential hook: ")
    add_bullet_item(doc, "Names the specific pain — the missed call that turns into a booked competitor.", bold_prefix="Concrete problem: ")
    add_bullet_item(doc, "Describes what it does in one sentence a plumber can picture.", bold_prefix="Plain description: ")
    add_bullet_item(doc, "Asks for 15 minutes, not a decision.", bold_prefix="Low ask: ")
    add_bullet_item(doc, "No features list. No pricing. No jargon.", bold_prefix="Nothing to gloss over: ")

    # ══════════════════════════════════════════════════════════════════════════
    # PART 3 — FOLLOW-UP EMAIL 1
    # ══════════════════════════════════════════════════════════════════════════
    add_h1(doc, "Part 3 — Follow-Up Email 1 (Day 3–4)")

    add_h2(doc, "Step 3 — Second Touch If No Response")
    add_body(doc,
        "Send this 3–4 days after the initial email if you get no response. "
        "Do not re-send the original email. Come from a different angle — lead with the ROI, not the feature. "
        "Even shorter than the first email."
    )

    add_h3(doc, "Follow-Up 1 Template")
    add_code(doc,
        "Subject: Re: [Original subject]\n"
        "\n"
        "Hi [Name],\n"
        "\n"
        "Just following up on this. Didn't want it to get buried.\n"
        "\n"
        "Here's the short version: for most [plumbers] I work with, recovering even one missed lead "
        "a month more than covers the cost. The average job value in [trade] is around $[X] — my "
        "monthly fee is $59 to $79.\n"
        "\n"
        "If the timing isn't right, no problem at all. But if you want to see what it looks like, "
        "I'm happy to show you in 15 minutes.\n"
        "\n"
        "[Your Name]"
    )

    add_h3(doc, "Notes on Follow-Up 1")
    add_bullet_item(doc, "\"Just following up\" — not pushy, acknowledges they're busy.")
    add_bullet_item(doc, "Lead with ROI math, not features — a different hook than the first email.")
    add_bullet_item(doc, "Fill in [X] with the actual average job value for their trade ($600 for plumbing, $500 for HVAC, etc.).")
    add_bullet_item(doc, "Give them an easy out — \"if the timing isn't right\" — so it doesn't feel like pressure.")

    # ══════════════════════════════════════════════════════════════════════════
    # PART 4 — FOLLOW-UP EMAIL 2
    # ══════════════════════════════════════════════════════════════════════════
    add_h1(doc, "Part 4 — Follow-Up Email 2 (Day 7–10)")

    add_h2(doc, "Step 4 — Final Touch Before Moving On")
    add_body(doc,
        "Send this 7–10 days after the initial email if still no response. "
        "This is the last email in the sequence. Make it clear it's the last one — it creates urgency "
        "without being aggressive. Offer something concrete (example sites) to give them a reason to reply."
    )

    add_h3(doc, "Follow-Up 2 Template")
    add_code(doc,
        "Subject: Last one from me — a couple of examples\n"
        "\n"
        "Hi [Name],\n"
        "\n"
        "Last follow-up — I won't keep pinging you after this.\n"
        "\n"
        "I wanted to leave you with something useful either way: here are two examples of what "
        "I've built for other [trade] businesses. Both are mobile-first, load fast, and have the "
        "lead alert system built in.\n"
        "\n"
        "[Link or description of example 1]\n"
        "[Link or description of example 2]\n"
        "\n"
        "If it ever makes sense down the road, feel free to reach out. And if now works, "
        "I'm happy to hop on a quick call — just reply and we'll find a time.\n"
        "\n"
        "[Your Name]"
    )

    add_h3(doc, "Notes on Follow-Up 2")
    add_bullet_item(doc, "\"Last one from me\" — signals respect for their time and creates a mild sense of finality.")
    add_bullet_item(doc, "Offering example sites gives them a reason to click even if they're not ready to buy.")
    add_bullet_item(doc, "Leaves the door open without pressure — they can come back months later and this email still makes sense.")
    add_bullet_item(doc, "After this, mark them as cold and move on. Do not follow up a fourth time.")

    # ══════════════════════════════════════════════════════════════════════════
    # PART 5 — TEXT MESSAGE TEMPLATE
    # ══════════════════════════════════════════════════════════════════════════
    add_h1(doc, "Part 5 — Text Message Template")

    add_h2(doc, "Step 5 — When to Use a Text")
    add_body(doc,
        "Text outreach works best for warm leads — someone who was referred to you, someone who "
        "visited your site, or someone you left a voicemail for. Do not cold-text strangers without "
        "permission. Keep it to one sentence plus one question. If they reply, you're in a conversation."
    )

    add_h3(doc, "Text Template — Post-Voicemail or Warm Lead")
    add_code(doc,
        "Hi [Name] — this is [Your Name]. I left you a voicemail about getting every missed call "
        "and form submission texted to your phone automatically. Happy to explain in 2 minutes if "
        "you've got a sec — just reply here or call me back at [Your Number]."
    )

    add_h3(doc, "Text Template — Referral or Introduction")
    add_code(doc,
        "Hi [Name] — [Referral Name] suggested I reach out. I help [trade] businesses in [City] "
        "make sure every new lead goes straight to their phone as a text — forms and missed calls. "
        "Worth a quick call? — [Your Name]"
    )

    add_h3(doc, "Text Template — After Sending a Cold Email (No Response)")
    add_code(doc,
        "Hi [Name] — sent you an email last week about automating your lead alerts. "
        "Easier to ask here: would a 15-min call this week work? — [Your Name]"
    )

    add_h3(doc, "Rules for Texting")
    add_bullet_item(doc, "Never text more than once without a reply. One text, then wait.")
    add_bullet_item(doc, "Keep it under 3 sentences. If you need more, use email.")
    add_bullet_item(doc, "Always include your name at the end — they don't have your number.")
    add_bullet_item(doc, "Do not send a features list over text. One benefit, one ask.")

    # ══════════════════════════════════════════════════════════════════════════
    # PART 6 — VOICEMAIL DROP SCRIPT
    # ══════════════════════════════════════════════════════════════════════════
    add_h1(doc, "Part 6 — Voicemail Drop Script")

    add_h2(doc, "Step 6 — What to Say When They Don't Pick Up")
    add_body(doc,
        "When a prospect doesn't answer, leaving the right voicemail is the difference between a callback "
        "and a delete. Keep it under 20 seconds. Say your name, one specific thing you do, and one "
        "reason to call back. Do not leave pricing or a full pitch. Leave curiosity."
    )

    add_h3(doc, "Voicemail Script — Version A (Website + Lead Management)")
    add_code(doc,
        "\"Hey [Name], this is [Your Name] — I work with [plumbers / electricians / contractors] "
        "in [City] to make sure missed calls and form submissions go straight to their phone as a "
        "text. Takes about 15 minutes to explain. I'll send you an email too — but feel free to "
        "call me back at [Your Number]. Thanks.\""
    )

    add_h3(doc, "Voicemail Script — Version B (Lead Intake Only)")
    add_code(doc,
        "\"Hey [Name], this is [Your Name]. Quick question about your website — I help [trade] "
        "businesses in [City] make sure their site is actually sending them leads. Happy to show "
        "you what that looks like in 15 minutes. Call me back at [Your Number] or I'll shoot you "
        "an email. Thanks.\""
    )

    add_h3(doc, "Voicemail Rules")
    add_bullet_item(doc, "Under 20 seconds. If you go over, they'll skip it.", bold_prefix="Length: ")
    add_bullet_item(doc, "Always follow a voicemail with an email the same day — reference the voicemail in the subject line.", bold_prefix="Always pair with email: ")
    add_bullet_item(doc, "Do not leave more than 2 voicemails total across an outreach sequence.", bold_prefix="Limit: ")
    add_bullet_item(doc, "Speak slowly on your name and phone number — they may be writing it down.", bold_prefix="Pace: ")

    # ══════════════════════════════════════════════════════════════════════════
    # PART 7 — OUTREACH SEQUENCE OVERVIEW
    # ══════════════════════════════════════════════════════════════════════════
    add_h1(doc, "Part 7 — Outreach Sequence Overview")

    add_h2(doc, "Step 7 — Day-by-Day Sequence")
    add_body(doc,
        "Use this sequence for every new cold prospect. The goal is three touches over 10 days — "
        "email, follow-up, final email — before marking them cold. Do not drag a prospect out "
        "longer than this. Move on and come back in 60–90 days if they were not interested."
    )

    add_table(doc, [
        ("__header__", ("Day", "Action")),
        ("Day 1",   "Send cold email (Part 2). If you also called and got voicemail, leave Voicemail Version A or B (Part 6)."),
        ("Day 1",   "If you left a voicemail, also send a text the same day referencing the voicemail (Part 5, Post-Voicemail template)."),
        ("Day 3–4", "Send Follow-Up Email 1 (Part 3) if no reply. Lead with ROI math — different angle from Day 1."),
        ("Day 7–10","Send Follow-Up Email 2 (Part 4) with example sites. Tell them it's the last one."),
        ("Day 10+", "Mark as cold. Do not follow up again. Set a reminder to revisit in 60–90 days if you want."),
    ])

    add_h3(doc, "What Counts as a Response")
    add_bullet_item(doc, "Any reply to an email — even 'not interested' — is a response. Acknowledge it and move on.")
    add_bullet_item(doc, "A callback from your voicemail — treat this as a warm lead and go straight to the pitch.")
    add_bullet_item(doc, "A reply to your text — you're now in a live conversation. Stop the sequence and talk to them.")
    add_bullet_item(doc, "Booking a call via Calendly — sequence is complete. Prep for the demo call.")

    add_h3(doc, "Quick Reference — Subject Lines That Work")
    add_table(doc, [
        ("__header__", ("Email", "Subject Line")),
        ("Cold email — no website",       "Missed calls turning into missed jobs?"),
        ("Cold email — has website",      "Is your website actually sending you leads?"),
        ("Follow-up 1",                   "Re: [original subject line]"),
        ("Follow-up 2",                   "Last one from me — a couple of examples"),
        ("Post-voicemail email",          "Just left you a voicemail — [Business Name]"),
    ])

    out_path = "/Users/creighbaby/LeadManagement/docs/Cold Outreach — Email & Text Templates.docx"
    doc.save(out_path)
    print(f"Saved: {out_path}")


if __name__ == "__main__":
    build()
