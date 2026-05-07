"""
My Business — Source of Truth
Generates the master reference document for services, pricing, and client-facing materials.
Pulled from: website (my-site/index.html) + Lead Management Value Proposition doc.
"""

from docx import Document
from docx.shared import Pt, Inches, RGBColor, Twips
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

# ── COLOR PALETTE ──────────────────────────────────────────────────────────────
DARK_GREEN   = RGBColor(0x1a, 0x3a, 0x2a)
NEAR_BLACK   = RGBColor(0x1c, 0x1c, 0x1a)
GOLD         = "c8a96e"
CREAM_LINE   = "e2e0d8"
TABLE_HDR_BG = "2c1f0e"
WHITE        = RGBColor(0xFF, 0xFF, 0xFF)
CODE_BG      = "F2F2F2"
BORDER_COLOR = "D0CEC6"


# ── HELPERS ────────────────────────────────────────────────────────────────────

def set_para_spacing(para, before=0, after=120):
    pf = para.paragraph_format
    pf.space_before = Twips(before)
    pf.space_after  = Twips(after)


def add_bottom_border(para, color_hex, sz=12, space=4):
    pPr = para._p.get_or_add_pPr()
    pBdr = OxmlElement("w:pBdr")
    bottom = OxmlElement("w:bottom")
    bottom.set(qn("w:val"),   "single")
    bottom.set(qn("w:sz"),    str(sz))
    bottom.set(qn("w:space"), str(space))
    bottom.set(qn("w:color"), color_hex)
    pBdr.append(bottom)
    pPr.append(pBdr)


def set_cell_shading(cell, fill_hex):
    tc   = cell._tc
    tcPr = tc.get_or_add_tcPr()
    shd  = OxmlElement("w:shd")
    shd.set(qn("w:val"),   "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"),  fill_hex)
    tcPr.append(shd)


def set_cell_margins(cell, top=80, bottom=80, left=115, right=115):
    tc   = cell._tc
    tcPr = tc.get_or_add_tcPr()
    tcMar = OxmlElement("w:tcMar")
    for side, val in [("top", top), ("bottom", bottom), ("left", left), ("right", right)]:
        el = OxmlElement(f"w:{side}")
        el.set(qn("w:w"),    str(val))
        el.set(qn("w:type"), "dxa")
        tcMar.append(el)
    tcPr.append(tcMar)


def set_table_borders(table, color_hex=BORDER_COLOR):
    tbl   = table._tbl
    tblPr = tbl.find(qn("w:tblPr"))
    if tblPr is None:
        tblPr = OxmlElement("w:tblPr")
        tbl.insert(0, tblPr)
    tblBdr = OxmlElement("w:tblBorders")
    for side in ("top", "left", "bottom", "right", "insideH", "insideV"):
        el = OxmlElement(f"w:{side}")
        el.set(qn("w:val"),   "single")
        el.set(qn("w:sz"),    "4")
        el.set(qn("w:space"), "0")
        el.set(qn("w:color"), color_hex)
        tblBdr.append(el)
    tblPr.append(tblBdr)


def add_h1(doc, text):
    para = doc.add_paragraph()
    run  = para.add_run(text)
    run.bold      = True
    run.font.name = "Arial"
    run.font.size = Pt(18)
    run.font.color.rgb = DARK_GREEN
    set_para_spacing(para, before=400, after=200)
    add_bottom_border(para, GOLD, sz=12, space=4)
    return para


def add_h2(doc, text):
    para = doc.add_paragraph()
    run  = para.add_run(text)
    run.bold      = True
    run.font.name = "Arial"
    run.font.size = Pt(14)
    run.font.color.rgb = DARK_GREEN
    set_para_spacing(para, before=300, after=160)
    add_bottom_border(para, CREAM_LINE, sz=4, space=4)
    return para


def add_h3(doc, text):
    para = doc.add_paragraph()
    run  = para.add_run(text)
    run.bold      = True
    run.font.name = "Arial"
    run.font.size = Pt(12)
    run.font.color.rgb = DARK_GREEN
    set_para_spacing(para, before=220, after=120)
    return para


def add_body(doc, text):
    para = doc.add_paragraph()
    run  = para.add_run(text)
    run.font.name = "Arial"
    run.font.size = Pt(11)
    run.font.color.rgb = NEAR_BLACK
    set_para_spacing(para, before=0, after=120)
    return para


def add_bullet_item(doc, text, bold_prefix=None):
    para = doc.add_paragraph(style="List Bullet")
    if bold_prefix:
        r1 = para.add_run(bold_prefix)
        r1.bold = True
        r1.font.name = "Arial"
        r1.font.size = Pt(11)
        r1.font.color.rgb = NEAR_BLACK
    run = para.add_run(text)
    run.font.name = "Arial"
    run.font.size = Pt(11)
    run.font.color.rgb = NEAR_BLACK
    set_para_spacing(para, before=0, after=80)
    return para


def add_code(doc, text):
    para = doc.add_paragraph()
    run  = para.add_run(text)
    run.font.name = "Courier New"
    run.font.size = Pt(9)
    run.font.color.rgb = NEAR_BLACK
    set_para_spacing(para, before=60, after=60)
    pPr = para._p.get_or_add_pPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"),   "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"),  CODE_BG)
    pPr.append(shd)
    return para


def add_table(doc, rows_data, col_widths=(1.75, 4.5)):
    table = doc.add_table(rows=0, cols=2)
    set_table_borders(table)

    for i, (label, detail) in enumerate(rows_data):
        row   = table.add_row()
        left  = row.cells[0]
        right = row.cells[1]

        left.width  = Inches(col_widths[0])
        right.width = Inches(col_widths[1])

        set_cell_margins(left)
        set_cell_margins(right)

        if label == "__header__":
            set_cell_shading(left,  TABLE_HDR_BG)
            set_cell_shading(right, TABLE_HDR_BG)

            lp = left.paragraphs[0]
            lr = lp.add_run(detail[0])
            lr.bold = True
            lr.font.name = "Arial"
            lr.font.size = Pt(11)
            lr.font.color.rgb = WHITE

            rp = right.paragraphs[0]
            rr = rp.add_run(detail[1])
            rr.bold = True
            rr.font.name = "Arial"
            rr.font.size = Pt(11)
            rr.font.color.rgb = WHITE
        else:
            lp = left.paragraphs[0]
            lr = lp.add_run(label)
            lr.bold = True
            lr.font.name = "Arial"
            lr.font.size = Pt(11)
            lr.font.color.rgb = DARK_GREEN

            rp = right.paragraphs[0]
            rr = rp.add_run(detail)
            rr.font.name = "Arial"
            rr.font.size = Pt(11)
            rr.font.color.rgb = NEAR_BLACK

    spacer = doc.add_paragraph()
    set_para_spacing(spacer, before=0, after=80)
    return table


def add_title_block(doc, line1, line2, subtitle):
    p1 = doc.add_paragraph()
    p1.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r1 = p1.add_run(line1)
    r1.bold = True
    r1.font.name = "Arial"
    r1.font.size = Pt(22)
    r1.font.color.rgb = DARK_GREEN
    set_para_spacing(p1, before=0, after=60)

    p2 = doc.add_paragraph()
    p2.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r2 = p2.add_run(line2)
    r2.bold = True
    r2.font.name = "Arial"
    r2.font.size = Pt(22)
    r2.font.color.rgb = DARK_GREEN
    set_para_spacing(p2, before=0, after=100)

    p3 = doc.add_paragraph()
    p3.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r3 = p3.add_run(subtitle)
    r3.italic = True
    r3.font.name = "Arial"
    r3.font.size = Pt(11)
    r3.font.color.rgb = NEAR_BLACK
    set_para_spacing(p3, before=0, after=300)


# ── BUILD DOCUMENT ─────────────────────────────────────────────────────────────

def build():
    doc = Document()

    for section in doc.sections:
        section.top_margin    = Inches(1.0)
        section.bottom_margin = Inches(1.0)
        section.left_margin   = Inches(1.1)
        section.right_margin  = Inches(1.1)

    # ── TITLE BLOCK
    add_title_block(
        doc,
        "MY BUSINESS",
        "SOURCE OF TRUTH",
        "Master reference for services, pricing, deliverables, and client-facing materials"
    )

    # ── WHAT'S INSIDE
    add_h2(doc, "What's inside this document")
    add_table(doc, [
        ("__header__", ("Section", "Contents")),
        ("Part 1 — The Business",                   "What we do, who we serve, and how we position in the market"),
        ("Part 2 — Service 1: Website + Lead Mgmt", "Full definition, feature list, pricing, and delivery for the complete package"),
        ("Part 3 — Service 2: Lead Intake + Mgmt",  "Full definition, feature list, pricing, and delivery for the add-on package"),
        ("Part 4 — Shared Deliverables",            "Every component included across both services — explained in detail"),
        ("Part 5 — Pricing, Terms & Guarantee",     "Setup fees, monthly fees, payment schedule, and the 60-day guarantee"),
        ("Part 6 — Target Market",                  "Who we sell to and why this market specifically"),
        ("Part 7 — Key Selling Points",             "The four differentiators that close deals"),
        ("Part 8 — Speed-to-Lead Research",         "The market data that underpins the entire value proposition"),
        ("Part 9 — Competitive Positioning",        "How we compare against agencies, freelancers, DIY platforms, and CRMs"),
        ("Part 10 — Client ROI Framework",          "The math that makes our pricing impossible to argue with"),
    ])

    add_h2(doc, "How to use this document")
    add_body(doc,
        "This is the single source of truth for the business. When pricing, service scope, or positioning "
        "is unclear, this document is the authority. All other materials — cold call scripts, onboarding guides, "
        "service agreements, and client-facing documents — are derived from this. If something changes "
        "(a price, a feature, a policy), update this document first, then propagate changes outward."
    )

    # ══════════════════════════════════════════════════════════════════════════
    # PART 1 — THE BUSINESS
    # ══════════════════════════════════════════════════════════════════════════
    add_h1(doc, "Part 1 — The Business")

    add_h2(doc, "Step 1 — What We Do")
    add_body(doc,
        "We build professional websites and lead management systems for small service businesses. "
        "The goal is simple: make sure every customer inquiry reaches the business owner within minutes, "
        "organized and ready to act on — whether they're on a job site, driving between calls, or unavailable. "
        "We handle all the technical work. The client handles the calls."
    )

    add_h3(doc, "The One-Sentence Pitch")
    add_code(doc,
        "We build professional websites and lead management systems for small businesses — "
        "so you never miss another customer."
    )

    add_h2(doc, "Step 2 — Who We Serve")
    add_body(doc,
        "Our primary market is small home service and trade businesses. These are owner-operated "
        "companies of 1–10 people who generate most of their revenue from local inbound leads. "
        "They are almost universally behind on digital infrastructure, skeptical of agencies, "
        "and receptive to a simple, affordable solution with fast delivery and no learning curve."
    )

    add_h3(doc, "Primary Target Verticals")
    add_table(doc, [
        ("__header__", ("Vertical", "Why They're a Strong Fit")),
        ("Plumbing",             "High job values ($250–$3,500+), frequent inbound urgency, almost always missed via slow response"),
        ("HVAC",                 "Seasonal demand spikes, high ticket ($500–$8,000+), strong need for same-day lead capture"),
        ("Electrical",           "Regulated work, customers research before calling, strong web presence = trust signal"),
        ("Landscaping / Lawn",   "Repeat business, strong referral culture, website legitimizes the business"),
        ("House cleaning",       "High repeat value, price-sensitive setup but recurring revenue worth capturing"),
        ("General contracting",  "Large job values, long sales cycle — a lead log is critical to track multiple open bids"),
        ("Painting",             "Heavily quote-driven, photo-forward sites convert well, fast turnaround matters"),
        ("Roofing",              "Emergency-driven leads, $400–$12,000+ per job, speed-to-lead is the deciding factor"),
    ])

    add_h3(doc, "Secondary Target Verticals")
    add_body(doc,
        "The contact form on our website also accepts inquiries from Photography/Videography, "
        "Events/Catering, Legal/Accounting/Finance, Auto services, and Health & Wellness. "
        "These are addressable but not the primary cold outreach focus. Lead with trades."
    )

    # ══════════════════════════════════════════════════════════════════════════
    # PART 2 — SERVICE 1: WEBSITE + LEAD MANAGEMENT
    # ══════════════════════════════════════════════════════════════════════════
    add_h1(doc, "Part 2 — Service 1: Website + Lead Management")

    add_h2(doc, "Step 3 — Product Definition")
    add_body(doc,
        "Website + Lead Management is our complete, full-service offering. It is for businesses "
        "that have no website, an outdated website, or a site that is functionally useless — "
        "and for businesses that want everything handled in one place. The client gets a "
        "professional one-page website with a lead system built in from the start. "
        "Everything is connected, tested, and live before we hand it off."
    )

    add_h3(doc, "Tagline")
    add_code(doc, "Everything you need to go from invisible to booked.")

    add_h3(doc, "Pricing")
    add_table(doc, [
        ("__header__", ("Item", "Detail")),
        ("Setup fee",        "$750 — one-time, covers build, configuration, and go-live"),
        ("Monthly fee",      "$79/month — keeps the lead system running and the site maintained"),
        ("Payment schedule", "50% deposit ($375) due on signing. Balance ($375) due before go-live"),
        ("Guarantee",        "60-day money-back on all monthly fees. Setup fee is non-refundable once work begins"),
        ("Delivery",         "Live in 3–5 business days from signed agreement"),
    ])

    add_h3(doc, "Complete Feature List")
    add_body(doc, "Confirmed from the live website. This is the definitive feature set for this package.")
    add_bullet_item(doc, "Professional one-page website — mobile-first, fast-loading, built on the client's domain")
    add_bullet_item(doc, "Lead intake form — built directly into the website, captures name, phone, email, service, notes")
    add_bullet_item(doc, "Instant email alert — full lead details delivered to the client's inbox within minutes of submission")
    add_bullet_item(doc, "Instant text (SMS) alert — same lead details sent to the client's phone; no app, no login required")
    add_bullet_item(doc, "Voicemail transcription — missed calls logged as leads automatically; caller ID, transcription, and SMS alert sent to client")
    add_bullet_item(doc, "Morning lead digest — daily summary email with all active leads, organized by recency")
    add_bullet_item(doc, "Google Sheets lead log — every lead auto-saved, permanent record, accessible from any device")
    add_bullet_item(doc, "Monthly lead summary report — sent on the 1st of each month with volume, timing, and observations")
    add_bullet_item(doc, "Google Business Profile audit — full GBP review conducted quarterly, flagging any gaps in visibility")
    add_bullet_item(doc, "Speed-to-lead facts sheet — one-page market research summary sent at onboarding")
    add_bullet_item(doc, "Form updates — any changes to the intake form are always included, no extra charge")
    add_bullet_item(doc, "Ongoing site edits — service updates, price changes, photo swaps, seasonal content; always included")
    add_bullet_item(doc, "Custom domain connected — client's domain (e.g., smithplumbing.com) goes live; domain stays in client's name")

    add_h3(doc, "Who This Is For")
    add_bullet_item(doc, "Has no website, or has a website that is outdated, broken, or not mobile-optimized")
    add_bullet_item(doc, "Gets most work from word of mouth but suspects they are losing Google searches to competitors")
    add_bullet_item(doc, "Misses calls and messages regularly because they are on the job and unreachable during the day")
    add_bullet_item(doc, "Has no organized system to track who contacted them, what they quoted, or who is still waiting to hear back")
    add_bullet_item(doc, "Does not have 40+ hours to build a site themselves or $5,000+ for an agency")

    # ══════════════════════════════════════════════════════════════════════════
    # PART 3 — SERVICE 2: LEAD INTAKE + LEAD MANAGEMENT
    # ══════════════════════════════════════════════════════════════════════════
    add_h1(doc, "Part 3 — Service 2: Lead Intake + Lead Management")

    add_h2(doc, "Step 4 — Product Definition")
    add_body(doc,
        "Lead Intake + Lead Management is our standalone lead system for businesses that already "
        "have a working website. We do not touch or rebuild the site — we add a properly built "
        "lead capture form and wire up the full management backend behind it. The client goes from "
        "having a passive brochure site to an active lead engine, without rebuilding anything."
    )

    add_h3(doc, "Tagline")
    add_code(doc, "Already have a site? Put it to work.")

    add_h3(doc, "Pricing")
    add_table(doc, [
        ("__header__", ("Item", "Detail")),
        ("Setup fee",        "$400 — one-time, covers form build, integration, system configuration, and testing"),
        ("Monthly fee",      "$59/month — keeps the lead system running, maintained, and monitored"),
        ("Payment schedule", "50% deposit ($200) due on signing. Balance ($200) due before go-live"),
        ("Guarantee",        "60-day money-back on all monthly fees. Setup fee is non-refundable once work begins"),
        ("Delivery",         "Live in 2–3 business days — no website build required"),
    ])

    add_h3(doc, "Complete Feature List")
    add_body(doc, "Confirmed from the live website. This is the definitive feature set for this package.")
    add_bullet_item(doc, "Lead intake form added to the existing site — captures name, phone, email, service needed, notes")
    add_bullet_item(doc, "Instant email alert — full lead details delivered to the client's inbox within minutes of submission")
    add_bullet_item(doc, "Instant text (SMS) alert — same lead details sent to the client's phone; no app, no login required")
    add_bullet_item(doc, "Voicemail transcription — missed calls logged as leads automatically; caller ID, transcription, and SMS alert sent to client")
    add_bullet_item(doc, "Morning lead digest — daily summary email with all active leads, organized by recency")
    add_bullet_item(doc, "Google Sheets lead log — every lead auto-saved, permanent record, accessible from any device")
    add_bullet_item(doc, "Monthly lead summary report — sent on the 1st of each month with volume, timing, and observations")
    add_bullet_item(doc, "Google Business Profile audit — full GBP review conducted quarterly, flagging any gaps in visibility")
    add_bullet_item(doc, "Speed-to-lead facts sheet — one-page market research summary sent at onboarding")
    add_bullet_item(doc, "Form updates — any changes to the intake form are always included, no extra charge")

    add_h3(doc, "Who This Is For")
    add_bullet_item(doc, "Has a working website (Squarespace, Wix, WordPress, or custom) — but it has no lead system behind it")
    add_bullet_item(doc, "Has a contact form — but it just sends an unstructured email to a general inbox, untracked and easy to miss")
    add_bullet_item(doc, "Suspects they are losing leads but has no way to confirm because nothing is being logged")
    add_bullet_item(doc, "Regularly responds to inquiries hours or a full day late — after the customer has already booked someone else")
    add_bullet_item(doc, "Does not want to rebuild their website — just wants the backend to actually work")
    add_bullet_item(doc, "Already paying for SEO, Google Ads, or other traffic — and losing leads before they convert")

    # ══════════════════════════════════════════════════════════════════════════
    # PART 4 — SHARED DELIVERABLES EXPLAINED
    # ══════════════════════════════════════════════════════════════════════════
    add_h1(doc, "Part 4 — Shared Deliverables Explained")

    add_h2(doc, "Step 5 — What Every Package Includes, in Detail")
    add_body(doc,
        "The following deliverables are included in both packages. These descriptions are the "
        "authoritative definitions — use them when explaining the service to prospects, writing "
        "agreements, or training on what to deliver."
    )

    add_h3(doc, "Lead Intake Form")
    add_body(doc,
        "A professionally built form that captures structured data from every inquiry: name, phone, "
        "email, service needed, and any notes. The form is maintained by us. If the client wants "
        "to add a service, change a field, or update a question, we handle it — no charge, always included."
    )

    add_h3(doc, "Instant Email Alert")
    add_body(doc,
        "The moment someone submits the form, a formatted email alert goes to the client with the "
        "full lead details — name, number, service, notes — ready to act on. No inbox digging, "
        "no delayed notifications. This is the core speed-to-lead mechanism."
    )

    add_h3(doc, "Instant SMS Alert")
    add_body(doc,
        "Simultaneously with the email, a text message is sent to the client's phone with the same "
        "lead information. No app to download, no login required. The client gets a buzz on their "
        "phone while they are on a job — and they can call the lead back in minutes."
    )

    add_h3(doc, "Voicemail Transcription")
    add_body(doc,
        "When a customer calls the client's Twilio phone number and no one answers, Twilio records "
        "the voicemail, transcribes it automatically, and sends the transcript to the same lead system "
        "as a form submission would. The lead lands in the Google Sheet, an email alert goes to the client, "
        "and an SMS alert is sent with the caller's number and the transcription text. "
        "Missed calls become logged leads — even at 11pm on a Saturday."
    )

    add_h3(doc, "Morning Lead Digest")
    add_body(doc,
        "A daily summary email delivered each morning with all leads from the prior 24 hours, "
        "organized by recency. This gives the client a clear, prioritized call list to start the day — "
        "even if they missed an alert the night before."
    )

    add_h3(doc, "Google Sheets Lead Log")
    add_body(doc,
        "Every lead is automatically saved to a private Google Sheet in the client's own Google account. "
        "The log captures: date and time, name, phone, email, service requested, and a status column. "
        "The client owns this Sheet. If they ever leave, they take the full lead history with them. "
        "The log is also the data source for the monthly report."
    )

    add_h3(doc, "Monthly Lead Summary Report")
    add_body(doc,
        "On the first of every month, we send the client a summary covering: total leads received, "
        "peak inquiry days and times, any observations or notes from our side. This is the client's "
        "proof that the system is working — no guessing, no logging in, just a report in their inbox."
    )

    add_h3(doc, "Google Business Profile Audit (Quarterly)")
    add_body(doc,
        "Google Business Profile is the single biggest driver of free local leads for home service "
        "businesses. Most owners set it up once and never revisit it. Every quarter, we review the "
        "client's GBP and flag anything that is incomplete, outdated, or reducing visibility: "
        "category accuracy, service areas, photos, hours, Q&A, and review responses. No extra charge."
    )

    add_h3(doc, "Speed-to-Lead Facts Sheet")
    add_body(doc,
        "Sent at onboarding. A one-page summary of the market research on lead response behavior "
        "in home services: how quickly homeowners make hiring decisions, what competitors are doing, "
        "and what the data shows about timing and close rates. Not a how-to — background intelligence "
        "on the market the client is competing in."
    )

    # ══════════════════════════════════════════════════════════════════════════
    # PART 5 — PRICING, TERMS & GUARANTEE
    # ══════════════════════════════════════════════════════════════════════════
    add_h1(doc, "Part 5 — Pricing, Terms & Guarantee")

    add_h2(doc, "Step 6 — Definitive Pricing Reference")
    add_table(doc, [
        ("__header__", ("Item", "Website + Lead Mgmt / Lead Intake + Lead Mgmt")),
        ("Setup fee",          "$750 one-time  /  $400 one-time"),
        ("Monthly fee",        "$79/month  /  $59/month"),
        ("Deposit at signing", "$375 (50% of setup)  /  $200 (50% of setup)"),
        ("Balance due",        "$375 at go-live  /  $200 at go-live"),
        ("Delivery timeline",  "3–5 business days  /  2–3 business days"),
        ("Guarantee",          "60-day money-back on monthly fees — both packages"),
    ])

    add_h2(doc, "Step 7 — The 60-Day Guarantee")
    add_body(doc,
        "If the client is not satisfied in the first 60 days for any reason, every monthly payment "
        "is refunded in full — no questions asked. Setup fees are non-refundable once work begins, "
        "because work was done. Monthly fees are fully at risk until day 60. Use this on every sales "
        "call — it removes almost all of the financial risk objection and is the single most powerful "
        "closing tool we have."
    )

    add_h3(doc, "Guarantee Language for Client Communications")
    add_code(doc,
        "60-day money-back guarantee on your monthly fee. Not happy in the first 60 days? "
        "We refund every monthly payment — no questions. Setup fees are non-refundable once "
        "work begins."
    )

    # ══════════════════════════════════════════════════════════════════════════
    # PART 6 — TARGET MARKET
    # ══════════════════════════════════════════════════════════════════════════
    add_h1(doc, "Part 6 — Target Market")

    add_h2(doc, "Step 8 — The Size of the Opportunity")
    add_body(doc,
        "Small home service and trade businesses represent one of the largest and most underserved "
        "segments in the small business economy. These businesses generate real revenue, operate in "
        "high-trust local markets, and are almost universally behind on digital infrastructure. "
        "The data below defines exactly why the opportunity exists."
    )

    add_h3(doc, "The Web Presence Gap")
    add_table(doc, [
        ("__header__", ("Statistic", "Source / Notes")),
        ("~27% of small businesses have no website",
         "GoDaddy / SCORE, 2023. Highest concentration in trades and home services."),
        ("60%+ of small business websites were built more than 3 years ago",
         "Clutch Small Business Survey, 2023. Most are not mobile-optimized and have no lead capture."),
        ("97% of consumers search online to find a local business",
         "BrightLocal Local Consumer Review Survey, 2023. Google is the dominant discovery channel."),
        ("76% of local searchers visit a business within 24 hours",
         "Google / Ipsos, Think With Google. Offline action follows online search quickly."),
        ("Businesses with websites earn 2x more revenue than those without",
         "Deloitte Connected Small Business report. Applies strongly to local service businesses."),
    ])

    add_h3(doc, "The Lead Management Gap")
    add_table(doc, [
        ("__header__", ("Statistic", "Source / Notes")),
        ("Average small business response time to a new web lead: 17+ hours",
         "Drift / Lead Response Management study. Most businesses are responding when conversion odds are near zero."),
        ("78% of customers buy from the first company to respond",
         "Velocify (Salesforce subsidiary). First-mover advantage is decisive in home services."),
        ("35–50% of sales go to the vendor who responds first",
         "Inside Sales / Harvard Business Review. The race is won in the first few minutes."),
        ("Odds of qualifying a lead drop 21x after 5 minutes vs. under 5 minutes",
         "Lead Response Management study, Dr. James Oldroyd / MIT. One of the most cited stats in B2C sales research."),
        ("62% of home service businesses have no web lead capture system",
         "Industry-wide finding. Most competitors are not capturing web leads systematically."),
        ("Average home service response time: 42+ hours",
         "Combined industry data. Most competitors respond when the customer has already booked someone else."),
    ])

    # ══════════════════════════════════════════════════════════════════════════
    # PART 7 — KEY SELLING POINTS
    # ══════════════════════════════════════════════════════════════════════════
    add_h1(doc, "Part 7 — Key Selling Points")

    add_h2(doc, "Step 9 — The Four Differentiators")
    add_body(doc,
        "These are the four advantages we have over every alternative at our price point. "
        "No single competitor replicates all four simultaneously. Lead with all of them."
    )

    add_bullet_item(doc,
        "Most clients are live within 3–5 business days (website) or 2–3 business days (intake only). "
        "Agencies take 6–12 weeks. Freelancers take 4–8 weeks. DIY takes 30–60 hours of the owner's time.",
        bold_prefix="Speed — Live in under a week: ")

    add_bullet_item(doc,
        "The client's domain is in their name. Their leads live in their Google Sheet. "
        "If they ever leave, they take everything with them. No lock-in, no proprietary platform.",
        bold_prefix="Ownership — They own everything: ")

    add_bullet_item(doc,
        "The client deals with one person — always. Not a support ticket, not a rotating team. "
        "Call, text, or email whenever something needs to change.",
        bold_prefix="Direct access — One person, always: ")

    add_bullet_item(doc,
        "The goal is not a pretty website — it's more customers. Every site is wired to capture, "
        "organize, and summarize leads automatically. The system works whether the client is on a job or asleep.",
        bold_prefix="Outcome-focused — Built to convert leads: ")

    # ══════════════════════════════════════════════════════════════════════════
    # PART 8 — SPEED-TO-LEAD RESEARCH
    # ══════════════════════════════════════════════════════════════════════════
    add_h1(doc, "Part 8 — Speed-to-Lead Research")

    add_h2(doc, "Step 10 — The Research Behind the Value Proposition")
    add_body(doc,
        "Speed-to-lead is the single most important variable in whether a service business converts "
        "an inbound inquiry into a booked job. This is documented at scale across multiple independent "
        "studies. Every claim below is sourced. Use these numbers in sales conversations and client materials."
    )

    add_h3(doc, "Key Studies")
    add_table(doc, [
        ("__header__", ("Finding", "Source")),
        ("Businesses that contact a lead within 1 hour are 7x more likely to have a meaningful "
         "conversation than those who wait 2+ hours — and 60x more likely than those who wait 24 hours.",
         "Harvard Business Review analysis of 1.25M sales leads over 3 years, July 2011."),
        ("The odds of qualifying a lead drop by 21x when response time goes from under 5 minutes to 30 minutes.",
         "Lead Response Management study, Dr. James Oldroyd (MIT) and David Elkington. 15,000 leads."),
        ("78% of customers buy from the first company to respond to their inquiry.",
         "Velocify (acquired by Salesforce). Replicated across consumer service verticals including home services."),
        ("The average small business takes 17 hours to respond to a new online lead.",
         "Drift Conversational Marketing report. Most businesses respond when odds of conversion are near zero."),
        ("Increasing lead response speed from 24 hours to under 1 hour can increase conversion by 391%.",
         "Velocify study on 3.5M leads. Published as 'Optimizing Lead-to-Revenue Management.'"),
    ])

    add_h3(doc, "Plain-Language Scenario")
    add_code(doc,
        "10:47am — Customer fills out form on client's website\n"
        "10:49am — Alert hits client's phone with full lead details\n"
        "10:51am — Client calls back during a 2-minute break on the job\n"
        "10:53am — Job is booked. Competitors haven't called yet.\n"
        "4:12pm  — Competitors call back. Customer says: 'Already booked someone, thanks.'"
    )

    # ══════════════════════════════════════════════════════════════════════════
    # PART 9 — COMPETITIVE POSITIONING
    # ══════════════════════════════════════════════════════════════════════════
    add_h1(doc, "Part 9 — Competitive Positioning")

    add_h2(doc, "Step 11 — Where We Sit in the Market")
    add_body(doc,
        "We are priced below every comparable agency and above DIY platforms — deliberately. "
        "This is the zone where a trade business owner can say yes without involving a partner "
        "or a board. The comparison below maps our positioning directly against alternatives."
    )

    add_h3(doc, "Website Build — Market Price Benchmarks")
    add_table(doc, [
        ("__header__", ("Option", "Typical Cost")),
        ("Large digital agency (5+ page site)",
         "$5,000 – $15,000 setup + $200 – $500/month retainer. 6 to 12 week delivery."),
        ("Mid-tier freelancer (5-page custom site)",
         "$2,000 – $5,000 setup. Rarely includes a lead system or ongoing support. 4–8 weeks."),
        ("Small freelancer / Upwork contractor",
         "$500 – $2,000 setup. Quality highly variable. Often no maintenance included."),
        ("DIY — Squarespace or Wix",
         "$16 – $49/month. No setup cost but 30–60 hours of owner time to build. No lead backend or alerts."),
        ("Website + Lead Management — $750 setup + $79/mo",
         "Below every freelancer and agency option. Includes lead system that DIY platforms don't offer. "
         "Done in under a week with zero owner time required."),
    ])

    add_h3(doc, "Lead Management / CRM Software — Market Price Benchmarks")
    add_table(doc, [
        ("__header__", ("Option", "Typical Cost")),
        ("GoHighLevel",    "$97+/month. Extensive setup. The owner manages everything themselves."),
        ("HubSpot Starter", "$20 – $90/month. Requires setup and training. Not designed for a one-person trade business."),
        ("Jobber",          "$59 – $349/month. Good for scheduling. No lead intake form, no morning digest, no web integration."),
        ("Thryv",           "$250+/month. The owner manages everything. Steep learning curve."),
        ("ServiceTitan",    "$398 – $600+/month. Enterprise-level. Requires dedicated training. Not realistic for a solo operator."),
        ("Housecall Pro",   "$65 – $169/month. Good scheduling tool. Requires active app management. No automated lead digest."),
        ("Lead Intake + Lead Management — $400 setup + $59/mo",
         "Half the cost of most alternatives. Fully managed — the client does nothing except receive leads. "
         "No app, no dashboard, no learning curve."),
    ])

    add_h3(doc, "The Most Important Comparison Point")
    add_body(doc,
        "Every competing platform — GoHighLevel, Jobber, Thryv, HubSpot, Housecall Pro — "
        "requires the client to learn it, configure it, and run it themselves. "
        "We are a service, not software. The client pays us, we set everything up, we keep it running, "
        "and when a lead comes in, their phone buzzes. No dashboard to learn. No software to maintain."
    )

    # ══════════════════════════════════════════════════════════════════════════
    # PART 10 — CLIENT ROI FRAMEWORK
    # ══════════════════════════════════════════════════════════════════════════
    add_h1(doc, "Part 10 — Client ROI Framework")

    add_h2(doc, "Step 12 — Average Job Values by Trade")
    add_body(doc,
        "Use these figures when building ROI arguments on sales calls. "
        "All values sourced from HomeAdvisor, Angi, and Thumbtack cost guides (2024). "
        "Use the conservative figure with skeptical clients."
    )

    add_table(doc, [
        ("__header__", ("Trade", "Conservative Job Value")),
        ("Plumbing",          "$600 average (range: $250 service call to $3,500+ fixture/pipe job)"),
        ("HVAC",              "$500 average (range: $150 service call to $8,000 system replacement)"),
        ("Electrical",        "$450 average (range: $150 outlet repair to $5,000 panel upgrade)"),
        ("Landscaping",       "$400 average (range: $200 one-time to $3,000 design/install)"),
        ("House cleaning",    "$180 average (range: $100 standard to $400 deep clean)"),
        ("General contracting","$3,500 average (range: $2,000 minor renovation to $50,000+)"),
        ("Painting (interior)","$1,200 average (range: $600 per room to $5,000 full house)"),
        ("Roofing",           "$2,500 average (range: $400 repair to $12,000 replacement)"),
    ])

    add_h2(doc, "Step 13 — The ROI Argument")
    add_body(doc,
        "Tradespeople think in jobs, not monthly fees. Translate our price into jobs. "
        "One saved job per year pays for the entire service. Frame it that way every time."
    )

    add_h3(doc, "Sample ROI — Plumber, Website + Lead Management")
    add_table(doc, [
        ("__header__", ("Variable", "Figure")),
        ("Monthly fee",                  "$79/month"),
        ("Annual cost (after setup)",    "$948/year"),
        ("Conservative avg job value",   "$600"),
        ("Leads needed to break even",   "2 saved leads per year — service pays for itself"),
        ("Realistic leads saved/month",  "3–5 based on average 17-hr response time closing to under 5 min"),
        ("Monthly revenue recovered",    "3 leads × $600 = $1,800/month"),
        ("Net gain per month",           "$1,800 recovered − $79 fee = $1,721 net gain"),
        ("Annual net gain",              "~$20,600 in recovered revenue against $948 in fees"),
    ])

    add_h3(doc, "The One-Sentence ROI Close — Both Packages")
    add_code(doc,
        "Website + Lead Management: \"One saved job per year pays for everything — and you own "
        "a professional website and a lead system that keeps working whether you're on a job or asleep.\""
    )
    add_code(doc,
        "Lead Intake + Lead Management: \"You already have a website. Right now it's collecting "
        "leads and losing them. This makes sure every one of them reaches your phone within minutes "
        "— for less than the cost of one service call per month.\""
    )

    out_path = "/Users/creighbaby/LeadManagement/docs/My Business_ Source-of-Truth.docx"
    doc.save(out_path)
    print(f"Saved: {out_path}")


if __name__ == "__main__":
    build()
