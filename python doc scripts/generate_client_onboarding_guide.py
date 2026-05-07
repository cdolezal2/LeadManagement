"""
Generates: Client Onboarding — End-to-End Operations Guide.docx
Updated to reflect the two services:
  - Website + Lead Management ($750 setup / $79/mo)
  - Lead Intake + Lead Management ($400 setup / $59/mo)
"""

from docx import Document
from docx.shared import Pt, Inches, RGBColor, Twips
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

# ── COLOR PALETTE ──────────────────────────────────────────────────────────────
DARK_GREEN   = RGBColor(0x1a, 0x3a, 0x2a)
NEAR_BLACK   = RGBColor(0x1c, 0x1c, 0x1a)
GOLD         = "c8a96e"
CREAM_LINE   = "e2e0d8"
TABLE_HDR_BG = "2c1f0e"
WHITE        = RGBColor(0xFF, 0xFF, 0xFF)
CODE_BG      = "F2F2F2"
BORDER_COLOR = "D0CEC6"


# ── HELPERS ────────────────────────────────────────────────────────────────────

def set_para_spacing(para, before=0, after=120):
    pf = para.paragraph_format
    pf.space_before = Twips(before)
    pf.space_after  = Twips(after)


def add_bottom_border(para, color_hex, sz=12, space=4):
    pPr = para._p.get_or_add_pPr()
    pBdr = OxmlElement("w:pBdr")
    bottom = OxmlElement("w:bottom")
    bottom.set(qn("w:val"),   "single")
    bottom.set(qn("w:sz"),    str(sz))
    bottom.set(qn("w:space"), str(space))
    bottom.set(qn("w:color"), color_hex)
    pBdr.append(bottom)
    pPr.append(pBdr)


def set_cell_shading(cell, fill_hex):
    tc   = cell._tc
    tcPr = tc.get_or_add_tcPr()
    shd  = OxmlElement("w:shd")
    shd.set(qn("w:val"),   "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"),  fill_hex)
    tcPr.append(shd)


def set_cell_margins(cell, top=80, bottom=80, left=115, right=115):
    tc   = cell._tc
    tcPr = tc.get_or_add_tcPr()
    tcMar = OxmlElement("w:tcMar")
    for side, val in [("top", top), ("bottom", bottom), ("left", left), ("right", right)]:
        el = OxmlElement(f"w:{side}")
        el.set(qn("w:w"),    str(val))
        el.set(qn("w:type"), "dxa")
        tcMar.append(el)
    tcPr.append(tcMar)


def set_table_borders(table, color_hex=BORDER_COLOR):
    tbl   = table._tbl
    tblPr = tbl.find(qn("w:tblPr"))
    if tblPr is None:
        tblPr = OxmlElement("w:tblPr")
        tbl.insert(0, tblPr)
    tblBdr = OxmlElement("w:tblBorders")
    for side in ("top", "left", "bottom", "right", "insideH", "insideV"):
        el = OxmlElement(f"w:{side}")
        el.set(qn("w:val"),   "single")
        el.set(qn("w:sz"),    "4")
        el.set(qn("w:space"), "0")
        el.set(qn("w:color"), color_hex)
        tblBdr.append(el)
    tblPr.append(tblBdr)


def add_h1(doc, text):
    para = doc.add_paragraph()
    run  = para.add_run(text)
    run.bold = True
    run.font.name = "Arial"
    run.font.size = Pt(18)
    run.font.color.rgb = DARK_GREEN
    set_para_spacing(para, before=400, after=200)
    add_bottom_border(para, GOLD, sz=12, space=4)
    return para


def add_h2(doc, text):
    para = doc.add_paragraph()
    run  = para.add_run(text)
    run.bold = True
    run.font.name = "Arial"
    run.font.size = Pt(14)
    run.font.color.rgb = DARK_GREEN
    set_para_spacing(para, before=300, after=160)
    add_bottom_border(para, CREAM_LINE, sz=4, space=4)
    return para


def add_h3(doc, text):
    para = doc.add_paragraph()
    run  = para.add_run(text)
    run.bold = True
    run.font.name = "Arial"
    run.font.size = Pt(12)
    run.font.color.rgb = DARK_GREEN
    set_para_spacing(para, before=220, after=120)
    return para


def add_body(doc, text):
    para = doc.add_paragraph()
    run  = para.add_run(text)
    run.font.name = "Arial"
    run.font.size = Pt(11)
    run.font.color.rgb = NEAR_BLACK
    set_para_spacing(para, before=0, after=120)
    return para


def add_bullet_item(doc, text, bold_prefix=None):
    para = doc.add_paragraph(style="List Bullet")
    if bold_prefix:
        r1 = para.add_run(bold_prefix)
        r1.bold = True
        r1.font.name = "Arial"
        r1.font.size = Pt(11)
        r1.font.color.rgb = NEAR_BLACK
    run = para.add_run(text)
    run.font.name = "Arial"
    run.font.size = Pt(11)
    run.font.color.rgb = NEAR_BLACK
    set_para_spacing(para, before=0, after=80)
    return para


def add_code(doc, text):
    para = doc.add_paragraph()
    run  = para.add_run(text)
    run.font.name = "Courier New"
    run.font.size = Pt(9)
    run.font.color.rgb = NEAR_BLACK
    set_para_spacing(para, before=60, after=60)
    pPr = para._p.get_or_add_pPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"),   "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"),  CODE_BG)
    pPr.append(shd)
    return para


def add_table(doc, rows_data, col_widths=(1.75, 4.5)):
    table = doc.add_table(rows=0, cols=2)
    set_table_borders(table)
    for label, detail in rows_data:
        row   = table.add_row()
        left  = row.cells[0]
        right = row.cells[1]
        left.width  = Inches(col_widths[0])
        right.width = Inches(col_widths[1])
        set_cell_margins(left)
        set_cell_margins(right)
        if label == "__header__":
            set_cell_shading(left,  TABLE_HDR_BG)
            set_cell_shading(right, TABLE_HDR_BG)
            lp = left.paragraphs[0]
            lr = lp.add_run(detail[0])
            lr.bold = True; lr.font.name = "Arial"; lr.font.size = Pt(11); lr.font.color.rgb = WHITE
            rp = right.paragraphs[0]
            rr = rp.add_run(detail[1])
            rr.bold = True; rr.font.name = "Arial"; rr.font.size = Pt(11); rr.font.color.rgb = WHITE
        else:
            lp = left.paragraphs[0]
            lr = lp.add_run(label)
            lr.bold = True; lr.font.name = "Arial"; lr.font.size = Pt(11); lr.font.color.rgb = DARK_GREEN
            rp = right.paragraphs[0]
            rr = rp.add_run(detail)
            rr.font.name = "Arial"; rr.font.size = Pt(11); rr.font.color.rgb = NEAR_BLACK
    spacer = doc.add_paragraph()
    set_para_spacing(spacer, before=0, after=80)
    return table


def add_title_block(doc, line1, line2, subtitle):
    p1 = doc.add_paragraph()
    p1.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r1 = p1.add_run(line1)
    r1.bold = True; r1.font.name = "Arial"; r1.font.size = Pt(22); r1.font.color.rgb = DARK_GREEN
    set_para_spacing(p1, before=0, after=60)

    p2 = doc.add_paragraph()
    p2.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r2 = p2.add_run(line2)
    r2.bold = True; r2.font.name = "Arial"; r2.font.size = Pt(22); r2.font.color.rgb = DARK_GREEN
    set_para_spacing(p2, before=0, after=100)

    p3 = doc.add_paragraph()
    p3.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r3 = p3.add_run(subtitle)
    r3.italic = True; r3.font.name = "Arial"; r3.font.size = Pt(11); r3.font.color.rgb = NEAR_BLACK
    set_para_spacing(p3, before=0, after=300)


# ── BUILD DOCUMENT ─────────────────────────────────────────────────────────────

def build():
    doc = Document()

    for section in doc.sections:
        section.top_margin    = Inches(1.0)
        section.bottom_margin = Inches(1.0)
        section.left_margin   = Inches(1.1)
        section.right_margin  = Inches(1.1)

    # ── TITLE
    add_title_block(
        doc,
        "CLIENT ONBOARDING",
        "End-to-End Operations Guide",
        "From cold outreach to live service, signed contract, and recurring payment"
    )

    # ── WHAT'S INSIDE
    add_h2(doc, "What's inside this document")
    add_table(doc, [
        ("__header__", ("Section", "Contents")),
        ("Part 1 — Finding Prospects",         "How to find warm leads and make first contact"),
        ("Part 2 — Discovery Call",            "Questions to ask, what to cover, and how to determine the right service"),
        ("Part 3 — Proposal, Contract, Signing","Sending the proposal, getting the contract signed"),
        ("Part 4 — Collecting Payment",        "Setting up setup fee and monthly subscription in Stripe"),
        ("Part 5 — Website Build",             "Applies to Website + Lead Management only — building and customizing the site"),
        ("Part 6 — GitHub Setup",              "Repo structure and pushing client files — Website + Lead Management only"),
        ("Part 7 — Domain and DNS",            "Domain purchase, DNS records, and Netlify connection — Website + Lead Management only"),
        ("Part 8 — Netlify Deployment",        "Deploying the client site to Netlify — Website + Lead Management only"),
        ("Part 9 — Snippet Deployment",        "Adding the lead intake form to a client's existing site — Lead Intake + Lead Management only"),
        ("Part 10 — Lead System Setup",        "Google Sheets and Apps Script webhook — applies to both services"),
        ("Part 11 — Google Business Profile",  "Getting manager access to the client's GBP listing"),
        ("Part 12 — Client Handoff",           "Handoff email, what the client can expect, and your monthly maintenance checklist"),
    ])

    add_h2(doc, "How to use this document")
    add_body(doc,
        "Work through each part in order for every new client. Once you determine which service the client "
        "is on (Part 2), follow Parts 5–9 based on that service — not both. Parts 10–12 apply to every client "
        "regardless of which service they are on."
    )
    add_table(doc, [
        ("__header__", ("Service", "Parts to Follow")),
        ("Website + Lead Management\n$750 setup / $79 per month",
         "Parts 1–8, then Part 10–12. Skip Part 9."),
        ("Lead Intake + Lead Management\n$400 setup / $59 per month",
         "Parts 1–4, then Part 9–12. Skip Parts 5–8."),
    ])

    # ══════════════════════════════════════════════════════════════════════════
    # PART 1 — FINDING PROSPECTS
    # ══════════════════════════════════════════════════════════════════════════
    add_h1(doc, "Part 1 — Finding Prospects and Cold Outreach")

    add_body(doc,
        "Your best prospects are small businesses that either have no website, have an outdated one, "
        "or have no Google Business Profile. Target service businesses — contractors, cleaners, "
        "landscapers, salons, repair shops."
    )

    add_h2(doc, "Step 1 — Finding Leads")
    add_bullet_item(doc, "Search Google Maps for your service category + city (e.g., 'plumber Chicago') — businesses with no website listed are your warmest leads")
    add_bullet_item(doc, "Check Yelp listings — many small businesses are active there but have no standalone site")
    add_bullet_item(doc, "Drive or walk local commercial areas and note businesses without web presence")
    add_bullet_item(doc, "Ask for referrals from anyone you know who owns a small business")

    add_h2(doc, "Step 2 — Cold Call Script")
    add_body(doc, "Keep it short. Your goal is not to sell on the call — it is to book a 20-minute meeting.")
    add_code(doc,
        "Hi, this is [YOUR NAME]. I help small businesses in [CITY] get a professional website\n"
        "and a system so you never miss a lead. I noticed [BUSINESS NAME] doesn't have a\n"
        "website yet — I'd love to show you what I can put together. Do you have 20 minutes\n"
        "this week for a quick call or screen share?"
    )
    add_bullet_item(doc, "If they say they have a website: 'That's great — I also help businesses make sure their lead forms are actually working and sending inquiries to their phone. Worth a quick look?'")
    add_bullet_item(doc, "If they say not interested: 'Totally understand — can I send you a one-pager in case it becomes useful down the road?'")
    add_bullet_item(doc, "If no answer: leave a voicemail and follow up with an email same day")

    add_h2(doc, "Step 3 — Follow-Up Email")
    add_code(doc,
        "Subject: Quick website idea for [BUSINESS NAME]\n\n"
        "Hi [NAME],\n\n"
        "I left you a voicemail earlier — I help small businesses in [CITY] get a clean,\n"
        "professional website and a lead notification system set up fast.\n\n"
        "I'd love to show you a quick demo — no commitment, just 20 minutes.\n"
        "Here's a link to my site so you can see what I do: [YOUR WEBSITE URL]\n\n"
        "Would [DAY] or [DAY] work for a quick call?\n\n"
        "[YOUR NAME]\n"
        "[YOUR GOOGLE VOICE NUMBER]\n"
        "[YOUR GOOGLE WORKSPACE EMAIL]"
    )

    # ══════════════════════════════════════════════════════════════════════════
    # PART 2 — DISCOVERY CALL
    # ══════════════════════════════════════════════════════════════════════════
    add_h1(doc, "Part 2 — Discovery Call")

    add_body(doc,
        "This is a 20–30 minute call or screen share. Your goal is to understand the business, "
        "confirm fit, determine which service is the right match, and set up a follow-up with a proposal."
    )

    add_h2(doc, "Step 1 — Meeting Setup")
    add_bullet_item(doc, "Schedule via Google Calendar — send a Google Meet link (free) or Zoom link")
    add_bullet_item(doc, "Google Meet link: meet.google.com/new — generate a new link and paste into the invite")
    add_bullet_item(doc, "Use your Google Workspace email for all scheduling so it looks professional")

    add_h2(doc, "Step 2 — Questions to Ask")
    add_bullet_item(doc, "What does your business do and who are your best customers?")
    add_bullet_item(doc, "How do people currently find you? (word of mouth, Google, Yelp, referrals?)")
    add_bullet_item(doc, "Do you have a website or Google Business Profile today?")
    add_bullet_item(doc, "When someone contacts you — how do they reach you and how do you track that?")
    add_bullet_item(doc, "What would a lead management system do for your business?")
    add_bullet_item(doc, "What's your timeline — are you looking to move quickly or just exploring?")

    add_h2(doc, "Step 3 — Determine the Right Service")
    add_body(doc,
        "The discovery call is where you figure out which service fits. Use these signals to guide the "
        "recommendation — then present it as a clear suggestion, not a menu of options."
    )
    add_table(doc, [
        ("__header__", ("Signal", "Right Service")),
        ("No website, weak website, or outdated site",
         "Website + Lead Management — they need a full rebuild and the lead system together"),
        ("Decent website already in place",
         "Lead Intake + Lead Management — add the form and notification system to what they have"),
        ("On the fence about rebuilding their site",
         "Lead Intake + Lead Management — lower commitment, solve the immediate problem first"),
        ("Specifically asking for a new website",
         "Website + Lead Management — confirm scope and set timeline expectations"),
    ])

    add_h2(doc, "Step 4 — What to Cover on Your End")
    add_bullet_item(doc, "Show your own website briefly — walk them through what you build")
    add_bullet_item(doc, "Explain the two services and which one you are recommending and why")
    add_bullet_item(doc, "Walk through what the lead system does — instant email + text alerts, Google Sheets log, monthly report, GBP audit")
    add_bullet_item(doc, "Give a rough timeline: live in under a week once you have their info")
    add_bullet_item(doc, "Tell them you will send a proposal after the call")
    add_bullet_item(doc, "Do NOT close on price during this call — send a written proposal instead")

    # ══════════════════════════════════════════════════════════════════════════
    # PART 3 — PROPOSAL, CONTRACT, SIGNING
    # ══════════════════════════════════════════════════════════════════════════
    add_h1(doc, "Part 3 — Proposal, Contract, and Signing")

    add_h2(doc, "Step 1 — Send the Proposal")
    add_body(doc,
        "Send a simple proposal via email within 24 hours of the discovery call. Name the specific "
        "service you are recommending, what is included, and next steps. Do not present both services "
        "as options — recommend one based on what you learned on the call."
    )
    add_code(doc,
        "Subject: Proposal for [BUSINESS NAME] — [Service Name]\n\n"
        "Hi [NAME],\n\n"
        "Great talking with you. Based on what you shared, I'd recommend the\n"
        "[Website + Lead Management / Lead Intake + Lead Management] service.\n\n"
        "Here's what's included:\n"
        "  [list the features for the service]\n\n"
        "Setup fee: $[750 or 400] (50% due on signing, balance at go-live)\n"
        "Monthly: $[69 or 59]/month — billed on the 1st via Stripe\n"
        "Timeline: live in 3–5 business days from when I have your info\n\n"
        "I'll send the service agreement next — let me know if you have any questions.\n\n"
        "[YOUR NAME]"
    )

    add_h2(doc, "Step 2 — The Contract")
    add_bullet_item(doc, "Use the Client Service Agreement — Template.docx from your docs folder")
    add_bullet_item(doc, "Fill in both parties' information, the service name, setup fee, and monthly amount")
    add_bullet_item(doc, "Send as a PDF — or use DocuSign ($10/month) for e-signature")
    add_bullet_item(doc, "Send the contract before collecting any payment")

    add_h2(doc, "Step 3 — Get the Signed Contract Back")
    add_bullet_item(doc, "Client signs and returns the contract (PDF or DocuSign completion email)")
    add_bullet_item(doc, "Save a copy to your client folder — create one folder per client in Google Drive")
    add_bullet_item(doc, "Do not begin work until contract is signed")

    # ══════════════════════════════════════════════════════════════════════════
    # PART 4 — PAYMENT
    # ══════════════════════════════════════════════════════════════════════════
    add_h1(doc, "Part 4 — Collecting Payment with Stripe")

    add_body(doc,
        "Collect the one-time setup fee before starting work. Set up the recurring monthly subscription "
        "at the same time so the client only has to enter their card once."
    )

    add_h2(doc, "Step 1 — Create a Payment Link for the Setup Fee")
    add_bullet_item(doc, "Log into Stripe Dashboard → Payment Links → Create new link")
    add_bullet_item(doc, "Set the product name to match the service — e.g., 'Website + Lead Management — Setup Fee' or 'Lead Intake + Lead Management — Setup Fee'")
    add_bullet_item(doc, "Set the price ($750 or $400) and quantity to 1")
    add_bullet_item(doc, "Copy the payment link and paste it into your proposal or a follow-up email")
    add_bullet_item(doc, "The client clicks the link, enters their card, and you receive a confirmation email")

    add_h2(doc, "Step 2 — Set Up the Monthly Subscription")
    add_bullet_item(doc, "In Stripe Dashboard → Subscriptions → + New subscription")
    add_bullet_item(doc, "Enter the client's email — Stripe will send them a hosted invoice to add their card")
    add_bullet_item(doc, "Set the product name — e.g., 'Website + Lead Management — Monthly' or 'Lead Intake + Lead Management — Monthly'")
    add_bullet_item(doc, "Set the monthly amount ($79 or $59), billing cycle to monthly, starting on the 1st of next month")
    add_bullet_item(doc, "Stripe will automatically charge and send receipts each month")

    add_h2(doc, "Step 3 — Confirm Payment Before Starting Work")
    add_bullet_item(doc, "Check Stripe Dashboard to confirm setup fee payment has cleared")
    add_bullet_item(doc, "Stripe payouts typically land in your bank account within 2 business days")
    add_bullet_item(doc, "Once setup fee is confirmed — begin the build")

    # ══════════════════════════════════════════════════════════════════════════
    # PART 5 — WEBSITE BUILD (Website + Lead Management only)
    # ══════════════════════════════════════════════════════════════════════════
    add_h1(doc, "Part 5 — Building the Client Website")

    add_body(doc,
        "Applies to: Website + Lead Management only. "
        "If the client is on Lead Intake + Lead Management, skip to Part 9."
    )

    add_h2(doc, "Step 1 — Gather Client Information")
    add_bullet_item(doc, "Business name, tagline, and a one-sentence description of what they do")
    add_bullet_item(doc, "Phone number, email, and service area or address")
    add_bullet_item(doc, "Services offered — get at least 3–5 specific services to list on the site")
    add_bullet_item(doc, "Any brand colors, logo files, or photos they want to use")
    add_bullet_item(doc, "Their domain name (purchased in Part 7) or the domain they want")

    add_h2(doc, "Step 2 — Customize the HTML Template")
    add_bullet_item(doc, "Open the Website + Lead Management base template from your clients/templates/ folder in the GitHub repo")
    add_bullet_item(doc, "Find and replace all placeholders: [CLIENT_NAME], [CLIENT_PHONE], [CLIENT_EMAIL], [CLIENT_CITY], [CLIENT_SERVICES]")
    add_bullet_item(doc, "Update the WEBHOOK_URL constant at the top with the Apps Script URL (set up in Part 10)")
    add_bullet_item(doc, "Adjust colors in the :root CSS block if the client has brand colors")
    add_bullet_item(doc, "Update the page title and meta description with the client's business name")
    add_bullet_item(doc, "Preview in a browser before pushing to GitHub")

    # ══════════════════════════════════════════════════════════════════════════
    # PART 6 — GITHUB (Website + Lead Management only)
    # ══════════════════════════════════════════════════════════════════════════
    add_h1(doc, "Part 6 — GitHub Setup")

    add_body(doc,
        "Applies to: Website + Lead Management only. "
        "All client websites live in one GitHub repository under a clients/ folder. "
        "Each client gets their own subfolder. This keeps everything organized and makes "
        "Netlify deployment per-client simple."
    )

    add_h2(doc, "Step 1 — Repo Structure")
    add_code(doc,
        "LeadManagement/\n"
        "  my-site/\n"
        "    index.html          \u2190 your own business website\n"
        "  clients/\n"
        "    client-name/\n"
        "      index.html        \u2190 each client's site\n"
        "    another-client/\n"
        "      index.html"
    )

    add_h2(doc, "Step 2 — Add the Client Folder")
    add_bullet_item(doc, "Open your LeadManagement repo folder on your computer")
    add_bullet_item(doc, "Inside the clients/ folder, create a new folder named after the client (lowercase, no spaces — e.g., rivera-plumbing)")
    add_bullet_item(doc, "Copy your base index.html template into that folder")
    add_bullet_item(doc, "Make your customizations (see Part 5)")

    add_h2(doc, "Step 3 — Push to GitHub")
    add_code(doc,
        "# In terminal from your LeadManagement folder:\n"
        "git add clients/rivera-plumbing/index.html\n"
        "git commit -m \"Add Rivera Plumbing client site\"\n"
        "git push"
    )
    add_bullet_item(doc, "Or use GitHub Desktop: stage the new file, write a commit message, click Commit, then Push")
    add_bullet_item(doc, "Verify the file appears in the GitHub repo before moving to Netlify setup")

    # ══════════════════════════════════════════════════════════════════════════
    # PART 7 — DOMAIN (Website + Lead Management only)
    # ══════════════════════════════════════════════════════════════════════════
    add_h1(doc, "Part 7 — Domain Purchase and DNS Setup")

    add_body(doc,
        "Applies to: Website + Lead Management only. "
        "The client purchases and owns their own domain. You help them pick it and walk them "
        "through the DNS settings needed to point it at Netlify."
    )

    add_body(doc,
        "There are two domain options — both are included at no extra cost. Ask the client on the "
        "discovery call which they prefer. Either way, you handle all the technical setup."
    )
    add_table(doc, [
        ("__header__", ("Option", "How It Works")),
        ("Option A — Client Owns the Domain",
         "Client purchases their own domain before the onboarding call. You connect it to Netlify on the call. "
         "They own and control it at all times. If they cancel, nothing to transfer."),
        ("Option B — Fully Managed (You Own the Domain)",
         "Client provides nothing — just their email and phone. You purchase and manage the domain on their behalf. "
         "No action required from the client at onboarding. If they cancel, you transfer the domain to them."),
    ])

    add_h2(doc, "Step 1 — Option A: Client Owns the Domain")

    add_h3(doc, "What to Send the Client Before the Onboarding Call")
    add_code(doc,
        "Before our call, go to Namecheap.com, create a free account, and search for\n"
        "[yourbusinessname].com. If it's available, purchase it — it's about $12 for the year.\n"
        "Keep your login handy. I'll take care of everything else on our call."
    )

    add_h3(doc, "Client Does (Before the Call)")
    add_bullet_item(doc, "Create a Namecheap account at namecheap.com")
    add_bullet_item(doc, "Search for and purchase their domain — ideal format: businessname.com or businessnamecity.com")
    add_bullet_item(doc, "Avoid hyphens and numbers — hard to say on the phone")
    add_bullet_item(doc, "Have their Namecheap login ready for the onboarding call")

    add_h3(doc, "You Do (On the Onboarding Call — Screen Share)")
    add_bullet_item(doc, "Ask the client to log into Namecheap and share their screen, or log in for them if they hand over credentials")
    add_bullet_item(doc, "Go to Domain List → Manage → Advanced DNS tab")
    add_bullet_item(doc, "Delete any default A records or CNAME records Namecheap added — they will conflict")
    add_bullet_item(doc, "Add A Record: Host = @  |  Value = 75.2.60.5  |  TTL = Automatic")
    add_bullet_item(doc, "Add CNAME Record: Host = www  |  Value = [their-netlify-site].netlify.app  |  TTL = Automatic")
    add_bullet_item(doc, "In Netlify → Site Settings → Domain Management → Add custom domain → enter their .com → Verify")
    add_bullet_item(doc, "Click Provision certificate — Netlify provisions HTTPS automatically (Let's Encrypt)")
    add_bullet_item(doc, "Site is fully live once the green padlock appears — usually within minutes to a few hours")

    add_h2(doc, "Step 2 — Option B: Fully Managed — You Own the Domain")

    add_h3(doc, "Client Does")
    add_bullet_item(doc, "Nothing — just provide their business name, phone number, and email address")
    add_bullet_item(doc, "No Namecheap account, no DNS settings, no action required at all")

    add_h3(doc, "You Do (Before the Onboarding Call)")
    add_bullet_item(doc, "Search Namecheap for the best available domain for their business — businessname.com preferred")
    add_bullet_item(doc, "Purchase the domain from your own Namecheap account (~$12/year) — keep it in your account")
    add_bullet_item(doc, "Enable auto-renewal on the domain so it never accidentally expires")
    add_bullet_item(doc, "Add a recurring annual reminder in your calendar: '[Client] domain renewal — [month/year]'")
    add_bullet_item(doc, "Note the domain in the client's folder in Google Drive")

    add_h3(doc, "You Do (DNS — Same as Option A, Just From Your Account)")
    add_bullet_item(doc, "In your Namecheap account → Domain List → Manage → Advanced DNS tab")
    add_bullet_item(doc, "Delete any default records Namecheap added")
    add_bullet_item(doc, "Add A Record: Host = @  |  Value = 75.2.60.5  |  TTL = Automatic")
    add_bullet_item(doc, "Add CNAME Record: Host = www  |  Value = [their-netlify-site].netlify.app  |  TTL = Automatic")
    add_bullet_item(doc, "In Netlify → add the custom domain, verify, and provision the SSL certificate")

    add_h3(doc, "If the Client Cancels — Domain Transfer")
    add_bullet_item(doc, "Ask the client to create a Namecheap account (or any registrar they prefer)")
    add_bullet_item(doc, "In your Namecheap account → Domain List → Manage → Sharing & Transfer → Transfer to another user")
    add_bullet_item(doc, "Enter the client's Namecheap username — they accept the transfer and the domain is theirs")
    add_bullet_item(doc, "The transfer is free and takes just a few minutes — no cost to the client")
    add_bullet_item(doc, "Confirm the transfer is complete before closing out the client account")

    add_h3(doc, "DNS Reference — Exact Records for Both Options")
    add_table(doc, [
        ("__header__", ("Record", "Details")),
        ("A Record",     "Host: @  |  Value: 75.2.60.5  |  TTL: Automatic"),
        ("CNAME Record", "Host: www  |  Value: [their-site].netlify.app  |  TTL: Automatic"),
    ])
    add_body(doc,
        "Note: replace [their-site].netlify.app with the actual Netlify subdomain for this client "
        "(found in Netlify → Site Settings → General → Site name)."
    )

    # ══════════════════════════════════════════════════════════════════════════
    # PART 8 — NETLIFY (Website + Lead Management only)
    # ══════════════════════════════════════════════════════════════════════════
    add_h1(doc, "Part 8 — Netlify Deployment")

    add_body(doc,
        "Applies to: Website + Lead Management only. "
        "Each client site is deployed as a separate Netlify site, pointing to their specific folder "
        "in your GitHub repo. This keeps sites independent — updating one client does not affect another."
    )

    add_h2(doc, "Step 1 — Create a New Netlify Site")
    add_bullet_item(doc, "Log into Netlify → Add new site → Import an existing project")
    add_bullet_item(doc, "Connect to GitHub → Select your LeadManagement repo")
    add_bullet_item(doc, "Set the Publish Directory to the client's folder path (e.g., clients/rivera-plumbing)")
    add_bullet_item(doc, "Leave Build Command blank — you are deploying plain HTML, no build step needed")
    add_bullet_item(doc, "Click Deploy Site")

    add_h2(doc, "Step 2 — Confirm the Site Is Live")
    add_bullet_item(doc, "Netlify will assign a random subdomain (e.g., luminous-fox-123.netlify.app)")
    add_bullet_item(doc, "Open that URL in a browser and verify the client's site loads correctly")
    add_bullet_item(doc, "Test the contact form — submit a test entry and check that the Google Sheet receives it")
    add_bullet_item(doc, "Check on mobile — resize your browser or use Chrome DevTools mobile view")

    add_h2(doc, "Step 3 — Connect the Custom Domain")
    add_bullet_item(doc, "In Netlify → Site Settings → Domain Management → Add custom domain")
    add_bullet_item(doc, "Enter the client's domain (e.g., riveraplumbing.com)")
    add_bullet_item(doc, "Netlify will prompt you to verify DNS — confirm the A record and CNAME from Part 7 are set")
    add_bullet_item(doc, "Enable HTTPS — Netlify provisions a free SSL certificate automatically (Let's Encrypt)")
    add_bullet_item(doc, "Site is fully live once HTTPS shows a green padlock in the browser")

    # ══════════════════════════════════════════════════════════════════════════
    # PART 9 — SNIPPET DEPLOYMENT (Lead Intake + Lead Management only)
    # ══════════════════════════════════════════════════════════════════════════
    add_h1(doc, "Part 9 — Snippet Deployment")

    add_body(doc,
        "Applies to: Lead Intake + Lead Management only. "
        "The client already has a website. You add a self-contained lead intake form snippet "
        "to their existing site. No Netlify, no GitHub repo changes needed for their site — "
        "just the embed code and the Apps Script webhook."
    )

    add_h2(doc, "Step 1 — Set Up the Apps Script First")
    add_body(doc,
        "Complete Part 10 (lead system setup) before providing the snippet code — you need the "
        "webhook URL to put into the snippet."
    )

    add_h2(doc, "Step 2 — Prepare the Snippet")
    add_bullet_item(doc, "Open clients/templates/snippet-lead-intake-only.html from your GitHub repo")
    add_bullet_item(doc, "Replace [CLIENT_WEBHOOK_URL] with the Apps Script deployment URL from Part 10")
    add_bullet_item(doc, "Update any placeholder text: phone number, business-specific confirmation message")
    add_bullet_item(doc, "Test the snippet locally by opening it in a browser and submitting a test entry")

    add_h2(doc, "Step 3 — Deliver the Snippet to the Client")
    add_body(doc,
        "Most Lead Intake clients have a web developer, a website platform (Squarespace, Wix, WordPress), "
        "or manage their own site. Provide the snippet in a way that works for their situation."
    )
    add_bullet_item(doc, "If they have a developer: send the three-part snippet (styles, HTML, script) with clear labels for where each piece goes")
    add_bullet_item(doc, "If they are on WordPress: walk them through using a Custom HTML block in their page editor")
    add_bullet_item(doc, "If they are on Squarespace or Wix: use a Code Block or Embed element on the page")
    add_bullet_item(doc, "If they want you to add it directly: get temporary access to their site and add it yourself")

    add_h2(doc, "Step 4 — Confirm the Form Is Live")
    add_bullet_item(doc, "Once the snippet is on their site, submit a test entry through the live form")
    add_bullet_item(doc, "Confirm the row appears in the Google Sheet within seconds")
    add_bullet_item(doc, "Confirm the notification email fires")
    add_bullet_item(doc, "Do not mark the client as onboarded until the live test passes")

    # ══════════════════════════════════════════════════════════════════════════
    # PART 10 — LEAD SYSTEM (both services)
    # ══════════════════════════════════════════════════════════════════════════
    add_h1(doc, "Part 10 — Lead System Setup (Google Sheets + Apps Script)")

    add_body(doc,
        "Applies to: both services. Each client gets their own Google Sheet for lead tracking "
        "and their own Apps Script webhook. This keeps client data completely separate."
    )

    add_h2(doc, "Step 1 — Create the Google Sheet")
    add_bullet_item(doc, "In Google Drive, create a new folder for the client")
    add_bullet_item(doc, "Inside the folder, create a new Google Sheet named: [Client Name] — Leads")
    add_bullet_item(doc, "Add column headers in Row 1: Timestamp, Name, Phone, Email, Message, Source")

    add_h2(doc, "Step 2 — Set Up the Apps Script Webhook")
    add_bullet_item(doc, "Open the Google Sheet → Extensions → Apps Script")
    add_bullet_item(doc, "Paste in your standard doPost() webhook script")
    add_bullet_item(doc, "Click Deploy → New Deployment → Web App")
    add_bullet_item(doc, "Set Execute as: Me — Who has access: Anyone")
    add_bullet_item(doc, "Copy the deployment URL — this is the WEBHOOK_URL for the client's form")
    add_bullet_item(doc, "Paste the webhook URL into the client's HTML template or snippet")
    add_bullet_item(doc, "For Website + Lead Management: push the updated HTML to GitHub (Netlify will auto-redeploy)")
    add_bullet_item(doc, "For Lead Intake + Lead Management: provide the updated snippet to the client")

    add_h2(doc, "Step 3 — Test the Full Lead Flow")
    add_bullet_item(doc, "Open the live client site or snippet and submit the contact form with test data")
    add_bullet_item(doc, "Confirm the row appears in the Google Sheet within seconds")
    add_bullet_item(doc, "Confirm the client receives a notification email")
    add_bullet_item(doc, "If the form does not submit, check the webhook URL and redeploy")

    # ══════════════════════════════════════════════════════════════════════════
    # PART 11 — GBP (both services)
    # ══════════════════════════════════════════════════════════════════════════
    add_h1(doc, "Part 11 — Google Business Profile Access")

    add_body(doc,
        "Applies to: both services — quarterly GBP audit is included in both packages. "
        "Refer to the full guide in your docs folder: Google Business Profile — Manager Access Guide.docx. "
        "Below is the quick checklist for onboarding."
    )
    add_bullet_item(doc, "Ask the client if they have a Google Business Profile set up")
    add_bullet_item(doc, "If yes: request Manager access (not Owner) — follow the GBP Manager Access Guide")
    add_bullet_item(doc, "If no: walk them through creating one, then request Manager access")
    add_bullet_item(doc, "Confirm your Manager access appears in their GBP dashboard before marking this complete")
    add_bullet_item(doc, "Add a recurring calendar reminder to do a quarterly audit for this client")

    # ══════════════════════════════════════════════════════════════════════════
    # PART 12 — HANDOFF AND ONGOING (both services)
    # ══════════════════════════════════════════════════════════════════════════
    add_h1(doc, "Part 12 — Client Handoff and Ongoing Support")

    add_body(doc,
        "Applies to: both services. Once the form is live, tested, and the lead system is confirmed "
        "working, send the client a handoff email."
    )

    add_h2(doc, "Step 1 — Send the Handoff Email")
    add_code(doc,
        "Subject: Your lead system is live — here's everything you need\n\n"
        "Hi [NAME],\n\n"
        "You're all set. Here's a quick summary of what's in place:\n\n"
        "  [Website: https://theirdomain.com]  ← include if Website + Lead Management\n"
        "  Lead Sheet: [link to their Google Sheet]\n"
        "  Lead form: tested and working — every submission goes to the sheet\n"
        "  Notifications: you'll get an email + text within minutes of each new lead\n"
        "  Monthly report: I'll send you a lead summary at the end of each month\n"
        "  GBP audit: I'll review your Google Business Profile quarterly\n\n"
        "For any updates or questions, just reply to this email.\n\n"
        "Your monthly billing of $[69 or 59] will continue on the 1st of each month\n"
        "via Stripe — you'll receive a receipt automatically.\n\n"
        "Thanks for trusting me with this.\n\n"
        "[YOUR NAME]\n"
        "[YOUR GOOGLE VOICE NUMBER]"
    )

    add_h2(doc, "Step 2 — What the Client Can Expect Each Month")
    add_bullet_item(doc, "Instant email + text notification every time someone submits the lead form")
    add_bullet_item(doc, "Every lead logged automatically in their Google Sheet — organized and permanent")
    add_bullet_item(doc, "Monthly lead summary report sent by you at the end of each month")
    add_bullet_item(doc, "Quarterly Google Business Profile audit — you review for accuracy and flag issues")
    add_bullet_item(doc, "Speed-to-lead facts sheet provided at onboarding — explains the value of responding fast")
    add_bullet_item(doc, "Form updates always included at no charge — just ask")
    add_bullet_item(doc, "Website + Lead Management clients: ongoing site tweaks and content edits included")

    add_h2(doc, "Step 3 — Your Monthly Maintenance Checklist")
    add_bullet_item(doc, "Submit a test entry on the client's form — confirm it lands in the Google Sheet")
    add_bullet_item(doc, "Review Netlify for any deploy errors or alerts (Website + Lead Management clients only)")
    add_bullet_item(doc, "Check Stripe for any failed payments — follow up with client within 3 business days")
    add_bullet_item(doc, "Apply any content updates or form changes the client has requested")
    add_bullet_item(doc, "Send the monthly lead summary report — pull from the Google Sheet and email it")
    add_bullet_item(doc, "Once per quarter: run the GBP audit and send a brief summary to the client")

    out_path = "/Users/creighbaby/LeadManagement/docs/Client Onboarding — End-to-End Operations Guide.docx"
    doc.save(out_path)
    print(f"Saved: {out_path}")


if __name__ == "__main__":
    build()
