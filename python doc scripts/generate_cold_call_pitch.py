"""
Cold Call Sales Pitch — How-To Guide
Generates a branded .docx following the LeadManagement Source-of-Truth style.
"""

from docx import Document
from docx.shared import Pt, Inches, RGBColor, Twips
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
import copy

# ── COLOR PALETTE ──────────────────────────────────────────────────────────────
DARK_GREEN   = RGBColor(0x1a, 0x3a, 0x2a)
NEAR_BLACK   = RGBColor(0x1c, 0x1c, 0x1a)
GOLD         = "c8a96e"
CREAM_LINE   = "e2e0d8"
TABLE_HDR_BG = "2c1f0e"
WHITE        = RGBColor(0xFF, 0xFF, 0xFF)
CODE_BG      = "F2F2F2"
BORDER_COLOR = "D0CEC6"


# ── HELPERS ────────────────────────────────────────────────────────────────────

def set_para_spacing(para, before=0, after=120):
    pf = para.paragraph_format
    pf.space_before = Twips(before)
    pf.space_after  = Twips(after)


def add_bottom_border(para, color_hex, sz=12, space=4):
    pPr = para._p.get_or_add_pPr()
    pBdr = OxmlElement("w:pBdr")
    bottom = OxmlElement("w:bottom")
    bottom.set(qn("w:val"),   "single")
    bottom.set(qn("w:sz"),    str(sz))
    bottom.set(qn("w:space"), str(space))
    bottom.set(qn("w:color"), color_hex)
    pBdr.append(bottom)
    pPr.append(pBdr)


def set_cell_shading(cell, fill_hex):
    tc   = cell._tc
    tcPr = tc.get_or_add_tcPr()
    shd  = OxmlElement("w:shd")
    shd.set(qn("w:val"),   "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"),  fill_hex)
    tcPr.append(shd)


def set_cell_margins(cell, top=80, bottom=80, left=115, right=115):
    tc   = cell._tc
    tcPr = tc.get_or_add_tcPr()
    tcMar = OxmlElement("w:tcMar")
    for side, val in [("top", top), ("bottom", bottom), ("left", left), ("right", right)]:
        el = OxmlElement(f"w:{side}")
        el.set(qn("w:w"),    str(val))
        el.set(qn("w:type"), "dxa")
        tcMar.append(el)
    tcPr.append(tcMar)


def set_table_borders(table, color_hex=BORDER_COLOR):
    tbl    = table._tbl
    tblPr  = tbl.find(qn("w:tblPr"))
    if tblPr is None:
        tblPr = OxmlElement("w:tblPr")
        tbl.insert(0, tblPr)
    tblBdr = OxmlElement("w:tblBorders")
    for side in ("top", "left", "bottom", "right", "insideH", "insideV"):
        el = OxmlElement(f"w:{side}")
        el.set(qn("w:val"),   "single")
        el.set(qn("w:sz"),    "4")
        el.set(qn("w:space"), "0")
        el.set(qn("w:color"), color_hex)
        tblBdr.append(el)
    tblPr.append(tblBdr)


def add_h1(doc, text):
    para = doc.add_paragraph()
    run  = para.add_run(text)
    run.bold      = True
    run.font.name = "Arial"
    run.font.size = Pt(18)
    run.font.color.rgb = DARK_GREEN
    set_para_spacing(para, before=400, after=200)
    add_bottom_border(para, GOLD, sz=12, space=4)
    return para


def add_h2(doc, text):
    para = doc.add_paragraph()
    run  = para.add_run(text)
    run.bold      = True
    run.font.name = "Arial"
    run.font.size = Pt(14)
    run.font.color.rgb = DARK_GREEN
    set_para_spacing(para, before=300, after=160)
    add_bottom_border(para, CREAM_LINE, sz=4, space=4)
    return para


def add_h3(doc, text):
    para = doc.add_paragraph()
    run  = para.add_run(text)
    run.bold      = True
    run.font.name = "Arial"
    run.font.size = Pt(12)
    run.font.color.rgb = DARK_GREEN
    set_para_spacing(para, before=220, after=120)
    return para


def add_body(doc, text):
    para = doc.add_paragraph()
    run  = para.add_run(text)
    run.font.name = "Arial"
    run.font.size = Pt(11)
    run.font.color.rgb = NEAR_BLACK
    set_para_spacing(para, before=0, after=120)
    return para


def add_bullet_item(doc, text, bold_prefix=None):
    para  = doc.add_paragraph(style="List Bullet")
    if bold_prefix:
        r1 = para.add_run(bold_prefix)
        r1.bold = True
        r1.font.name = "Arial"
        r1.font.size = Pt(11)
        r1.font.color.rgb = NEAR_BLACK
    run = para.add_run(text)
    run.font.name = "Arial"
    run.font.size = Pt(11)
    run.font.color.rgb = NEAR_BLACK
    set_para_spacing(para, before=0, after=80)
    return para


def add_code(doc, text):
    para = doc.add_paragraph()
    run  = para.add_run(text)
    run.font.name = "Courier New"
    run.font.size = Pt(9)
    run.font.color.rgb = NEAR_BLACK
    set_para_spacing(para, before=60, after=60)
    # Gray background via shading on the paragraph's run — approximate via table trick not needed;
    # apply shading to the paragraph's pPr instead
    pPr  = para._p.get_or_add_pPr()
    shd  = OxmlElement("w:shd")
    shd.set(qn("w:val"),   "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"),  CODE_BG)
    pPr.append(shd)
    return para


def add_table(doc, rows_data, col_widths=(1.75, 4.5)):
    """
    rows_data: list of (label, detail) tuples.
    First tuple is treated as header if label == '__header__'.
    """
    table = doc.add_table(rows=0, cols=2)
    set_table_borders(table)

    for i, (label, detail) in enumerate(rows_data):
        row   = table.add_row()
        left  = row.cells[0]
        right = row.cells[1]

        left.width  = Inches(col_widths[0])
        right.width = Inches(col_widths[1])

        set_cell_margins(left)
        set_cell_margins(right)

        if label == "__header__":
            # Dark brown header row
            set_cell_shading(left,  TABLE_HDR_BG)
            set_cell_shading(right, TABLE_HDR_BG)

            lp = left.paragraphs[0]
            lr = lp.add_run(detail[0])
            lr.bold = True
            lr.font.name = "Arial"
            lr.font.size = Pt(11)
            lr.font.color.rgb = WHITE

            rp = right.paragraphs[0]
            rr = rp.add_run(detail[1])
            rr.bold = True
            rr.font.name = "Arial"
            rr.font.size = Pt(11)
            rr.font.color.rgb = WHITE
        else:
            lp = left.paragraphs[0]
            lr = lp.add_run(label)
            lr.bold = True
            lr.font.name = "Arial"
            lr.font.size = Pt(11)
            lr.font.color.rgb = DARK_GREEN

            rp = right.paragraphs[0]
            rr = rp.add_run(detail)
            rr.font.name = "Arial"
            rr.font.size = Pt(11)
            rr.font.color.rgb = NEAR_BLACK

    # Spacer after table
    spacer = doc.add_paragraph()
    set_para_spacing(spacer, before=0, after=80)
    return table


def add_title_block(doc, line1, line2, subtitle):
    p1 = doc.add_paragraph()
    p1.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r1 = p1.add_run(line1)
    r1.bold = True
    r1.font.name = "Arial"
    r1.font.size = Pt(22)
    r1.font.color.rgb = DARK_GREEN
    set_para_spacing(p1, before=0, after=60)

    p2 = doc.add_paragraph()
    p2.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r2 = p2.add_run(line2)
    r2.bold = True
    r2.font.name = "Arial"
    r2.font.size = Pt(22)
    r2.font.color.rgb = DARK_GREEN
    set_para_spacing(p2, before=0, after=100)

    p3 = doc.add_paragraph()
    p3.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r3 = p3.add_run(subtitle)
    r3.italic = True
    r3.font.name = "Arial"
    r3.font.size = Pt(11)
    r3.font.color.rgb = NEAR_BLACK
    set_para_spacing(p3, before=0, after=300)


# ── BUILD DOCUMENT ─────────────────────────────────────────────────────────────

def build():
    doc = Document()

    # Page margins
    for section in doc.sections:
        section.top_margin    = Inches(1.0)
        section.bottom_margin = Inches(1.0)
        section.left_margin   = Inches(1.1)
        section.right_margin  = Inches(1.1)

    # ── TITLE BLOCK
    add_title_block(
        doc,
        "COLD CALL SALES PITCH",
        "HOW-TO GUIDE",
        "A ready-to-use script for selling website + lead management services to small service businesses"
    )

    # ── WHAT'S INSIDE TABLE
    add_h2(doc, "What's inside this document")
    add_table(doc, [
        ("__header__", ("Section", "Contents")),
        ("Part 1 — What You're Selling",   "Product overview and elevator pitch in plain language"),
        ("Part 2 — The 30-Second Opener",  "The exact words to say in the first 30 seconds of a call"),
        ("Part 3 — Why They Should Buy",   "Client pain points, benefits, and what sets you apart"),
        ("Part 4 — Return on Investment",  "ROI breakdown — how to show the math to a skeptical client"),
        ("Part 5 — The Full Call Script",  "Word-for-word pitch from opener to close"),
        ("Part 6 — Handling Objections",   "The most common pushbacks and how to respond"),
        ("Part 7 — Closing the Call",      "How to end every call with a clear next step"),
    ])

    # ── HOW TO USE
    add_h2(doc, "How to use this document")
    add_body(doc,
        "This guide is built for cold calls to busy tradespeople — plumbers, electricians, HVAC techs, "
        "landscapers, and other service businesses. These clients are often on a job site, driving between "
        "calls, or covered in grease. You have 20 seconds before they hang up. Every section in this guide "
        "is written with that reality in mind. Read Part 5 out loud before your first call. Personalize the "
        "opener with the prospect's business name when you can. The rest is a framework — not a script you "
        "have to follow word-for-word."
    )

    # ══════════════════════════════════════════════════════════════════════════
    # PART 1 — WHAT YOU'RE SELLING
    # ══════════════════════════════════════════════════════════════════════════
    add_h1(doc, "Part 1 — What You're Selling")

    add_h2(doc, "Step 1 — Know Your Product Cold")
    add_body(doc,
        "You offer two services. The first is for businesses that need everything — a professional website "
        "and a full lead management system built in. The second is for businesses that already have a website "
        "but are losing leads because nothing is capturing or tracking them. Both services include the same "
        "lead management backbone: instant alerts, an organized lead log, monthly reports, and ongoing support."
    )

    add_h3(doc, "The Two Services")
    add_table(doc, [
        ("__header__", ("Service", "What It Does — In Plain Language")),
        ("Website + Lead Management\n$750 setup + $79/mo",
         "Builds a professional mobile-ready website and wires it up so that every time someone submits the "
         "contact form OR calls and misses you, the owner gets a text on their phone within minutes — "
         "name, number, and what they need. Every lead is automatically saved to a Google Sheet. "
         "Daily digest email every morning. Monthly report on the 1st. "
         "Quarterly Google Business Profile review. Ongoing site edits always included."),
        ("Lead Intake + Lead Management\n$400 setup + $59/mo",
         "No new website needed. Adds a lead capture form to their existing site and wires up the same "
         "alert system — every form submission and every missed call sends a text to the owner's phone "
         "within minutes. Everything logged automatically in Google Sheets. "
         "Same daily digest, monthly report, and GBP audit as the full package."),
    ])

    add_h3(doc, "Which Service to Lead With")
    add_body(doc,
        "Before the call, do a quick Google search on the prospect. If they have no website or a weak one, "
        "lead with Website + Lead Management. If they already have a decent website, lead with "
        "Lead Intake + Lead Management — it's a lower barrier to entry and solves a more immediate problem."
    )

    add_h3(doc, "The Guarantee")
    add_body(doc,
        "60-day money-back guarantee on monthly fees. If the client is not happy in the first 60 days, "
        "every monthly payment is refunded — no questions asked. Setup fees are non-refundable once work "
        "begins. Use this on every call — it removes almost all of the risk objection."
    )

    add_h3(doc, "The Timeline")
    add_body(doc,
        "Most clients are live in 3 to 5 business days. This is a massive differentiator. Agencies take "
        "6 to 12 weeks. You take less than a week. Lead with this."
    )

    # ══════════════════════════════════════════════════════════════════════════
    # PART 2 — THE 30-SECOND OPENER
    # ══════════════════════════════════════════════════════════════════════════
    add_h1(doc, "Part 2 — The 30-Second Opener")

    add_h2(doc, "Step 2 — Hook Them Before They Hang Up")
    add_body(doc,
        "A plumber on a job site has no patience for a sales pitch. Your opener must do three things in "
        "under 30 seconds: (1) say who you are, (2) say exactly what you do in one sentence, and "
        "(3) tie it directly to their problem. Do not ask 'is this a good time?' — it gives them an "
        "easy out. Instead, acknowledge you know they're busy right up front."
    )

    add_h3(doc, "The Opener — Read This Out Loud")
    add_code(doc,
        '"Hey [Name], my name is [Your Name] — I work with plumbers and contractors in [City] to get them '
        'a professional website and a system that sends every new customer inquiry straight to their phone '
        'within minutes. I know you\'re probably in the middle of a job — this will take 60 seconds. '
        'Can I have just a minute?"'
    )

    add_h3(doc, "Why This Works")
    add_bullet_item(doc, "You name their industry — plumber, electrician, HVAC — so they feel seen, not cold-called.", bold_prefix="Personalized: ")
    add_bullet_item(doc, "\"Sends leads straight to your phone\" — that's the hook. They've missed calls before.", bold_prefix="Specific benefit: ")
    add_bullet_item(doc, "You ask for 60 seconds, not 20 minutes. Low commitment.", bold_prefix="Low ask: ")
    add_bullet_item(doc, "\"I know you're in the middle of a job\" — shows respect for their time.", bold_prefix="Empathy: ")

    add_h3(doc, "If They Say 'I'm Busy Right Now'")
    add_code(doc,
        '"Totally understand — when\'s a better time? I can call back this afternoon or tomorrow morning '
        'before your day gets going."'
    )

    # ══════════════════════════════════════════════════════════════════════════
    # PART 3 — WHY THEY SHOULD BUY
    # ══════════════════════════════════════════════════════════════════════════
    add_h1(doc, "Part 3 — Why They Should Buy")

    add_h2(doc, "Step 3 — Hit the Real Pain Points")
    add_body(doc,
        "Most small service businesses have the same core problems. Knowing these before you call "
        "lets you speak directly to what's already keeping them up at night. You're not selling a website "
        "— you're selling fewer missed jobs and more money in their pocket."
    )

    add_h3(doc, "The Four Pain Points")
    add_table(doc, [
        ("__header__", ("Pain Point", "What It Costs Them")),
        ("No website or a bad one",
         "Customers Google them, find nothing, and call a competitor. They lose jobs they never even know they lost."),
        ("Leads come in while they're on a job",
         "A customer fills out a form or calls while the owner is under a sink. They miss it. "
         "By the time they call back — hours later — the customer already booked someone else."),
        ("Missed calls with no follow-up system",
         "They get a voicemail. They listen to it at 6pm, forget by morning, and never call back. "
         "That was a $600 job. Gone."),
        ("No system to track who called",
         "They're juggling names and numbers in their head or in a notes app. Leads fall through the cracks every week."),
        ("Invisible on Google",
         "Their competitor with a Google Business Profile shows up first. They don't. The phone goes to the other guy."),
    ])

    add_h3(doc, "Your Value Proposition — One Sentence")
    add_code(doc,
        '"I give you a professional website that works while you\'re on the job, and every lead that '
        'comes in goes straight to your inbox within minutes — organized and ready to call back."'
    )

    add_h3(doc, "What Sets You Apart")
    add_bullet_item(doc, "Live in under a week — not 6 weeks like an agency.", bold_prefix="Speed: ")
    add_bullet_item(doc, "Covers both form submissions AND missed calls — both log as leads and both send a text alert to the owner.", bold_prefix="Voicemail-to-lead: ")
    add_bullet_item(doc, "Their domain, their Google Sheet, their data. They own everything and can walk away any time.", bold_prefix="Ownership: ")
    add_bullet_item(doc, "They deal with one person — not a support ticket or a rotating team.", bold_prefix="One person: ")
    add_bullet_item(doc, "60 days to decide if it's working. Monthly fees come back if they're not satisfied.", bold_prefix="Guarantee: ")
    add_bullet_item(doc, "$79/month for the full website + lead management package. $59/month if they already have a site.", bold_prefix="Price: ")

    # ══════════════════════════════════════════════════════════════════════════
    # PART 4 — RETURN ON INVESTMENT
    # ══════════════════════════════════════════════════════════════════════════
    add_h1(doc, "Part 4 — Return on Investment")

    add_h2(doc, "Step 4 — Show Them the Math")
    add_body(doc,
        "Tradespeople think in jobs, not monthly fees. Translate your price into jobs. One saved job "
        "pays for 12 months of service. When you frame it that way, the math is impossible to argue with."
    )

    add_h3(doc, "The ROI Calculation")
    add_table(doc, [
        ("__header__", ("Scenario", "Numbers")),
        ("Average plumbing job value",       "$400 – $1,500 (use $600 as a conservative estimate)"),
        ("Leads missed per month (typical)", "3 – 5 leads lost to slow follow-up or no website presence"),
        ("Revenue lost per month",           "3 missed leads x $600 = $1,800 lost every month"),
        ("Website + Lead Management cost",   "$79/month after a one-time $750 setup fee"),
        ("Lead Intake + Lead Management",    "$59/month after a one-time $400 setup fee"),
        ("Break-even point",                 "One saved job covers a full year of monthly fees"),
        ("ROI if only 1 lead saved/month",  "$600 recovered – $79 cost = $521 net gain, every month"),
    ])

    add_h3(doc, "Say This on the Call")
    add_code(doc,
        '"Here\'s the honest math — if you\'re a plumber and the average job is $600, and you recover '
        'even one lead a month that you would have missed otherwise, that\'s $600 coming in against $79 '
        'going out. The service pays for itself the first time your phone buzzes with a lead while '
        'you\'re under a sink."'
    )

    add_h3(doc, "The One-Year Frame")
    add_code(doc,
        '"One saved job pays for an entire year of service. That\'s the deal. And if it doesn\'t work '
        'in the first 60 days, you get every monthly payment back. I take all the risk."'
    )

    # ══════════════════════════════════════════════════════════════════════════
    # PART 5 — FULL CALL SCRIPT
    # ══════════════════════════════════════════════════════════════════════════
    add_h1(doc, "Part 5 — The Full Cold Call Script")

    add_h2(doc, "Step 5 — The Complete Word-for-Word Pitch")
    add_body(doc,
        "This script is designed to run in under 3 minutes if the client is engaged. Read it out loud "
        "before your first call. Adapt the business type and city to match the prospect. Anything in "
        "[brackets] is a variable you fill in live."
    )

    add_h3(doc, "The Opener (0:00 – 0:30)")
    add_code(doc,
        '"Hey [Name], my name is [Your Name]. I work with [plumbers / electricians / landscapers] in '
        '[City] to get them a professional website and a system that sends every new customer lead '
        'straight to their phone within minutes. I know you\'re probably in the middle of a job — '
        'can I have just 60 seconds?"'
    )

    add_h3(doc, "The Hook (0:30 – 1:00)")
    add_code(doc,
        '"Here\'s the thing — most [plumbers] I talk to are losing 3 or 4 jobs a month and they don\'t '
        'even know it. A customer Googles them, finds nothing — or finds a site that looks like it\'s from '
        '2009 — and calls the next guy. Or they fill out a form while you\'re on a job and by the time '
        'you call back, they\'ve already booked someone else. Or they leave a voicemail, and you hear it '
        'at 7pm, forget about it by morning, and that was a $600 job gone. That\'s real money walking out '
        'the door every single month."'
    )

    add_h3(doc, "The Solution (1:00 – 1:45)")
    add_code(doc,
        '"What I do is build you a professional website that looks great on a phone — because that\'s how '
        'people are searching — and wire up a system so that every lead that comes in goes straight to '
        'your phone as a text within minutes. Form submission, missed call, voicemail — all of it. '
        'Your phone buzzes and it says: \'New lead — Mike R., leaking pipe under kitchen sink, Naperville. '
        'Call back today.\' You\'re not checking a dashboard. You\'re not logging into anything. '
        'You just get a text. You\'re live in under a week. It\'s $79 a month after a one-time setup."'
    )

    add_h3(doc, "The ROI Moment (1:45 – 2:15)")
    add_code(doc,
        '"Here\'s the math — if your average job is $600 and you recover even one lead a month you '
        'would have missed, you\'re up $531 that month alone. One saved job covers a full year of my '
        'fee. And if it doesn\'t work in the first 60 days — if you\'re not getting leads '
        'or you\'re not happy for any reason — I refund every monthly payment, no questions. I take '
        'all the risk on this."'
    )

    add_h3(doc, "The Close (2:15 – 3:00)")
    add_code(doc,
        '"I\'m not asking you to commit to anything today. What I\'d like to do is get on a quick call '
        'with you for 15 minutes when you\'re off the job — I\'ll show you what these sites actually '
        'look like, walk you through how the lead system works, and give you a no-pressure quote. '
        'If it makes sense, great. If not, no hard feelings. When\'s a good time — would later today '
        'work, or is tomorrow morning better?"'
    )

    # ══════════════════════════════════════════════════════════════════════════
    # PART 6 — HANDLING OBJECTIONS
    # ══════════════════════════════════════════════════════════════════════════
    add_h1(doc, "Part 6 — Handling Objections")

    add_h2(doc, "Step 6 — Turn Pushback into Progress")
    add_body(doc,
        "Every objection is a question in disguise. Below are the most common pushbacks you'll hear "
        "on a cold call to a tradesperson, and exactly how to respond. The goal is not to argue — "
        "it's to lower the risk and keep the conversation moving."
    )

    add_h3(doc, "Objection: 'I already have a website.'")
    add_code(doc,
        '"That\'s great — a lot of the folks I work with had a website too. The question is whether '
        'it\'s actually bringing in leads. Is it connected to a system that texts you when someone '
        'fills out the form? If not, you might be getting traffic but losing the customer before '
        'you even know they called. That\'s the piece I fix."'
    )

    add_h3(doc, "Objection: 'I get all my work from referrals.'")
    add_code(doc,
        '"Referrals are the best kind of work — no argument there. But here\'s the thing: even your '
        'referral customers Google you before they call. If they find nothing — or a bad site — '
        'you\'ve already lost some of them. A good web presence makes your referrals convert better. '
        'It\'s like a business card that never runs out."'
    )

    add_h3(doc, "Objection: 'I can\'t afford it right now.'")
    add_code(doc,
        '"I hear you — and that\'s exactly why I want to show you the math. If one saved lead a month '
        'is worth $600 to you, and my fee is $79 a month, you\'re net positive by over $500 every single '
        'month. And if it doesn\'t work in 60 days, you get every monthly payment back. You\'re not '
        'risking the money — you\'re risking 60 days of your time."'
    )

    add_h3(doc, "Objection: 'I don\'t have time for this.'")
    add_code(doc,
        '"That\'s the whole reason I built this the way I did. You don\'t manage anything. I build '
        'it, I maintain it, I make updates when you need them. All you do is get a notification when '
        'a lead comes in and decide if you want to call them back. That\'s it. The setup call takes '
        'about 15 minutes."'
    )

    add_h3(doc, "Objection: 'I need to think about it.'")
    add_code(doc,
        '"Totally fair — I\'d never want you to rush a decision. Can I ask — is there a specific '
        'part you\'re unsure about? Sometimes there\'s one question I can answer that makes the '
        'rest clear. And if you just want to see what the sites actually look like before deciding, '
        'I can send you three examples right now. Would that help?"'
    )

    add_h3(doc, "Objection: 'I had a bad experience with a web company before.'")
    add_code(doc,
        '"I\'m not surprised — a lot of people have. Big agencies take months, charge a fortune, '
        'and then you can\'t even make a simple change without paying extra. I\'m one person — you '
        'deal with me directly. Your domain is in your name. Your leads go to your Google Sheet. '
        'If you ever leave, you take everything with you. I don\'t lock you in."'
    )

    # ══════════════════════════════════════════════════════════════════════════
    # PART 7 — CLOSING THE CALL
    # ══════════════════════════════════════════════════════════════════════════
    add_h1(doc, "Part 7 — Closing the Call")

    add_h2(doc, "Step 7 — Always End with a Specific Next Step")
    add_body(doc,
        "The biggest mistake on a cold call is ending without a clear next action. 'I'll think about it' "
        "and 'send me some info' are not next steps — they're polite exits. Every call should end with "
        "a booked time on the calendar, a specific follow-up call agreed upon, or a clear no so you "
        "can move on. Never leave a call open-ended."
    )

    add_h3(doc, "The Soft Close — For Engaged Prospects")
    add_code(doc,
        '"It sounds like this could be a good fit. The next step is just a quick 15-minute call where '
        'I show you exactly what the system looks like and give you a quote. Does [Day] at [Time] work '
        'for you, or is [alternate time] better?"'
    )

    add_h3(doc, "The Follow-Up Close — For 'I Need to Think' Responses")
    add_code(doc,
        '"No problem at all. I\'ll send you a couple of examples so you can see what other service '
        'businesses\' sites look like. Can I follow up with you on [specific day] once you\'ve had '
        'a chance to look them over? What\'s the best number to reach you?"'
    )

    add_h3(doc, "The Clean Exit — If They Are Clearly Not Interested")
    add_code(doc,
        '"Totally understood — I appreciate you taking a minute. If anything changes down the road '
        'and you want to revisit it, feel free to reach out. Have a good one."'
    )

    add_h3(doc, "Post-Call Checklist")
    add_bullet_item(doc, "Log the call: name, business, phone, outcome, and follow-up date in your tracker.")
    add_bullet_item(doc, "If they agreed to a demo call: send a calendar invite within 10 minutes.")
    add_bullet_item(doc, "If they said 'send me info': email 2–3 site examples the same day — not a week later.")
    add_bullet_item(doc, "If they said 'call me back': set a reminder for the specific day and time they gave you.")
    add_bullet_item(doc, "If they said no: mark them as cold and move on. Don't follow up more than once.")

    add_h3(doc, "Quick Reference — Pricing Cheat Sheet")
    add_table(doc, [
        ("__header__", ("Service", "Setup / Monthly")),
        ("Website + Lead Management",        "$750 setup — $79/month"),
        ("Lead Intake + Lead Management",    "$400 setup — $59/month"),
        ("Payment terms",                    "50% deposit on signing, balance before go-live"),
        ("Guarantee",                        "60-day money-back on monthly fees — no questions asked"),
        ("Timeline",                         "Live in 3–5 business days from signing"),
    ])

    out_path = "/Users/creighbaby/LeadManagement/docs/Cold Call Sales Pitch — How-To Guide.docx"
    doc.save(out_path)
    print(f"Saved: {out_path}")


if __name__ == "__main__":
    build()
