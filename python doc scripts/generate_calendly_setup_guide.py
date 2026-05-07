"""
Client Calendly Setup — How-To Guide
Generates a branded .docx following the LeadManagement Source-of-Truth style.
"""

from docx import Document
from docx.shared import Pt, Inches, RGBColor, Twips
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

DARK_GREEN   = RGBColor(0x1a, 0x3a, 0x2a)
NEAR_BLACK   = RGBColor(0x1c, 0x1c, 0x1a)
GOLD         = "c8a96e"
CREAM_LINE   = "e2e0d8"
TABLE_HDR_BG = "2c1f0e"
WHITE        = RGBColor(0xFF, 0xFF, 0xFF)
CODE_BG      = "F2F2F2"
BORDER_COLOR = "D0CEC6"


def set_para_spacing(para, before=0, after=120):
    pf = para.paragraph_format
    pf.space_before = Twips(before)
    pf.space_after  = Twips(after)

def add_bottom_border(para, color_hex, sz=12, space=4):
    pPr = para._p.get_or_add_pPr()
    pBdr = OxmlElement("w:pBdr")
    bottom = OxmlElement("w:bottom")
    bottom.set(qn("w:val"), "single"); bottom.set(qn("w:sz"), str(sz))
    bottom.set(qn("w:space"), str(space)); bottom.set(qn("w:color"), color_hex)
    pBdr.append(bottom); pPr.append(pBdr)

def set_cell_shading(cell, fill_hex):
    tc = cell._tc; tcPr = tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"), "clear"); shd.set(qn("w:color"), "auto"); shd.set(qn("w:fill"), fill_hex)
    tcPr.append(shd)

def set_cell_margins(cell, top=80, bottom=80, left=115, right=115):
    tc = cell._tc; tcPr = tc.get_or_add_tcPr()
    tcMar = OxmlElement("w:tcMar")
    for side, val in [("top", top), ("bottom", bottom), ("left", left), ("right", right)]:
        el = OxmlElement(f"w:{side}"); el.set(qn("w:w"), str(val)); el.set(qn("w:type"), "dxa")
        tcMar.append(el)
    tcPr.append(tcMar)

def set_table_borders(table):
    tbl = table._tbl; tblPr = tbl.find(qn("w:tblPr"))
    if tblPr is None: tblPr = OxmlElement("w:tblPr"); tbl.insert(0, tblPr)
    tblBdr = OxmlElement("w:tblBorders")
    for side in ("top","left","bottom","right","insideH","insideV"):
        el = OxmlElement(f"w:{side}")
        el.set(qn("w:val"),"single"); el.set(qn("w:sz"),"4")
        el.set(qn("w:space"),"0"); el.set(qn("w:color"), BORDER_COLOR)
        tblBdr.append(el)
    tblPr.append(tblBdr)

def add_h1(doc, text):
    para = doc.add_paragraph()
    run = para.add_run(text)
    run.bold = True; run.font.name = "Arial"; run.font.size = Pt(18); run.font.color.rgb = DARK_GREEN
    set_para_spacing(para, before=400, after=200); add_bottom_border(para, GOLD, sz=12, space=4)

def add_h2(doc, text):
    para = doc.add_paragraph()
    run = para.add_run(text)
    run.bold = True; run.font.name = "Arial"; run.font.size = Pt(14); run.font.color.rgb = DARK_GREEN
    set_para_spacing(para, before=300, after=160); add_bottom_border(para, CREAM_LINE, sz=4, space=4)

def add_h3(doc, text):
    para = doc.add_paragraph()
    run = para.add_run(text)
    run.bold = True; run.font.name = "Arial"; run.font.size = Pt(12); run.font.color.rgb = DARK_GREEN
    set_para_spacing(para, before=220, after=120)

def add_body(doc, text):
    para = doc.add_paragraph()
    run = para.add_run(text)
    run.font.name = "Arial"; run.font.size = Pt(11); run.font.color.rgb = NEAR_BLACK
    set_para_spacing(para, before=0, after=120)

def add_bullet(doc, text, bold_prefix=None):
    para = doc.add_paragraph(style="List Bullet")
    if bold_prefix:
        r1 = para.add_run(bold_prefix); r1.bold = True
        r1.font.name = "Arial"; r1.font.size = Pt(11); r1.font.color.rgb = NEAR_BLACK
    run = para.add_run(text)
    run.font.name = "Arial"; run.font.size = Pt(11); run.font.color.rgb = NEAR_BLACK
    set_para_spacing(para, before=0, after=80)

def add_code(doc, text):
    para = doc.add_paragraph()
    run = para.add_run(text)
    run.font.name = "Courier New"; run.font.size = Pt(9); run.font.color.rgb = NEAR_BLACK
    set_para_spacing(para, before=60, after=60)
    pPr = para._p.get_or_add_pPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"),"clear"); shd.set(qn("w:color"),"auto"); shd.set(qn("w:fill"), CODE_BG)
    pPr.append(shd)

def add_table(doc, rows, col_widths=(1.75, 4.5)):
    table = doc.add_table(rows=0, cols=2); set_table_borders(table)
    for label, detail in rows:
        row = table.add_row(); left = row.cells[0]; right = row.cells[1]
        left.width = Inches(col_widths[0]); right.width = Inches(col_widths[1])
        set_cell_margins(left); set_cell_margins(right)
        if label == "__header__":
            set_cell_shading(left, TABLE_HDR_BG); set_cell_shading(right, TABLE_HDR_BG)
            lp = left.paragraphs[0]; lr = lp.add_run(detail[0])
            lr.bold = True; lr.font.name = "Arial"; lr.font.size = Pt(11); lr.font.color.rgb = WHITE
            rp = right.paragraphs[0]; rr = rp.add_run(detail[1])
            rr.bold = True; rr.font.name = "Arial"; rr.font.size = Pt(11); rr.font.color.rgb = WHITE
        else:
            lp = left.paragraphs[0]; lr = lp.add_run(label)
            lr.bold = True; lr.font.name = "Arial"; lr.font.size = Pt(11); lr.font.color.rgb = DARK_GREEN
            rp = right.paragraphs[0]; rr = rp.add_run(detail)
            rr.font.name = "Arial"; rr.font.size = Pt(11); rr.font.color.rgb = NEAR_BLACK
    doc.add_paragraph(); return table

def add_title_block(doc, line1, line2, subtitle):
    for line, after in [(line1, 60), (line2, 100)]:
        p = doc.add_paragraph(); p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        r = p.add_run(line); r.bold = True; r.font.name = "Arial"
        r.font.size = Pt(22); r.font.color.rgb = DARK_GREEN
        set_para_spacing(p, before=0, after=after)
    p3 = doc.add_paragraph(); p3.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r3 = p3.add_run(subtitle); r3.italic = True; r3.font.name = "Arial"
    r3.font.size = Pt(11); r3.font.color.rgb = NEAR_BLACK
    set_para_spacing(p3, before=0, after=300)


def build():
    doc = Document()
    for section in doc.sections:
        section.top_margin    = Inches(1.0)
        section.bottom_margin = Inches(1.0)
        section.left_margin   = Inches(1.1)
        section.right_margin  = Inches(1.1)

    add_title_block(doc,
        "CLIENT CALENDLY SETUP",
        "ONBOARDING GUIDE",
        "Step-by-step process for setting up scheduling on behalf of a client — done entirely by you"
    )

    add_h2(doc, "What's inside this document")
    add_table(doc, [
        ("__header__", ("Section", "Contents")),
        ("Part 1 — Before the Call",        "What to gather and prepare before the onboarding session"),
        ("Part 2 — Create the Account",     "Creating the Calendly account for the client from scratch"),
        ("Part 3 — Connect Google Calendar","Getting calendar authorization during the call — step by step"),
        ("Part 4 — Build the Event Type",   "Setting up the booking event (name, duration, availability)"),
        ("Part 5 — Get the Calendly URL",   "Finding the embed URL to drop into the client's website code"),
        ("Part 6 — Add to the Website",     "Swapping the placeholder in the HTML template with the real URL"),
        ("Part 7 — Store Credentials",      "Logging the client's login info securely in your credentials sheet"),
        ("Part 8 — Test Everything",        "Final checklist before closing the onboarding call"),
        ("Quick Reference",                 "One-page summary of all steps and placeholders"),
    ])

    add_h2(doc, "How to use this document")
    add_body(doc,
        "Run through this guide during the client onboarding call. Parts 1 and 7 happen off the call. "
        "Parts 2 through 6 and Part 8 are done live with the client — they do not need to touch anything "
        "except clicking 'Allow' when connecting Google Calendar. The whole process takes 15 to 20 minutes."
    )

    # ── PART 1 ──────────────────────────────────────────────────────────────
    add_h1(doc, "Part 1 — Before the Call")

    add_h2(doc, "Step 1 — Gather What You Need")
    add_body(doc,
        "Have the following ready before the onboarding call starts. You will need these "
        "to set up the account without interruptions."
    )
    add_table(doc, [
        ("__header__", ("Item", "Details")),
        ("Client's business email",
         "The email address you will use to create their Calendly account. "
         "Should be their professional email — e.g. mike@smithplumbing.com. "
         "Ask for this in your pre-call confirmation message."),
        ("Password to create",
         "Create a strong password on their behalf. Use a format like: BusinessName2025! "
         "You will store this in the Client Credentials Sheet after the call."),
        ("Their availability hours",
         "What days and times they want to accept bookings. e.g. Mon–Fri, 8am–4pm. "
         "Ask this in the pre-call intake form or confirm at the start of the call."),
        ("Their Google account email",
         "The Gmail address connected to their Google Calendar. Usually the same as their "
         "business email or a personal Gmail. Needed to connect the calendar."),
        ("The client's website HTML file open",
         "Have the client's HTML template open in your editor so you can paste the "
         "Calendly URL in immediately after getting it in Part 5."),
    ])

    add_h3(doc, "Pre-Call Message to Send the Client")
    add_code(doc,
        "Hi [Client Name] — before our call, can you send me:\n"
        "  1. Your business email address (the one you check most)\n"
        "  2. The Gmail address tied to your Google Calendar\n"
        "  3. What days/hours you want to accept estimate bookings\n\n"
        "That's it — I'll handle the rest on the call. See you at [time]."
    )

    # ── PART 2 ──────────────────────────────────────────────────────────────
    add_h1(doc, "Part 2 — Create the Calendly Account")

    add_h2(doc, "Step 2 — Set Up the Account (You Do This, Not the Client)")
    add_body(doc,
        "Create the account yourself before or at the start of the call. "
        "The client does not need to be watching for this step."
    )
    add_bullet(doc, "Go to calendly.com and click Sign Up.")
    add_bullet(doc, "Enter the client's business email address.")
    add_bullet(doc, "Create a password using your standard format — e.g. BusinessName2025!")
    add_bullet(doc, "Select the Free plan. One event type is all they need.")
    add_bullet(doc, "Skip any onboarding wizard prompts — you will configure everything manually.")
    add_bullet(doc, "You are now inside their Calendly account. Do not close this tab.")

    add_h3(doc, "Account Name and Profile")
    add_bullet(doc, "Set the display name to their business name — e.g. Smith Plumbing.")
    add_bullet(doc, "Set the Calendly URL slug to a clean version of their business name "
               "— e.g. calendly.com/smith-plumbing. This is what shows in the booking link.")
    add_bullet(doc, "Skip the profile photo for now unless you have one ready.")

    # ── PART 3 ──────────────────────────────────────────────────────────────
    add_h1(doc, "Part 3 — Connect Google Calendar")

    add_h2(doc, "Step 3 — Get Calendar Authorization (Do This Live on the Call)")
    add_body(doc,
        "This is the one step that requires the client to take an action. "
        "It takes about 60 seconds. Walk them through it verbally while screen sharing."
    )

    add_h3(doc, "Script for the Call")
    add_code(doc,
        '"[Client name], I need you to do one quick thing — I\'m going to share my screen\n'
        'so you can see what I\'m clicking. In about 30 seconds a Google login window is\n'
        'going to pop up on YOUR screen. All you need to do is log in with your Google\n'
        'account and click Allow. That connects your Google Calendar so bookings go\n'
        'straight to it. Ready?"'
    )

    add_h3(doc, "Step-by-Step — What You Click")
    add_bullet(doc, "In Calendly, go to Account Settings → Calendar Connections.")
    add_bullet(doc, "Click Connect next to Google Calendar.")
    add_bullet(doc, "A Google OAuth window opens — this appears on the client's screen if you are screen sharing. "
               "Direct them to log in with their Google account.")
    add_bullet(doc, "They click Allow on the permissions screen.")
    add_bullet(doc, "Calendly confirms the calendar is connected. You will see a green checkmark.")
    add_bullet(doc, "Confirm with the client: 'Perfect — your Google Calendar is connected. "
               "From now on, any booking that comes through your site will show up on your calendar automatically.'")

    add_h3(doc, "If They Have Multiple Google Calendars")
    add_body(doc,
        "In Calendly → Calendar Connections, you can choose which calendar to write new events to "
        "and which calendars to check for conflicts. Set 'Add to calendar' to their primary work calendar. "
        "Set 'Check for conflicts' to include any personal calendars they want blocked off."
    )

    # ── PART 4 ──────────────────────────────────────────────────────────────
    add_h1(doc, "Part 4 — Build the Event Type")

    add_h2(doc, "Step 4 — Create the Booking Event")
    add_body(doc,
        "On Calendly's free plan, one event type is allowed. Create it now. "
        "This is the event customers will see when they open the calendar on the client's website."
    )
    add_table(doc, [
        ("__header__", ("Field", "What to Set")),
        ("Event name",         "Free Estimate — or whatever the client calls it. Keep it short."),
        ("Duration",           "30 minutes is standard for a free estimate. 60 minutes for larger consults."),
        ("Location",           "Set to 'I will call my invitee' or 'In-person' depending on the business. "
                               "For tradespeople, 'In-person at your location' is typical."),
        ("Description",        "Optional — add a short note like 'We'll come to you, assess the job, and "
                               "give you a quote on the spot. No pressure.'"),
        ("Availability",       "Set to the client's working hours from Step 1. "
                               "e.g. Monday–Friday, 8:00am–4:00pm. Block weekends."),
        ("Minimum notice",     "Set to at least 2 hours so same-day bookings don't catch the client off guard. "
                               "4–24 hours is more common for field service businesses."),
        ("Buffer time",        "Add 30 minutes after each event so back-to-back bookings are impossible."),
        ("Max events per day", "Optional — set a limit if the client can only handle a certain number of estimates."),
    ])

    add_h3(doc, "Confirmation and Reminder Emails")
    add_body(doc,
        "Calendly sends automatic confirmation emails to the customer when they book, and reminder emails "
        "24 hours and 1 hour before the appointment. These are on by default — leave them on. "
        "Under Notifications, add the client's email as an additional notification recipient "
        "so they also get an email when someone books."
    )

    # ── PART 5 ──────────────────────────────────────────────────────────────
    add_h1(doc, "Part 5 — Get the Calendly URL")

    add_h2(doc, "Step 5 — Copy the Event Link")
    add_body(doc,
        "Once the event type is created, you need the direct booking URL to embed in the website. "
        "This is the URL that powers the calendar widget on the client's site."
    )
    add_bullet(doc, "From the Calendly dashboard, go to Event Types.")
    add_bullet(doc, "Click the Share button on the event you just created.")
    add_bullet(doc, "Copy the link. It will look like: calendly.com/smith-plumbing/free-estimate")
    add_bullet(doc, "This is the CLIENT_CALENDLY_URL you will paste into the website code.")

    add_h3(doc, "Also Note the Brand Color Hex")
    add_body(doc,
        "The calendar widget on the client's site passes a primary_color parameter to Calendly "
        "so the booking modal matches the site's color scheme. Note the hex code used in the "
        "client's CSS (the --accent variable in the HTML file) — you will need to paste it "
        "without the # symbol as the CLIENT_ACCENT_HEX placeholder."
    )
    add_code(doc,
        "Example:\n"
        "  CLIENT_CALENDLY_URL  = https://calendly.com/smith-plumbing/free-estimate\n"
        "  CLIENT_ACCENT_HEX   = 3d72a8   (no # symbol — just the 6 hex digits)"
    )

    # ── PART 6 ──────────────────────────────────────────────────────────────
    add_h1(doc, "Part 6 — Add the URL to the Website")

    add_h2(doc, "Step 6 — Swap the Placeholders in the HTML")
    add_body(doc,
        "Open the client's HTML file in your editor. Search for the two placeholders below "
        "and replace them with the values from Step 5. Both appear in the JavaScript block "
        "near the bottom of the file."
    )
    add_table(doc, [
        ("__header__", ("Placeholder", "Replace With")),
        ("[CLIENT_CALENDLY_URL]",
         "The full Calendly booking URL copied in Step 5. "
         "e.g. https://calendly.com/smith-plumbing/free-estimate"),
        ("[CLIENT_ACCENT_HEX]",
         "The 6-character hex code from the client's --accent CSS variable — without the # symbol. "
         "e.g. 3d72a8"),
    ])

    add_h3(doc, "Where to Find the Placeholders in the File")
    add_code(doc,
        "Search the HTML file for: CLIENT_CALENDLY_URL\n\n"
        "You will find it in a JavaScript block that looks like:\n\n"
        "  var CALENDLY = '[CLIENT_CALENDLY_URL]';\n"
        "  var HEX      = '[CLIENT_ACCENT_HEX]';\n\n"
        "Replace both values and save the file. Then redeploy the site to Netlify."
    )

    add_h3(doc, "Redeploy to Netlify")
    add_bullet(doc, "If Netlify is connected to GitHub: git add, commit, and push. The deploy triggers automatically.")
    add_bullet(doc, "If deploying manually: drag the updated site folder into the Netlify deploy dropzone.")
    add_bullet(doc, "Visit the live site and test the calendar before ending the call.")

    # ── PART 7 ──────────────────────────────────────────────────────────────
    add_h1(doc, "Part 7 — Store Credentials Securely")

    add_h2(doc, "Step 7 — Log Everything in the Client Credentials Sheet")
    add_body(doc,
        "Immediately after the call, open your Client Credentials Google Sheet and log the following. "
        "This sheet is your only record of how to access each client's Calendly account if they ever "
        "need changes. Never store passwords in email or chat."
    )
    add_table(doc, [
        ("__header__", ("Column", "What to Enter")),
        ("Client Name",             "First and last name of the primary contact"),
        ("Business Name",           "Full legal or trade name of the business"),
        ("Calendly Email",          "The email address used to create their Calendly account"),
        ("Calendly Password",       "The password you created — store this carefully"),
        ("Calendly URL",            "The full booking link — e.g. calendly.com/smith-plumbing/free-estimate"),
        ("Calendly Slug",           "The account slug — e.g. smith-plumbing (the part after calendly.com/)"),
        ("Google Calendar Connected","Yes or No — and note which Google account was connected"),
        ("Date Setup",              "The date the setup was completed"),
        ("Accent Hex",              "The hex code used for the calendar color (no #) — e.g. 3d72a8"),
        ("Notes",                   "Anything unusual about the setup — e.g. 'client uses shared Google calendar'"),
    ])

    add_h3(doc, "Keeping the Sheet Secure")
    add_bullet(doc, "Keep the Client Credentials Sheet in your own Google Drive — not shared with anyone.")
    add_bullet(doc, "Do not share client passwords via email, text, or Slack.")
    add_bullet(doc, "If a client ever needs their login, share it via a secure method — "
               "call them and read it out, or use a one-time secure link service like privnote.com.")

    # ── PART 8 ──────────────────────────────────────────────────────────────
    add_h1(doc, "Part 8 — Test Everything Before Ending the Call")

    add_h2(doc, "Step 8 — Final Checklist")
    add_body(doc,
        "Run through every item below before you wrap up the onboarding call. "
        "Do not skip this — it is much easier to catch a problem now than after the client has gone."
    )
    add_bullet(doc, "Open the live client website and scroll to the scheduling section.")
    add_bullet(doc, "Confirm the calendar loads and shows the current month.")
    add_bullet(doc, "Click a future date — confirm the Calendly modal opens with the correct event.")
    add_bullet(doc, "Confirm the time slots shown match the availability you configured.")
    add_bullet(doc, "Book a test appointment using your own email address.")
    add_bullet(doc, "Confirm the client receives a notification email for the test booking.")
    add_bullet(doc, "Confirm the booking appears on the client's Google Calendar.")
    add_bullet(doc, "Cancel the test appointment from Calendly's dashboard.")
    add_bullet(doc, "Confirm the client's credentials are logged in the Client Credentials Sheet.")
    add_bullet(doc, "Walk the client through what they will see when a real booking comes in.")

    add_h3(doc, "Script for Closing the Call")
    add_code(doc,
        '"Everything is set up. When a customer visits your site and picks a day on the calendar,\n'
        'they will book directly onto your Google Calendar — you will get a notification email\n'
        'and the appointment shows up just like any other meeting. You do not need to log into\n'
        'Calendly for anything. If you ever want to change your availability or block off time,\n'
        'just let me know and I will handle it. Any questions before we wrap up?"'
    )

    # ── QUICK REFERENCE ─────────────────────────────────────────────────────
    add_h1(doc, "Quick Reference — Complete Setup Checklist")

    add_h2(doc, "All Steps at a Glance")
    add_table(doc, [
        ("__header__", ("Step", "Action")),
        ("Before call",    "Gather client email, Google account, availability hours. Open their HTML file."),
        ("Step 1",         "Create Calendly account at calendly.com using client's business email."),
        ("Step 2",         "Set account display name and URL slug to their business name."),
        ("Step 3",         "Go to Calendar Connections → Connect Google Calendar."),
        ("Step 4",         "Client clicks Allow on Google OAuth popup (only thing client does)."),
        ("Step 5",         "Create event type: name, duration, location, availability, buffer time."),
        ("Step 6",         "Enable client email notifications for new bookings."),
        ("Step 7",         "Copy the booking URL from Event Types → Share."),
        ("Step 8",         "Paste URL into HTML file: replace [CLIENT_CALENDLY_URL]."),
        ("Step 9",         "Paste accent hex into HTML file: replace [CLIENT_ACCENT_HEX] (no # symbol)."),
        ("Step 10",        "Save file and redeploy site to Netlify."),
        ("Step 11",        "Test: click a date, book a test appointment, confirm on client's calendar."),
        ("Step 12",        "Cancel test appointment. Log all credentials in Client Credentials Sheet."),
    ])

    add_h3(doc, "Placeholder Reference")
    add_code(doc,
        "In the HTML file, search for and replace these two placeholders:\n\n"
        "  [CLIENT_CALENDLY_URL]   Full booking URL (e.g. https://calendly.com/smith-plumbing/free-estimate)\n"
        "  [CLIENT_ACCENT_HEX]    Brand color hex without # (e.g. 3d72a8)\n\n"
        "Both are in the <script> block near the bottom of the HTML file."
    )

    out = "/Users/creighbaby/LeadManagement/Client Calendly Setup — Onboarding Guide.docx"
    doc.save(out)
    print(f"Saved: {out}")

if __name__ == "__main__":
    build()
