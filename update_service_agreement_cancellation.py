"""
Updates Section 7 (Cancellation) of the Client Service Agreement to:
1. Clarify the website stays live — client paid for the build
2. Correct that lead data is in provider's Google account (will be exported + sent)
3. Add language about the form being updated to email-only forwarding as a courtesy
"""

from docx import Document
from docx.shared import Pt, RGBColor
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
import copy

DOC_PATH = '/Users/creighbaby/LeadManagement/docs/Client Service Agreement — Template.docx'

DARK_GREEN = RGBColor(0x1a, 0x3a, 0x2a)
NEAR_BLACK = RGBColor(0x1c, 0x1c, 0x1a)

def set_run_style(run, bold=False, color=NEAR_BLACK, size=11):
    run.bold = bold
    run.font.color.rgb = color
    run.font.size = Pt(size)
    run.font.name = 'Arial'

def set_para_spacing(para, before=0, after=120):
    pPr = para._p.get_or_add_pPr()
    spacing = OxmlElement('w:spacing')
    spacing.set(qn('w:before'), str(before))
    spacing.set(qn('w:after'),  str(after))
    pPr.append(spacing)

doc = Document(DOC_PATH)

# ── Locate the paragraphs to update ────────────────────────────────────────────
# We'll scan for the exact text of the three bullets we're replacing/adding to.

TARGET_WEBSITE  = 'Your website may be taken offline after the end of the final paid period'
TARGET_SHEET    = 'Your lead data in Google Sheets stays in your Google account'
TARGET_SECTION  = 'What Happens to Your Stuff When We Part Ways'

# New bullet texts
NEW_WEBSITE = (
    'Your website remains live and is yours to keep. '
    'The website build is a one-time service you have already paid for — '
    'canceling the monthly subscription does not take it down.'
)
NEW_SHEET = (
    'You will receive a full export of your lead data (CSV) at the time of cancellation. '
    'The Google Sheet will then be closed and your data deleted from my account.'
)
NEW_FORM = (
    'Your lead intake form will be updated to forward new submissions directly to your email '
    'as an offboarding courtesy — so you are not left with a broken form on your site. '
    'The full managed service (instant text alerts, organized lead log, monthly reports, '
    'and GBP audits) stops with the subscription.'
)

section_found = False
form_bullet_inserted = False

for i, para in enumerate(doc.paragraphs):
    text = para.text.strip()

    # Track when we enter the "What Happens" subsection
    if TARGET_SECTION in text:
        section_found = True

    # Replace the website bullet
    if section_found and TARGET_WEBSITE in text:
        para.clear()
        run = para.add_run(NEW_WEBSITE)
        set_run_style(run)
        set_para_spacing(para, before=0, after=80)

    # Replace the sheet bullet and insert the new form bullet right after
    if section_found and TARGET_SHEET in text:
        para.clear()
        run = para.add_run(NEW_SHEET)
        set_run_style(run)
        set_para_spacing(para, before=0, after=80)

        # Insert the new form bullet immediately after this paragraph
        # by cloning the paragraph's XML element and inserting a new one after it
        new_para = OxmlElement('w:p')

        # Copy paragraph properties (style — List Bullet) from the current para
        pPr_src = para._p.find(qn('w:pPr'))
        if pPr_src is not None:
            new_pPr = copy.deepcopy(pPr_src)
            new_para.append(new_pPr)

        # Add the run
        r = OxmlElement('w:r')
        rPr = OxmlElement('w:rPr')
        rFonts = OxmlElement('w:rFonts')
        rFonts.set(qn('w:ascii'), 'Arial')
        rFonts.set(qn('w:hAnsi'), 'Arial')
        rPr.append(rFonts)
        sz = OxmlElement('w:sz')
        sz.set(qn('w:val'), '22')  # 11pt = 22 half-points
        rPr.append(sz)
        szCs = OxmlElement('w:szCs')
        szCs.set(qn('w:val'), '22')
        rPr.append(szCs)
        color_el = OxmlElement('w:color')
        color_el.set(qn('w:val'), '1c1c1a')
        rPr.append(color_el)
        r.append(rPr)
        t = OxmlElement('w:t')
        t.set('{http://www.w3.org/XML/1998/namespace}space', 'preserve')
        t.text = NEW_FORM
        r.append(t)
        new_para.append(r)

        # Insert after the current paragraph in the document body
        para._p.addnext(new_para)
        form_bullet_inserted = True
        break

doc.save(DOC_PATH)
print('Section 7 updated successfully.')
if form_bullet_inserted:
    print('Form email-only offboarding bullet inserted.')
